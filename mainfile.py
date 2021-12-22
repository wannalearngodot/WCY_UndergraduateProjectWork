# -*- coding: utf-8 -*-
#########################################################通用库
import os
import sip
import math
import cv2 
import sys
import xlrd
import pymysql
import numpy as np
import matplotlib.pyplot as plt
import scipy
import copy as Xcopy

from pylab import *

from PyQt5 import QtGui
from PyQt5 import QtCore
from PyQt5 import QtWidgets
from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
from PyQt5.QtWidgets import QMessageBox
from PyQt5.QtWidgets import (QMainWindow, QWidget, QApplication)
#from PIL import Image, ExifTags
#和excel相关
import xlrd
import xlwt
import re
from xlutils.copy  import copy

#maskrcmnn检测函数
import detect_function

from tkinter import *
from tkinter import messagebox

from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5 import NavigationToolbar2QT as NavigationToolbar
# ################################################################################窗口导入区

#已确认
from Ui_MainWindow import Ui_MainWindow
from Ui_XPictureAnalyze import Ui_XPictureAnalyze
from Ui_SpineFusion import Ui_SpineFusion
from Ui_DlgPatientInfoInput import Ui_DlgPatientInfoInput
from Ui_DlgStatisticsAnalyzing import Ui_DlgStatisticsAnalyzing

#未确认
from YH_DataPreprocess import fRotationAngleCalWith4BkMarker
from YH_DataPreprocess import fRotationMatrixCal, fTransMatrixCal
from YH_DataPreprocess import getpoints, dataprocess_main, plot_contour_map
from YH_DataPreprocess import fPtMovetoCenterofIliaca
import YH_DataPreprocess 
from Ui_FootPressAnalyze import CFootPressAnalyze
from Ui_DataPreprocessing import CDataPreprocessing

from Ui_ResultShowFirst import CResultShowFirst
from Ui_ResultShowSecond import CResultShowSecond
from Ui_ResultShowThird import CResultShowThird
from Ui_Coronal_c_bend import CCoronal_c_bend
from Ui_Result_dynamic import CResult_dynamic
import YH_Result_dynamic

# #################################### 全局变量定义区
g_strAllSysDataPath="D:\灵犀脊柱形态评价系统\数据文件"  #全部数据文件所在完整路径
g_strDataFolderNameOutHospital="" # 当前出院被检者文件夹名，只包含患者名字的文件夹名
g_strPatientFullPath="" #当前出院被检者文件夹名+系统数据文件路径，
                        #即g_strAllSysDataPath+g_strPatientFullPath
g_fCenterPxyz=[0, 0, 0] #系统世界坐标原点坐标
g_fRXYZCam=[0, 0, 0] #相机三个方向倾斜角，弧度制
# 检测模式选择
# 0----"全标记点模式")
# 1----"无标记点模式")
# 2----"疗效对比模式" 
g_strCheckModeChosenValue="全标记点模式"
# 相机姿态矫正模式选择
# 0----"背景标记点坐标相机姿态矫正模式")
# 1----"三个旋转方向角相机姿态矫正模式")
g_strCamPosAdjModeChosenValue="背景标记点坐标相机姿态矫正模式"
g_fArrayBKPtCorEnd=0
#g_fArraySpineMarkerPtCorEnd=0
px=[]
py=[]
iNumPChose=0
g_fLineAngleForCobb=[0, 0, 0]
g_centerx=0
g_centery=0
g_R=0
g_fMainCStartPtScale=0.1
g_fMainCEndPtScale=0.9
fBKCamAngleXYZ=[0,0,0]
#X光图片文件路径
XimgName=''
# #################################### 全局变量定义区end

class MainForm(QMainWindow,Ui_MainWindow):    
    
    def __init__(self):               
        
        super(MainForm, self).__init__()
        print(__name__)
        # 主窗口初始化时实现主窗口布局
        self.setupUi(self)    
        self.setWindowState(Qt.WindowMaximized)
	 
		# ###################################
        # 主窗口初始化时实例化每个子窗口

        #已确认
        self.DlgPatientInfoInput = Ui_DlgPatientInfoInput()#患者信息录入
        self.XPictureAnalyze=Ui_XPictureAnalyze()
        self.SpineFusion=Ui_SpineFusion()
        self.DlgStatisticsAnalyzing=Ui_DlgStatisticsAnalyzing()
		#未确认
        self.Ui_FootPressAnalyze = CFootPressAnalyze()  
        self.Ui_DataPreprocessing = CDataPreprocessing()
        self.Ui_ResultShowFirst = CResultShowFirst()
        self.Ui_ResultShowSecond = CResultShowSecond()
        self.Ui_ResultShowThird = CResultShowThird()
        self.Ui_Coronal_c_bend = CCoronal_c_bend()
        self.Ui_Result_dynamic = CResult_dynamic()        
        # 主窗口控件布局，在主窗口的QSplitter里添加子窗口
        self.splitter.addWidget(self.DlgPatientInfoInput)  
        ####################################控件响应消息函数列表
        # QToolBox 的 QToolBoxButton 按钮切换时发出的信号；也是信号与槽的连接
        self.toolBox.currentChanged.connect(self.MainWndFn_MenuLinktoSubWnd)
		
        ######################患者信息录入窗口控件函数
		#点击选择入院数据按钮
        self.DlgPatientInfoInput.BN_ChooseDataOutHospital.clicked.connect(self.DlgFn_BNclick_InfoInput_ChooseDataOutHospital)#获取列表中鼠标所指的文件夹名称
        #点击选择出院数据按钮
        self.DlgPatientInfoInput.BN_ChooseDataInHospital.clicked.connect(self.DlgFn_BNclick_InfoInput_ChooseDataInHospital)#获取列表中鼠标所指的文件夹名称
        #点击读取数据并处理按钮触发函数链接
        self.DlgPatientInfoInput.BN_ReadDataAndProcess.clicked.connect(self.DlgFn_BNclick_InfoInput_ReadDataAndProcess)#获取列表中鼠标所指的文件夹名称

		#######################x光片处理窗口控件函数
        self.BN_ReadXImg.clicked.connect(self.DlgFn_BNclick_Ximg_ReadXImg)   #读入图像
        self.BN_FlipXImg.clicked.connect(self.DlgFn_BNclick_Ximg_FlipXImg)       #翻转图像
        self.BN_MarkScale.clicked.connect(self.DlgFn_BNclick_Ximg_MarkScale)    #标记比例尺
        self.CoronalMarkerIndex.currentIndexChanged.connect(self.DlgFn_IndexChange_Ximg_CoronalMarker) #显示冠状面提示信息
        self.PelvisMarkerIndex.currentIndexChanged.connect(self.DlgFn_IndexChange_Ximg_PelvisMarker) #显示骨盆提示信息
        self.SagittalMarkerIndex.currentIndexChanged.connect(self.DlgFn_IndexChange_Ximg_SagittalMarker) #显示矢状面提示信息
        self.BN_CoronalAddMarker.clicked.connect(self.DlgFn_BNclick_Ximg_CoronalAddMarker) #冠状面指标标记
        self.BN_CoronalClearMarker.clicked.connect(self.DlgFn_BNclick_Ximg_CoronalClearMarker) #冠状面标记清除
        self.BN_CoronalCalculateMarker.clicked.connect(self.DlgFn_BNclick_Ximg_CoronalCalculateMarker) #冠状面标记计算
        self.BN_PelvisAddMarker.clicked.connect(self.DlgFn_BNclick_Ximg_PelvisAddMarker) #骨盆指标标记
        self.BN_PelvisClearMarker.clicked.connect(self.DlgFn_BNclick_Ximg_PelvisClearMarker) #骨盆指标清除
        self.BN_PelvisCalculateMarker.clicked.connect(self.DlgFn_BNclick_Ximg_PelvisCalculateMarker) #骨盆指标计算
        self.BN_SagittalAddMarker.clicked.connect(self.DlgFn_BNclick_Ximg_SagittalAddMarker) #矢状面指标标记
        self.BN_SagittalClearMarker.clicked.connect(self.DlgFn_BNclick_Ximg_SagittalClearMarker) #矢状面标记清除
        self.BN_SagittalCalculateMarker.clicked.connect(self.DlgFn_BNclick_Ximg_SagittalCalculateMarker) #矢状面标记计算
        self.BN_AddSpineMarkerLn.clicked.connect(self.DlgFn_BNclick_Ximg_AddSpineMarkerLn) #添加脊柱标记线（一个十字）
        self.BN_AddSpineMarkerPt.clicked.connect(self.DlgFn_BNclick_Ximg_AddSpineMarkerPt) #添加脊柱标记点（中心点）
        self.BN_ClearSpineMarkerLn.clicked.connect(self.DlgFn_BNclick_Ximg_ClearSpineMarkerLn) #清除脊柱标记线
        self.BN_ClearSpineMarkerPt.clicked.connect(self.DlgFn_BNclick_Ximg_ClearSpineMarkerPt) #清除脊柱标记点
        self.BN_CalculateSpineMarkerPt.clicked.connect(self.DlgFn_BNclick_Ximg_CalculateSpineMarkerPt) #计算脊柱曲线
        self.BN_CalculateMorph.clicked.connect(self.DlgFn_BNclick_Ximg_CalculateMorph) #计算椎体楔形变
        self.BN_MaskrcnnDetect.clicked.connect(self.DlgFn_BNclick_Ximg_MaskrcnnDetect)  # 识别脊柱截断
        self.BN_MaskAdd.clicked.connect(self.DlgFn_BNclick_Ximg_MaskAdd) # 添加识别截断
        self.BN_MaskRemove.clicked.connect(self.DlgFn_BNclick_Ximg_MaskRemove)  # 去除识别截断
        self.BN_MidCalculate.clicked.connect(self.DlgFn_BNclick_Ximg_MidCalculate)  # 中线自动计算
        self.BN_MidAdd.clicked.connect(self.DlgFn_BNclick_Ximg_MidAdd)  # 添加中线
        self.BN_MidRemove.clicked.connect(self.DlgFn_BNclick_Ximg_MidRemove)  # 去除中线
        self.BN_EndCalculate.clicked.connect(self.DlgFn_BNclick_Ximg_EndCalculate)  # 边缘自动计算
        self.BN_EndAdd.clicked.connect(self.DlgFn_BNclick_Ximg_EndAdd)  # 添加边缘
        self.BN_EndRemove.clicked.connect(self.DlgFn_BNclick_Ximg_EndRemove)  # 去除边缘
        self.BN_PelvisCalculate.clicked.connect(self.DlgFn_BNclick_Ximg_PelvisCalculate)  # 边缘自动计算
        self.BN_PelvisAdd.clicked.connect(self.DlgFn_BNclick_Ximg_PelvisAdd)  # 添加边缘
        self.BN_PelvisRemove.clicked.connect(self.DlgFn_BNclick_Ximg_PelvisRemove)  # 去除边缘


       # self.XPictureAnalyze.BN_IndexCalculate.clicked.connect(self.DlgFn_BNclick_Ximg_IndexCalculate)
        self.BN_SaveImg.clicked.connect(self.DlgFn_BNclick_Ximg_SaveImg)
        self.XPictureAnalyze.label_imgView = myLabel(self.XPictureAnalyze.groupBox)
        self.XPictureAnalyze.label_imgView.clicked.connect(self.DlgFn_LBclick_Hint_Change)
        #self.label_imgView = MyLabel(self.centralWidget)
        self.SpineFusion.BN_Confirm.clicked.connect(self.WgtFn_BNclick_Ximg_FusionConfirm)
        self.SpineFusion.BN_Cancel.clicked.connect(self.WgtFn_BNclick_Ximg_FusionCancel)
		######################统计分析窗口控件映射函数
		#点击读取原始数据按钮
        self.DlgStatisticsAnalyzing.BN_ReadOrgData.clicked.connect(self.DlgFn_BNclick_Statis_ReadOrgData)#获取列表中鼠标所指的文件夹名称
		
		# ###########未确认，待用区
		#保存当前患者信息到数据库按钮
        self.DlgPatientInfoInput.BN_SaveCurrentInfoToDatabase.clicked.connect(self.save_data)#数据入库(被测者信息入库)
        #点击开始采集三维数据按钮
        self.DlgPatientInfoInput.BN_Start3DDataCapture.clicked.connect(self.DlgFn_BNclick_InfoInput_CallCaptureExe)#调用采集系统exe(启动采集界面)
        #启动分析按钮
        self.pushButton_Start_Data_Analyze.clicked.connect(self.DlgFn_BNclick_InfoInput_StartAnalyze)#计算指标(启动数据分析)
                        #点击检索历史数据按钮
        self.DlgPatientInfoInput.BN_IndexHistoryData.clicked.connect(self.DlgFn_BNclick_InfoInput_IndexHisRecord)
		
        self.DlgPatientInfoInput.RB_FullMarkerMode.clicked.connect(self.DlgFn_RBclick_InfoInput_ChooseBackFullMarkerMode)
        self.DlgPatientInfoInput.RB_NoMarkerMode.clicked.connect(self.DlgFn_RBclick_InfoInput_ChooseNoMarkerMode)
        self.DlgPatientInfoInput.RB_ContrastMode.clicked.connect(self.DlgFn_RBclick_InfoInput_ChooseContrastMode)
		
        self.DlgPatientInfoInput.RB_BackroundMarkerCorMode.clicked.connect(self.DlgFn_RBclick_InfoInput_ChooseBackroundMarkerCorMode)
        self.DlgPatientInfoInput.RB_RotationAngleMode.clicked.connect(self.DlgFn_RBclick_InfoInput_ChooseRotationAngleMode)


		# ################################### 实际操作区
        #已确认
		#显示系统信息
       	
        strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
        self.DlgPatientInfoInput.strMessageOut.setPlainText("欢迎使用灵犀脊柱评价系统！请选择待分析文件夹\n"+strtem)
		
        self.DlgPatientInfoInput.LE_StartScale.setText("0.1")
        self.DlgPatientInfoInput.LE_EndScale.setText("0.9")
		
        self.DlgPatientInfoInput.RB_FullMarkerMode.setChecked(True)
        self.DlgFn_RBclick_InfoInput_ChooseBackFullMarkerMode()
		#选择背景旋转角模式
        self.DlgPatientInfoInput.RB_BackroundMarkerCorMode.setChecked(False)
        self.DlgPatientInfoInput.RB_RotationAngleMode.setChecked(True)
        self.DlgFn_RBclick_InfoInput_ChooseRotationAngleMode()
        #self.DlgFn_RBclick_InfoInput_ChooseBackroundMarkerCorMode()
		#读取并显示当前系统目录中已有的数据文件
        self.HistoryDataList = self.DlgFn_InfoInput_GetFile(g_strAllSysDataPath)  
        #未确认
        self.Ui_Coronal_c_bend.CBendStart_EndPoint_pushButton.clicked.connect(self.CBendStart_EndPoint)
        self.Ui_Coronal_c_bend.Refresh_pushButton.clicked.connect(self.CBendRefresh)
        self.Ui_Coronal_c_bend.delet_pushButton.clicked.connect(self.CBenddelet)        
        self.pushButton_plot_original_points.clicked.connect(self.plot_original_points)#显示原始带背景图        
        plt.cla()#去掉之前画的图
	    # self.Ui_DataPreprocessing.pushButton_get_points.clicked.connect(getpoints)  # 按钮触发选择四个点
        # self.Ui_DataPreprocessing.pushButton_plot_Coronal_plane.clicked.connect(self.plot_Coronal_and_Sagittal_plane_with_BG)#点击显示矫正后的冠状面和矢状面
        # self.Ui_DataPreprocessing.pushButton_plot_Index_chart.clicked.connect(self.plot_Index_chart)
        # self.Ui_DataPreprocessing.pushButton_Correcting_the_spine_points.clicked.connect(self.Correcting_the_spine_points)  # 按钮触发选择四个点
        self.pushButton_get_points.clicked.connect(getpoints)  # 按钮触发选择四个点
        self.pushButton_Attitude_correction.clicked.connect(dataprocess_main)  # 姿态矫正#20190530去除按钮20190605重加按钮
        self.pushButton_plot_Coronal_plane.clicked.connect(self.plot_Coronal_and_Sagittal_plane_with_BG)#点击显示矫正后的冠状面和矢状面
        self.pushButton_plot_Index_chart.clicked.connect(self.plot_Index_chart)
        self.pushButton_Correcting_the_spine_points.clicked.connect(self.Correcting_the_spine_points)  # 按钮触发选择四个点
        self.pushButton_refreshfile.clicked.connect(self.refreshfile)
        self.pushButton_Redraw_clouds_point.clicked.connect(self.Redraw_clouds_point)#重绘等高线点云图
        self.pushButton_save_to_word.clicked.connect(self.save_to_word)
        self.DlgPatientInfoInput.checkBox_SetDefaultMode.stateChanged.connect(self.Changed_attitude)#坐姿站姿的切换
        #self.slider_distance.valueChanged.connect(lambda: self.on_change_func(self.slider_distance))
		
        
    #单选钮关联触发函数区
    # 统计分析窗口按钮回调函数区
    def DlgFn_BNclick_Statis_ReadOrgData(self):        

        #路径前加 r，读取的文件路径
        file_path = r'整体变直.xlsx'

        #文件路径的中文转码
 #       file_path = file_path.decode('utf-8')

        #获取数据
        data = xlrd.open_workbook(file_path)

        #获取sheet
        table = data.sheet_by_name('Sheet1')

        #获取总行数
        nrows = table.nrows
        #获取总列数
        ncols = table.ncols

        self.DlgStatisticsAnalyzing.OrgDatatableWidget.set

        #获取一行的数值，例如第5行
        #rowvalue = table.row_values(5)

        #获取一列的数值，例如第6列
        col_values = table.col_values(0)

        #获取一个单元格的数值，例如第5行第6列
        #cell_value = table.cell(5,6).value



    #################################
    def DlgFn_LBclick_Hint_Change(self):
        if self.XPictureAnalyze.label_imgView.stateflag==3 and self.XPictureAnalyze.label_imgView.count[3]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==1:
            self.BN_MarkScale.setText("标记比例尺（已标）")
            self.XPictureAnalyze.TE_Hint.setText("提示信息：比例尺标记完成")
        if self.XPictureAnalyze.label_imgView.stateflag==4 and self.XPictureAnalyze.label_imgView.count[4]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==5 and self.XPictureAnalyze.label_imgView.count[5]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==6 and self.XPictureAnalyze.label_imgView.count[6]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==7 and self.XPictureAnalyze.label_imgView.count[7]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==8 and self.XPictureAnalyze.label_imgView.count[8]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==9 and self.XPictureAnalyze.label_imgView.count[9]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==10 and self.XPictureAnalyze.label_imgView.count[10]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制终点直线（下方），以两直线夹角标记cobb角")
        if self.XPictureAnalyze.label_imgView.stateflag==24 and self.XPictureAnalyze.label_imgView.count[24]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记CSVL")
        if self.XPictureAnalyze.label_imgView.stateflag==22 and self.XPictureAnalyze.label_imgView.count[22]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记上方偏移最大的脊柱截断")
        if self.XPictureAnalyze.label_imgView.stateflag==23 and self.XPictureAnalyze.label_imgView.count[23]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记下方偏移最大的脊柱截断")
        if self.XPictureAnalyze.label_imgView.stateflag==17 and self.XPictureAnalyze.label_imgView.count[13]==0 and self.XPictureAnalyze.label_imgView.count[16]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记骶骨连线倾角")
        if self.XPictureAnalyze.label_imgView.stateflag==25 and self.XPictureAnalyze.label_imgView.count[25]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记L5")
        if self.XPictureAnalyze.label_imgView.stateflag==18 and self.XPictureAnalyze.label_imgView.count[18]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记骨垢终点")
        if self.XPictureAnalyze.label_imgView.stateflag==18 and self.XPictureAnalyze.label_imgView.count[18]==4:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记骨垢当前点")
        if self.XPictureAnalyze.label_imgView.stateflag==26 and self.XPictureAnalyze.label_imgView.count[26]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎中线(直线长度和与边缘长度一致)")
        if self.XPictureAnalyze.label_imgView.stateflag==26 and self.XPictureAnalyze.label_imgView.count[26]==1:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎右边缘(直线长度和与边缘长度一致)")
        if self.XPictureAnalyze.label_imgView.stateflag==26 and self.XPictureAnalyze.label_imgView.count[26]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个或者两个点，以点标记椎弓根（如果只有一个椎弓根，就只画一个点）")
        if self.XPictureAnalyze.label_imgView.stateflag==27 and self.XPictureAnalyze.label_imgView.count[27]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎中线(直线长度和与边缘长度一致)")
        if self.XPictureAnalyze.label_imgView.stateflag==27 and self.XPictureAnalyze.label_imgView.count[27]==1:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎右边缘(直线长度和与边缘长度一致)")
        if self.XPictureAnalyze.label_imgView.stateflag==27 and self.XPictureAnalyze.label_imgView.count[27]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个或者两个点，以点标记椎弓根（如果只有一个椎弓根，就只画一个点）")
        if self.XPictureAnalyze.label_imgView.stateflag==28 and self.XPictureAnalyze.label_imgView.count[28]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎中线(直线长度和与边缘长度一致)")
        if self.XPictureAnalyze.label_imgView.stateflag==28 and self.XPictureAnalyze.label_imgView.count[28]==1:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎右边缘(直线长度和与边缘长度一致)")
        if self.XPictureAnalyze.label_imgView.stateflag==28 and self.XPictureAnalyze.label_imgView.count[28]==2:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个或者两个点，以点标记椎弓根（如果只有一个椎弓根，就只画一个点）")
        if self.XPictureAnalyze.label_imgView.stateflag==30 and self.XPictureAnalyze.label_imgView.count[30]==0:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记骶椎中心线")
    def DlgFn_BNclick_Ximg_MarkScale(self):
        if self.XPictureAnalyze.label_imgView.scaleCount<=0:
           self.XPictureAnalyze.label_imgView.stateflag=1
           self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，两个端点落在X光片比例尺上相距50的两点上")
        else:
           self.XPictureAnalyze.label_imgView.ScaleClear()
           self.BN_MarkScale.setText("标记比例尺（未标）")
    
    def DlgFn_IndexChange_Ximg_CoronalMarker(self):
        index=self.CoronalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            pixmap = QPixmap(r"任意位置cobb角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #胸段cobb角
        if index == 2:
            pixmap = QPixmap(r"任意位置cobb角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #胸腰段cobb角
        if index == 3:
            pixmap = QPixmap(r"任意位置cobb角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #腰段cobb角
        if index == 4:
            pixmap = QPixmap(r"任意位置cobb角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #锁骨连线倾角
        if index == 5:
            pixmap = QPixmap(r"锁骨连线倾角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #T1倾角
        if index == 6:
            pixmap = QPixmap(r"T1倾角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #顶椎上弯偏移距离
        if index == 7:
            pixmap = QPixmap(r"顶椎上弯偏移距离.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #底椎下弯偏移距离
        if index == 8:
            pixmap = QPixmap(r"底椎下弯偏移距离.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #冠状面失衡度
        if index == 9:
            pixmap = QPixmap(r"冠状面失衡度.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #顶椎旋转
        if index == 10:
            pixmap = QPixmap(r"顶椎旋转.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #上端椎旋转
        if index == 11:
            pixmap = QPixmap(r"顶椎旋转.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #下端椎旋转
        if index == 12:
            pixmap = QPixmap(r"顶椎旋转.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #椎体融合
        if index == 13:
            pixmap = QPixmap(r"椎体融合.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
    
    def DlgFn_IndexChange_Ximg_PelvisMarker(self):
        index=self.PelvisMarkerIndex.currentIndex()
        #股骨头连线倾角
        if index == 1:
            pixmap = QPixmap(r"股骨头连线倾角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #髂骨连线倾角
        if index == 2:
            pixmap = QPixmap(r"髂骨连线倾角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #坐骨连线倾角
        if index == 3:
            pixmap = QPixmap(r"坐骨连线倾角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #骶骨连线倾角
        if index == 4:
            pixmap = QPixmap(r"骶骨连线倾角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #骶骨连线相对倾角
        if index == 5:
            pixmap = QPixmap(r"骶骨相对倾斜角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #Risser参数
        if index == 6:
            pixmap = QPixmap(r"Risser参数.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #腰骶角
        if index == 7:
            pixmap = QPixmap(r"腰骶角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        
    def DlgFn_IndexChange_Ximg_SagittalMarker(self):
        index=self.SagittalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            pixmap = QPixmap(r"任意位置cobb角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #腰椎前凸角
        if index == 2:
            pixmap = QPixmap(r"腰椎前凸角.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
        #矢状面平衡
        if index == 3:
            pixmap = QPixmap(r"矢状面平衡.jpg")
            scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
            self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
    
    def DlgFn_BNclick_Ximg_CoronalAddMarker(self):
        index=self.CoronalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            self.XPictureAnalyze.label_imgView.stateflag=3
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制起点直线（上方）")
        #胸段cobb角
        if index == 2:
            self.XPictureAnalyze.label_imgView.stateflag=4
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制起点直线（上方）")
        #胸腰段cobb角
        if index == 3:
            self.XPictureAnalyze.label_imgView.stateflag=5
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制起点直线（上方）")
        #腰段cobb角
        if index == 4:
            self.XPictureAnalyze.label_imgView.stateflag=6
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制起点直线（上方）")
        #锁骨连线倾角
        if index == 5:
            self.XPictureAnalyze.label_imgView.stateflag=19
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记锁骨")
        #T1倾角
        if index == 6:
            self.XPictureAnalyze.label_imgView.stateflag=20
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记T1")
        #顶椎上弯偏移距离
        if index == 7:
            self.XPictureAnalyze.label_imgView.stateflag=22
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记C7PL")
        #底椎下弯偏移距离
        if index == 8:
            self.XPictureAnalyze.label_imgView.stateflag=23
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记CSVL")
        #冠状面失衡度
        if index == 9:
            self.XPictureAnalyze.label_imgView.stateflag=24
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记C7PL")
        #顶椎旋转
        if index == 10:
            self.XPictureAnalyze.label_imgView.stateflag=26
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎左边缘(直线长度和与边缘长度一致)")
        #上端椎旋转
        if index == 11:
            self.XPictureAnalyze.label_imgView.stateflag=27
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎左边缘(直线长度和与边缘长度一致)")
        #下端椎旋转
        if index == 12:
            self.XPictureAnalyze.label_imgView.stateflag=28
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记顶椎左边缘(直线长度和与边缘长度一致)")
        #椎体融合
        if index == 13:
            self.XPictureAnalyze.TE_Hint.setText("标记提示：观察X光片，在弹出的对话框内选择融合的锥体")
            self.SpineFusion.show()
    
    def DlgFn_BNclick_Ximg_PelvisAddMarker(self):
        index=self.PelvisMarkerIndex.currentIndex()
        #股骨头连线倾角
        if index == 1:
            self.XPictureAnalyze.label_imgView.stateflag=13
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记股骨头")
        #髂骨连线倾角
        if index == 2:
            self.XPictureAnalyze.label_imgView.stateflag=14
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记髂骨顶点")
        #坐骨连线倾角
        if index == 3:
            self.XPictureAnalyze.label_imgView.stateflag=15
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记坐骨底点")
        #骶骨连线倾角
        if index == 4:
            self.XPictureAnalyze.label_imgView.stateflag=16
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记骶骨连线倾角")
        #骶骨连线相对倾角
        if index == 5:
            if self.XPictureAnalyze.label_imgView.count[13]<1:
              self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记股骨头")
              self.XPictureAnalyze.label_imgView.stateflag=17
            elif self.XPictureAnalyze.label_imgView.count[16]<1:
              self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点标记骶骨连线倾角")
              self.XPictureAnalyze.label_imgView.stateflag=17
            else: 
               self.XPictureAnalyze.TE_Hint.setText("股骨头和骶骨的标记已完成，请点击指标计算按钮进行计算") 
        #Risser参数
        if index == 6:
            self.XPictureAnalyze.label_imgView.stateflag=18
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一个点，用这个点标记骨垢起点")
        #腰骶角
        if index == 7:
               self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用两个端点连接腰5和腰4中心或连接腰3和腰4中心")
               self.XPictureAnalyze.label_imgView.stateflag=30 
    
    def DlgFn_BNclick_Ximg_SagittalAddMarker(self):
        index=self.SagittalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            self.XPictureAnalyze.label_imgView.stateflag=9
            self.XPictureAnalyze.TE_Hint.setText("标记提示：请绘制起点直线（上方）")
        #腰椎前凸角
        if index == 2:
            self.XPictureAnalyze.label_imgView.stateflag=21
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画一条直线，用这条直线标记腰椎前凸角（垂线）")
        #矢状面平衡
        if index == 3:
            self.XPictureAnalyze.label_imgView.stateflag=25
            self.XPictureAnalyze.TE_Hint.setText("标记提示：画两个点，用这一个点标记C7")
    
    def DlgFn_BNclick_Ximg_AddSpineMarkerLn(self):
        if self.XPictureAnalyze.label_imgView.stateflag!=0:
         self.XPictureAnalyze.label_imgView.stateflag=0
         self.BN_AddSpineMarkerLn.setText("添加脊柱标记线")
         self.BN_AddSpineMarkerPt.setText("添加脊柱标记点")
         self.XPictureAnalyze.TE_Hint.setText("标记提示：已停止添加脊柱标记线")
        else :
         self.XPictureAnalyze.label_imgView.stateflag= 2
         self.BN_AddSpineMarkerLn.setText("停止添加标记线")  
         self.XPictureAnalyze.TE_Hint.setText("标记提示：在每个脊柱截断用两条直线标记它的对角线，从上到下依次标记颈7到腰5")
         pixmap = QPixmap(r"椎体楔形变.jpg")
         scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
         self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
         
         
    def DlgFn_BNclick_Ximg_AddSpineMarkerPt(self):
        if self.XPictureAnalyze.label_imgView.stateflag!=0:
         self.XPictureAnalyze.label_imgView.stateflag=0
         self.BN_AddSpineMarkerPt.setText("添加脊柱标记点")
         self.BN_AddSpineMarkerLn.setText("添加脊柱标记线")
         self.XPictureAnalyze.TE_Hint.setText("标记提示：已停止添加脊柱标记点")
        else :
         self.XPictureAnalyze.label_imgView.stateflag= 29
         self.BN_AddSpineMarkerPt.setText("停止添加标记点")  
         self.XPictureAnalyze.TE_Hint.setText("标记提示：在每个脊柱截断用一个点标记它的截断中心，从上到下依次标记颈7到腰5")
         pixmap = QPixmap(r"脊柱标记点.jpg")
         scaredPixmap = pixmap.scaled(400, 400, aspectRatioMode=Qt.KeepAspectRatio) #等比例缩放
         self.XPictureAnalyze.GV_Hint.setPixmap(scaredPixmap)
    
    def DlgFn_BNclick_Ximg_ClearSpineMarkerLn(self):
        self.XPictureAnalyze.label_imgView.MorphClear()
        self.BN_AddSpineMarkerLn.setText("添加脊柱标记线")
        
    def DlgFn_BNclick_Ximg_ClearSpineMarkerPt(self):
        self.XPictureAnalyze.label_imgView.PointClear()
        self.BN_AddSpineMarkerPt.setText("添加脊柱标记点")

    
    def DlgFn_BNclick_Ximg_CoronalClearMarker(self):
        index=self.CoronalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            self.XPictureAnalyze.label_imgView.LineClear()
            self.XPictureAnalyze.LE_RandomCobbResult.clear()
        #胸段cobb角
        if index == 2:
            self.XPictureAnalyze.label_imgView.CoronalCobb1Clear()
            self.XPictureAnalyze.LE_CoronalCobb1Result.clear()
        #胸腰段cobb角
        if index == 3:
            self.XPictureAnalyze.label_imgView.CoronalCobb2Clear()
            self.XPictureAnalyze.LE_CoronalCobb2Result.clear()
        #腰段cobb角
        if index == 4:
            self.XPictureAnalyze.label_imgView.CoronalCobb3Clear()
            self.XPictureAnalyze.LE_CoronalCobb3Result.clear()
        #锁骨连线倾角
        if index == 5:
            self.XPictureAnalyze.label_imgView.CollarboneClear()
            self.XPictureAnalyze.LE_CollarboneResult.clear()
        #T1倾角
        if index == 6:
            self.XPictureAnalyze.label_imgView.T1Clear()
            self.XPictureAnalyze.LE_T1Result.clear()
        #顶椎上弯偏移距离
        if index == 7:
            self.XPictureAnalyze.label_imgView.TVMUperClear()
            self.XPictureAnalyze.LE_TVMUperResult.clear()
        #底椎下弯偏移距离
        if index == 8:
            self.XPictureAnalyze.label_imgView.TVMLowerClear()
            self.XPictureAnalyze.LE_TVMLowerResult.clear()        
        #冠状面失衡度
        if index == 9:
            self.XPictureAnalyze.label_imgView.CoronalClear()
            self.XPictureAnalyze.LE_CoronalResult.clear()
        #顶椎旋转
        if index == 10:
            self.XPictureAnalyze.label_imgView.TVSClear()
            self.XPictureAnalyze.LE_TVSResult.clear()
        #上端椎旋转
        if index == 11:
            self.XPictureAnalyze.label_imgView.UVSClear()
            self.XPictureAnalyze.LE_UVSResult.clear()
        #下端椎旋转
        if index == 12:
            self.XPictureAnalyze.label_imgView.LVSClear()
            self.XPictureAnalyze.LE_LVSResult.clear()
        #椎体融合
        if index == 13:
            self.XPictureAnalyze.LE_FusionNumResult.clear()
            self.XPictureAnalyze.LE_Fusion1Result.clear()
            self.XPictureAnalyze.LE_Fusion2Result.clear()
            self.XPictureAnalyze.LE_Fusion3Result.clear()
    
    def DlgFn_BNclick_Ximg_PelvisClearMarker(self):
        index=self.PelvisMarkerIndex.currentIndex()
        #股骨头连线倾角
        if index == 1:
            self.XPictureAnalyze.label_imgView.FemurClear()
            self.XPictureAnalyze.LE_FemurAngleResult.clear()
        #髂骨连线倾角
        if index == 2:
            self.XPictureAnalyze.label_imgView.IliumClear()
            self.XPictureAnalyze.LE_IliumResult.clear()
        #坐骨连线倾角
        if index == 3:
            self.XPictureAnalyze.label_imgView.IschiumClear()
            self.XPictureAnalyze.LE_IschiumResult.clear()
        #骶骨连线倾角
        if index == 4:
            self.XPictureAnalyze.label_imgView.SacrumClear()
            self.XPictureAnalyze.LE_SacrumResult.clear()
        #骶骨连线相对倾角
        if index == 5:
            self.XPictureAnalyze.label_imgView.RelSacrumClear()
            self.XPictureAnalyze.LE_RelSacrumResult.clear()
        #Risser参数
        if index == 6:
            self.XPictureAnalyze.label_imgView.RsClear()
            self.XPictureAnalyze.LE_RsResult.clear()
            self.XPictureAnalyze.label_imgView.CalRs=1
        #腰骶角
        if index == 7:
            self.XPictureAnalyze.label_imgView.RelL5Clear()
            self.XPictureAnalyze.LE_RelL5Result.clear()
            self.XPictureAnalyze.LE_SacrumResult.clear()
    
    def DlgFn_BNclick_Ximg_SagittalClearMarker(self):
        index=self.SagittalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            self.XPictureAnalyze.label_imgView.SagittalCobb1Clear()
            self.XPictureAnalyze.LE_SagittalCobb1Result.clear()
        #腰椎前凸角
        if index == 2:
            self.XPictureAnalyze.label_imgView.LumbarClear()
            self.XPictureAnalyze.LE_LumbarResult.clear()
        #矢状面平衡
        if index == 3:
            self.XPictureAnalyze.label_imgView.SagittalClear()
            self.XPictureAnalyze.LE_SagittalResult.clear()
    
    def DlgFn_BNclick_Ximg_CoronalCalculateMarker(self):
        index=self.CoronalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            self.XPictureAnalyze.label_imgView.CalculateAngle()
            if self.XPictureAnalyze.label_imgView.count[3]>=2:
               self.XPictureAnalyze.LE_RandomCobbResult.setText(str(self.XPictureAnalyze.label_imgView.result[0]))
               self.XPictureAnalyze.LE_RandomCobbResult.setCursorPosition(0)
        #胸段cobb角
        if index == 2:
            self.XPictureAnalyze.label_imgView.CalculateCoronalCobb1()
            if self.XPictureAnalyze.label_imgView.count[4]>=2:
               self.XPictureAnalyze.LE_CoronalCobb1Result.setText(str(self.XPictureAnalyze.label_imgView.result[1]))
               self.XPictureAnalyze.LE_CoronalCobb1Result.setCursorPosition(0)
        #胸腰段cobb角
        if index == 3:
            self.XPictureAnalyze.label_imgView.CalculateCoronalCobb2()
            if self.XPictureAnalyze.label_imgView.count[5]>=2:
               self.XPictureAnalyze.LE_CoronalCobb2Result.setText(str(self.XPictureAnalyze.label_imgView.result[2]))
               self.XPictureAnalyze.LE_CoronalCobb2Result.setCursorPosition(0)
        #腰段cobb角
        if index == 4:
            self.XPictureAnalyze.label_imgView.CalculateCoronalCobb3()
            if self.XPictureAnalyze.label_imgView.count[6]>=2:
               self.XPictureAnalyze.LE_CoronalCobb3Result.setText(str(self.XPictureAnalyze.label_imgView.result[3]))
               self.XPictureAnalyze.LE_CoronalCobb3Result.setCursorPosition(0)
        #锁骨连线倾角
        if index == 5:
            self.XPictureAnalyze.label_imgView.CalculateCollarbone()
            if self.XPictureAnalyze.label_imgView.count[19]>=1:
               self.XPictureAnalyze.LE_CollarboneResult.setText(str(self.XPictureAnalyze.label_imgView.result[16]))
               self.XPictureAnalyze.LE_CollarboneResult.setCursorPosition(0) 
        #T1倾角
        if index == 6:
            self.XPictureAnalyze.label_imgView.CalculateT1()
            if self.XPictureAnalyze.label_imgView.count[20]>=1:
               self.XPictureAnalyze.LE_T1Result.setText(str(self.XPictureAnalyze.label_imgView.result[17]))
               self.XPictureAnalyze.LE_T1Result.setCursorPosition(0)
        #顶椎上弯偏移距离
        if index == 7:
            if self.XPictureAnalyze.label_imgView.scaleCount>=1:
               self.XPictureAnalyze.label_imgView.CalculateTVMUper()
               if self.XPictureAnalyze.label_imgView.count[22]>=4:
                  self.XPictureAnalyze.LE_TVMUperResult.setText(str(self.XPictureAnalyze.label_imgView.result[19]))
                  self.XPictureAnalyze.LE_TVMUperResult.setCursorPosition(0)
            else:
                self.XPictureAnalyze.TE_Hint.setText("提示信息：请先标记比例尺")
        #底椎下弯偏移距离
        if index == 8:
            if self.XPictureAnalyze.label_imgView.scaleCount>=1:
               self.XPictureAnalyze.label_imgView.CalculateTVMLower()
               if self.XPictureAnalyze.label_imgView.count[23]>=4:
                  self.XPictureAnalyze.LE_TVMLowerResult.setText(str(self.XPictureAnalyze.label_imgView.result[20]))
                  self.XPictureAnalyze.LE_TVMLowerResult.setCursorPosition(0)
            else:
                self.XPictureAnalyze.TE_Hint.setText("提示信息：请先标记比例尺")
        #冠状面失衡度
        if index == 9:
            if self.XPictureAnalyze.label_imgView.scaleCount>=1:
               self.XPictureAnalyze.label_imgView.CalculateCoronal()
               if self.XPictureAnalyze.label_imgView.count[24]>=4:
                  self.XPictureAnalyze.LE_CoronalResult.setText(str(self.XPictureAnalyze.label_imgView.result[21]))
                  self.XPictureAnalyze.LE_CoronalResult.setCursorPosition(0)
            else:
                self.XPictureAnalyze.TE_Hint.setText("提示信息：请先标记比例尺")
        #顶椎旋转
        if index == 10:
            self.XPictureAnalyze.label_imgView.CalculateTVS()
            if self.XPictureAnalyze.label_imgView.count[26]>=4:
                self.XPictureAnalyze.LE_TVSResult.setText("第"+str(self.XPictureAnalyze.label_imgView.result[23])+"级")
                self.XPictureAnalyze.LE_TVSResult.setCursorPosition(0)
        #上端椎旋转
        if index == 11:
            self.XPictureAnalyze.label_imgView.CalculateUVS()
            if self.XPictureAnalyze.label_imgView.count[27]>=4:
                self.XPictureAnalyze.LE_UVSResult.setText("第"+str(self.XPictureAnalyze.label_imgView.result[24])+"级")
                self.XPictureAnalyze.LE_UVSResult.setCursorPosition(0) 
        #下端椎旋转
        if index == 12:
            self.XPictureAnalyze.label_imgView.CalculateLVS()
            if self.XPictureAnalyze.label_imgView.count[28]>=4:
                self.XPictureAnalyze.LE_LVSResult.setText("第"+str(self.XPictureAnalyze.label_imgView.result[25])+"级")
                self.XPictureAnalyze.LE_LVSResult.setCursorPosition(0)
    
    def DlgFn_BNclick_Ximg_PelvisCalculateMarker(self):
        index=self.PelvisMarkerIndex.currentIndex()
        #股骨头连线倾角
        if index == 1:
            self.XPictureAnalyze.label_imgView.CalculateFemur()
            if self.XPictureAnalyze.label_imgView.count[13]>=1:
               self.XPictureAnalyze.LE_FemurAngleResult.setText(str(self.XPictureAnalyze.label_imgView.result[10],))
               self.XPictureAnalyze.LE_FemurAngleResult.setCursorPosition(0)
        #髂骨连线倾角
        if index == 2:
            self.XPictureAnalyze.label_imgView.CalculateIlium()
            if self.XPictureAnalyze.label_imgView.count[14]>=1:
               self.XPictureAnalyze.LE_IliumResult.setText(str(self.XPictureAnalyze.label_imgView.result[11]))
               self.XPictureAnalyze.LE_IliumResult.setCursorPosition(0)
        #坐骨连线倾角
        if index == 3:
            self.XPictureAnalyze.label_imgView.CalculateIschium()
            if self.XPictureAnalyze.label_imgView.count[15]>=1:
               self.XPictureAnalyze.LE_IschiumResult.setText(str(self.XPictureAnalyze.label_imgView.result[12]))
               self.XPictureAnalyze.LE_IschiumResult.setCursorPosition(0)
        #骶骨连线倾角
        if index == 4:
            self.XPictureAnalyze.label_imgView.CalculateSacrum()
            if self.XPictureAnalyze.label_imgView.count[16]>=1:
               self.XPictureAnalyze.LE_SacrumResult.setText(str(self.XPictureAnalyze.label_imgView.result[13]))
               self.XPictureAnalyze.LE_SacrumResult.setCursorPosition(0)   
        #骶骨连线相对倾角
        if index == 5:
            self.XPictureAnalyze.label_imgView.CalculateRelSacrumCobb()
            if self.XPictureAnalyze.label_imgView.count[13]>=1 and self.XPictureAnalyze.label_imgView.count[16]>=1:
               self.XPictureAnalyze.LE_FemurAngleResult.setText(str(self.XPictureAnalyze.label_imgView.result[10]))
               self.XPictureAnalyze.LE_SacrumResult.setText(str(self.XPictureAnalyze.label_imgView.result[13]))
               self.XPictureAnalyze.LE_RelSacrumResult.setText(str(self.XPictureAnalyze.label_imgView.result[14]))
               self.XPictureAnalyze.LE_FemurAngleResult.setCursorPosition(0)
               self.XPictureAnalyze.LE_SacrumResult.setCursorPosition(0)
               self.XPictureAnalyze.LE_RelSacrumResult.setCursorPosition(0) 
        #Risser参数
        if index == 6:
            self.XPictureAnalyze.label_imgView.CalculateRs()
            if self.XPictureAnalyze.label_imgView.CalRs==0 :
                self.XPictureAnalyze.TE_Hint.setText("标记提示：三点共线，请重新绘制")
            else:
                if self.XPictureAnalyze.label_imgView.count[18]>=6:
                   result=self.XPictureAnalyze.label_imgView.result[15]
                   if result<0.25:
                       result_str='1级'
                   elif result< 0.5:
                       result_str='2级'
                   elif result< 0.75:
                       result_str='3级' 
                   elif result< 0.95:
                       result_str='4级' 
                   else:
                       result_str='5级' 
                   self.XPictureAnalyze.LE_RsResult.setText(result_str+'('+str(self.XPictureAnalyze.label_imgView.result[15])+')')
                   self.XPictureAnalyze.LE_RsResult.setCursorPosition(0)
        #腰骶角
        if index == 7:
            self.XPictureAnalyze.label_imgView.CalculateRelL5()
            if self.XPictureAnalyze.label_imgView.count[30]>=2:
               self.XPictureAnalyze.LE_RelL5Result.setText(str(self.XPictureAnalyze.label_imgView.result[26]))
               self.XPictureAnalyze.LE_RelL5Result.setCursorPosition(0) 
    
    def DlgFn_BNclick_Ximg_SagittalCalculateMarker(self):
        index=self.SagittalMarkerIndex.currentIndex()
        #任意位置cobb角
        if index == 1:
            self.XPictureAnalyze.label_imgView.CalculateSagittalCobb1()
            if self.XPictureAnalyze.label_imgView.count[9]>=2:
               self.XPictureAnalyze.LE_SagittalCobb1Result.setText(str(self.XPictureAnalyze.label_imgView.result[6]))
               self.XPictureAnalyze.LE_SagittalCobb1Result.setCursorPosition(0)
        #腰椎前凸角
        if index == 2:
            self.XPictureAnalyze.label_imgView.CalculateLumbar()
            if self.XPictureAnalyze.label_imgView.count[21]>=1:
                self.XPictureAnalyze.LE_LumbarResult.setText(str(self.XPictureAnalyze.label_imgView.result[18]))
                self.XPictureAnalyze.LE_LumbarResult.setCursorPosition(0)
        #矢状面平衡
        if index == 3:
            if self.XPictureAnalyze.label_imgView.scaleCount>=1:
               self.XPictureAnalyze.label_imgView.CalculateSagittal()
               if self.XPictureAnalyze.label_imgView.count[25]>=4:
                  self.XPictureAnalyze.LE_SagittalResult.setText(str(self.XPictureAnalyze.label_imgView.result[22]))
                  self.XPictureAnalyze.LE_SagittalResult.setCursorPosition(0)
            else:
                self.XPictureAnalyze.TE_Hint.setText("提示信息：请先标记比例尺")
    
    
    def DlgFn_BNclick_Ximg_CalculateSpineMarkerPt(self):
        self.XPictureAnalyze.label_imgView.CalculatePoint()

            

    def DlgFn_BNclick_Ximg_CalculateMorph(self):
        notpass=['颈7', '胸1', '胸2','胸3','胸4','胸5','胸6','胸7','胸8','胸9','胸10','胸11','胸12','腰1','腰2','腰3','腰4','腰5']
        if self.XPictureAnalyze.label_imgView.scaleCount>=1:
           self.XPictureAnalyze.label_imgView.CalculateMorph()
           if self.XPictureAnalyze.label_imgView.groupcount>=18:
               result_str='不合格的截断：'
               for i in range(0, len(self.XPictureAnalyze.label_imgView.morphResult)):
                   result_str=result_str+notpass[self.XPictureAnalyze.label_imgView.morphResult[i]]+' '
               result_str=result_str+'需要结合临床判断的截断：'
               for i in range(0, len(self.XPictureAnalyze.label_imgView.morphQues)):
                   result_str=result_str+notpass[self.XPictureAnalyze.label_imgView.morphQues[i]]+' '
               self.XPictureAnalyze.LE_MorphLocResult.setText(result_str)
               self.XPictureAnalyze.LE_MorphResult.setText(str(self.XPictureAnalyze.label_imgView.morphNum))
               self.XPictureAnalyze.LE_MorphResult.setCursorPosition(0)
               self.XPictureAnalyze.LE_MorphLocResult.setCursorPosition(0)
        else:
                self.XPictureAnalyze.TE_Hint.setText("提示信息：请先标记比例尺")

    #函数功能：返回点集中第一个点的y坐标
    def takeFirstY(self, elem):
        return elem[0][1]

    # 函数功能：返回y坐标
    def takeY(self, elem):
            return elem[1]

    # 函数功能：返回绝对值
    def takeAbs(self, elem):
            return abs(elem)

    def DlgFn_BNclick_Ximg_MaskrcnnDetect(self):
        #设置模型和权重路径
        ROOT_DIR=os.getcwd()
        MODEL_DIR = ROOT_DIR+"\logs\manchu20200526T2143"
        COCO_MODEL_PATH = MODEL_DIR+"\mask_rcnn_manchu_0020.h5"  # （自己模型的文件名称）
        if self.XPictureAnalyze.label_imgView.ImgCount>=1:
            self.XPictureAnalyze.TE_Hint.setText("提示信息：已开始识别脊柱截断，请耐心等待计算完成")
            self.XPictureAnalyze.label_imgView.MaskClear()
            #获取缩放后的图片
            img=self.XPictureAnalyze.label_imgView.rsimg
            #调用自动识别函数
            vert_set, rect_set, box_set, list_center = detect_function.mask_detect(img, MODEL_DIR, COCO_MODEL_PATH)
            #将vert_set按takeFirstY的返回值排序
            vert_set.sort(key=self.takeFirstY)
            #将box_set按takeFirstY的返回值排序（
            box_set.sort(key=self.takeFirstY)
            #将list_center按takeY的返回值排序
            list_center.sort(key=self.takeY)
            #将xy坐标分开保存成numpy数组，方便进行曲线拟合
            center=np.array(list_center,dtype=np.int0)
            centerx=center[:,0]
            centery=center[:,1]
            #拟合曲线
            spacey=np.linspace(min(centery),max(centery),500)
            f_yx=scipy.interpolate.interp1d(centery,centerx, kind='cubic')
            spacex=f_yx(spacey)
            self.spacex=spacex
            self.spacey=spacey
            # 计算拐点
            iflc = []
            self.XPictureAnalyze.label_imgView.inflection_auto=[]
            para = np.polyfit(spacey, spacex, 7)
            for n in range(1, len(spacey) - 1):
                fdrv2 = 42 * para[0] * pow(spacey[n - 1], 5) + 30 * para[1] * pow(spacey[n - 1], 4) + 20 * para[
                    2] * pow(spacey[n - 1], 3) + 12 * para[3] * pow(spacey[n - 1], 2) + 6 * para[4] * spacey[
                            n - 1] + 2 * para[5]
                drv2 = 42 * para[0] * pow(spacey[n], 5) + 30 * para[1] * pow(spacey[n], 4) + 20 * para[2] * pow(
                    spacey[n], 3) + 12 * para[3] * pow(spacey[n], 2) + 6 * para[4] * spacey[n] + 2 * para[5]
                bdrv2 = 42 * para[0] * pow(spacey[n + 1], 5) + 30 * para[1] * pow(spacey[n + 1], 4) + 20 * para[
                    2] * pow(spacey[n + 1], 3) + 12 * para[3] * pow(spacey[n + 1], 2) + 6 * para[4] * spacey[
                            n + 1] + 2 * para[5]
                if abs(drv2) <= 0.02 and (bool(fdrv2 >= 0) != bool(bdrv2 >= 0)):
                    iflc.append(n)
            for o in range(0, len(iflc)):
                iflcIndex = iflc[o]
                self.XPictureAnalyze.label_imgView.inflection_auto.append(iflcIndex)

            #将识别结果赋值给图片label
            self.XPictureAnalyze.label_imgView.spacex = spacex
            self.XPictureAnalyze.label_imgView.spacey = spacey
            self.XPictureAnalyze.label_imgView.vert_set=vert_set
            self.XPictureAnalyze.label_imgView.box_set = box_set
            self.XPictureAnalyze.label_imgView.rect_set = rect_set
            self.XPictureAnalyze.label_imgView.list_center = list_center
            self.XPictureAnalyze.label_imgView.update()
            self.XPictureAnalyze.TE_Hint.setText("提示信息：已完成计算，请查看图片上的结果")

    def DlgFn_BNclick_Ximg_MaskAdd(self):
        if self.XPictureAnalyze.label_imgView.stateflag==32:
            self.XPictureAnalyze.label_imgView.stateflag = 0
            self.BN_MaskAdd.setText('脊柱截断添加')
        else:
            self.BN_MaskAdd.setText('停止截断添加')
            self.XPictureAnalyze.label_imgView.stateflag=32

    def DlgFn_BNclick_Ximg_MaskRemove(self):
        if self.XPictureAnalyze.label_imgView.stateflag==31:
            self.XPictureAnalyze.label_imgView.stateflag = 0
            self.BN_MaskRemove.setText('脊柱截断移除')
        else:
            self.BN_MaskRemove.setText('停止截断移除')
            self.XPictureAnalyze.label_imgView.RemoveMask()
            self.XPictureAnalyze.label_imgView.stateflag=31

    def mask_analyze(self,img, box_set):
        # 中点的二维数组索引
        midi = []
        midj = []
        # 中点的xy坐标（图片坐标）
        midxy = []
        # 每个截断的斜率和截距和角度
        SegSlope = []
        SegDistance = []
        SegAngle = []
        #中线坐标
        midlines=[]
        #中线中点坐标
        midcenter=[]


        for box in box_set:
            CalculatedAngle=[]
            # 单个截断的纵向边缘提取
            lstBoxi = []
            lstBoxj = []
            # 存放纵向边缘坐标的两个列表
            lstVerContouri = []
            lstVerContourj = []

            for i in range(0, len(box)):
                # 注意，我接下来是直接操作矩阵里的像素，它们的索引和cv2的图片坐标是不一样的。所以这里是x对应j，y对应i
                lstBoxj.append(box[i][0])
                lstBoxi.append(box[i][1])



            # 每个截断边缘提取的遍历范围,这里去掉了一定的范围，目的是去掉上下边缘，防止干扰
            imin = min(lstBoxi)
            imax = max(lstBoxi)
            iDis = imax - imin
            imin = int(imin + 1 / 5 * iDis)
            imax = int(imax - 1 / 5 * iDis)
            jmin = min(lstBoxj)
            jmax = max(lstBoxj)

            # 控制遍历范围，防止范围超过图片范围
            if imax > img.shape[1] - 3:
                imax = img.shape[1] - 3
            if jmax > img.shape[0] - 3:
                jmax = img.shape[0] - 3
            if imin < 3:
                imin = 3
            if jmin < 3:
                jmin = 3

            # 边缘提取,左右仅有一边和color颜色相同的视为边缘
            for i in range(imin - 2, imax + 2):
                for j in range(jmin - 2, jmax + 2):
                    if bool(img[i, j - 1, 0] == 0 and img[i, j - 1, 1] == 0 and img[i, j - 1, 2] == 0
                            ) != bool(img[i, j + 1, 0] == 0 and img[i, j + 1, 1] == 0 and img[i, j + 1, 2] == 0):
                        lstVerContouri.append(i)
                        lstVerContourj.append(j)

            # 给box的四个顶点按y坐标的大小排序，这样大的两个是上边，小的两个是下边,存为orderbox
            # 复制box，使排序不影响box
            orderbox = Xcopy.deepcopy(box)
            for i in range(0, len(orderbox) - 1):
                for j in range(0, len(orderbox) - 1 - i):
                    if orderbox[j][1] > orderbox[j + 1][1]:
                        # 交换x
                        temp = orderbox[j][0]
                        orderbox[j][0] = orderbox[j + 1][0]
                        orderbox[j + 1][0] = temp
                        # 交换y
                        temp = orderbox[j][1]
                        orderbox[j][1] = orderbox[j + 1][1]
                        orderbox[j + 1][1] = temp
                        # 这种是特殊情况，是中间的两个y坐标相等，我们就需要x坐标来帮助计算斜率
            if orderbox[1][1] == orderbox[2][1] and orderbox[1][0] < orderbox[2][0]:
                Slope = (orderbox[2][1] - orderbox[0][1]) / (orderbox[2][0] - orderbox[0][0])
            else:
                Slope = (orderbox[1][1] - orderbox[0][1]) / (orderbox[1][0] - orderbox[0][0])  # 利用第一个点和第二个点算的斜率
            HoriDis = abs(int(orderbox[1][0] - orderbox[0][0]))

            # 确定截距范围，就是带入y坐标最大和最小的点，这里的截距范围上下各去掉了五分之一（0.2）
            distance = []
            distance.append(orderbox[0][1] - Slope * orderbox[0][0])
            distance.append(orderbox[3][1] - Slope * orderbox[3][0])
            distanceDif = max(distance) - min(distance)
            distanceMax = int(max(distance) - 0.2 * distanceDif) + 1
            distanceMin = int(min(distance) + 0.2 * distanceDif) + 1
            # 利用截距范围进行遍历，然后找中点
            for k in range(distanceMin, distanceMax):
                lstTempi = []
                lstTempj = []
                count = 0
                for l in range(0, len(lstVerContouri)):
                    # 我们的点坐标都是整数，很难完全满足直线方程，所以这里定了一个阈值0.4判断点时候在直线上
                    # 如果改动这个阈值，可能出现整个截断找不到中点的情况
                    if abs(lstVerContouri[l] - lstVerContourj[l] * Slope - k) <= 0.4:
                        count = count + 1
                        lstTempi.append(lstVerContouri[l])
                        lstTempj.append(lstVerContourj[l])
                if count >= 2:
                    midi.append(0.5 * (max(lstTempi) + min(lstTempi)))
                    midj.append(0.5 * (max(lstTempj) + min(lstTempj)))
                    midxy.append([0.5 * (max(lstTempj) + min(lstTempj)), 0.5 * (max(lstTempi) + min(lstTempi))])
            aryMidj = np.array(midj)
            meanj = np.mean(aryMidj)

            for m in range(0, len(midj)):
                if abs(aryMidj[m] - meanj) <= 0.25 * HoriDis:
                    midxy.append([midj[m], midi[m]])

            # 列表转数组，之后套用函数
            aryMidxy = np.array(midxy)
            # cosa和sina是直线和坐标轴的夹角，(pointx,pointy)是直线上的一个点
            try:
                cosa, sina, pointx, pointy = cv2.fitLine(aryMidxy, cv2.DIST_L2, 0, 0.01, 0.01)
                # SingleAngle=math.acos(cosa)
                # 防止斜率无穷大的情况，00021.png里没出现，不影响输出的角度
                if sina == 0:
                    SingleSlope = 0
                    SingleDistance = 0
                    XMax = int(pointx)
                    XMin = int(pointx)
                    YMax = int(XMax * Slope + distanceMax)
                    YMin = int(XMin * Slope + distanceMin)
                else:
                    SingleSlope = sina / cosa
                    SingleDistance = pointy - SingleSlope * pointx
                    XMax = int((distanceMax - SingleDistance) / (SingleSlope - Slope))
                    YMax = int(XMax * Slope + distanceMax)
                    XMin = int((distanceMin - SingleDistance) / (SingleSlope - Slope))
                    YMin = int(XMin * Slope + distanceMin)
                cv2.line(img, (XMax, YMax), (XMin, YMin), (255, 255, 255))
                midlines.append([XMax, YMax, XMin, YMin])
                midcenter.append([int((XMin+XMax)/2),int((YMin+YMax)/2)])
                SegSlope.append(SingleSlope)
                SegDistance.append(SingleDistance)
                if XMax == XMin:
                    CalculatedAngle.append(0.00)
                elif YMax == YMin:
                    CalculatedAngle.append(90.00)
                else:
                    CalculatedAngle.append(round(math.atan(-1 / ((YMin - YMax) / (XMax - XMin))) / 3.1415926 * 180, 2))
            except Exception:
                CalculatedAngle.append(9999.99)
                print('无法拟合直线')
            CalculatedAngle = np.array(CalculatedAngle)
            SegAngle.append(np.mean(CalculatedAngle))
        return SegAngle,midlines,midcenter



    def DlgFn_BNclick_Ximg_MidCalculate(self):
        print('MidCalculate')
        if self.XPictureAnalyze.label_imgView.ImgCount>=1:
            self.XPictureAnalyze.TE_Hint.setText("提示信息：已开始计算中线，请耐心等待计算完成")
            self.XPictureAnalyze.label_imgView.MidClear()
            #新建与缩放后图片大小相同的黑色图片
            img=np.zeros(shape=self.XPictureAnalyze.label_imgView.rsimg.shape)
            box_set=self.XPictureAnalyze.label_imgView.box_set
            vert_set = self.XPictureAnalyze.label_imgView.vert_set
            color=(0,0,255)
            count=0
            for i in range(0, len(vert_set)):
               vert = np.empty((len(vert_set[i]), 2), dtype=np.int32)
               for j in range(0, len(vert_set[i])):
                  vert[j, 0] = round(vert_set[i][j][0])
                  vert[j, 1] = round(vert_set[i][j][1])
               cv2.fillPoly(img, [vert], color)

            angle, midlines, midcenter=self.mask_analyze(img,box_set)
            self.XPictureAnalyze.label_imgView.mid=midlines
            self.XPictureAnalyze.label_imgView.midcenter = midcenter
            angle_max=max(angle)
            angle_min=min(angle)
            print(str(angle_max-angle_min))
            self.XPictureAnalyze.LE_RandomCobbResult.setText(str(angle_max-angle_min))
            self.XPictureAnalyze.LE_RandomCobbResult.setCursorPosition(0)
            #cv2.imwrite("./cut/result.jpg", img)
            self.XPictureAnalyze.label_imgView.update()
            self.XPictureAnalyze.TE_Hint.setText("提示信息：中线计算完成")

    def DlgFn_BNclick_Ximg_MidRemove(self):
        if self.XPictureAnalyze.label_imgView.stateflag==33:
            self.XPictureAnalyze.label_imgView.stateflag = 0
            self.BN_MidRemove.setText('中线移除')
        else:
            self.BN_MidRemove.setText('停止移除')
            self.XPictureAnalyze.label_imgView.RemoveMid()
            self.XPictureAnalyze.label_imgView.stateflag=33

    def DlgFn_BNclick_Ximg_MidAdd(self):
        if self.XPictureAnalyze.label_imgView.stateflag==34:
            self.XPictureAnalyze.label_imgView.stateflag = 0
            self.BN_MidAdd.setText('中线添加')
        else:
            self.BN_MidAdd.setText('停止添加')
            self.XPictureAnalyze.label_imgView.stateflag=34

    # ...........................................................................................
    def calc_sigt(self,I, threshval):
        M, N = I.shape
        ulim = np.uint8(np.max(I))
        N1 = np.count_nonzero(I > threshval)
        N2 = np.count_nonzero(I <= threshval)
        w1 = np.float64(N1) / (M * N)
        w2 = np.float64(N2) / (M * N)
        # print N1,N2,w1,w2
        try:
            u1 = np.sum(
                i * np.count_nonzero(np.multiply(I > i - 0.5, I <= i + 0.5)) / N1 for i in range(threshval + 1, ulim))
            u2 = np.sum(i * np.count_nonzero(np.multiply(I > i - 0.5, I <= i + 0.5)) / N2 for i in range(threshval + 1))
            uT = u1 * w1 + u2 * w2
            sigt = w1 * w2 * (u1 - u2) ** 2
        # print u1,u2,uT,sigt
        except:
            return 0
        return sigt

    # ...........................................................................................
    def get_threshold(self,I):
        max_sigt = 0
        opt_t = 0
        ulim = np.uint8(np.max(I))
        # print(ulim)
        for t in range(ulim + 1):
            sigt = self.calc_sigt(I, t)
            # print t, sigt
            if sigt > max_sigt:
                max_sigt = sigt
                opt_t = t
        # print('optimal high threshold: ', opt_t)
        return opt_t


    def EndAnalyse(self,img, edge_t, edge_l):
        lineNum = 0
        denoise = cv2.fastNlMeansDenoising(img, None, 15, 2, 5)
        denoise=cv2.cvtColor(denoise, cv2.COLOR_BGR2GRAY)
        denoise = np.array(denoise, dtype='uint8')
        clahe = cv2.createCLAHE(clipLimit=4.0, tileGridSize=(2, 2))
        cl1 = clahe.apply(denoise)
        #cl1=denoise
        th = self.get_threshold(cl1)
        edge1 = cv2.Canny(cl1, th / 3, th)
        x = edge1.shape[0]
        y = edge1.shape[1]
        crossline = math.sqrt(x * x + y * y)
        # 霍夫变换检测直线
        for j in range(100, 20, -1):
            lines = cv2.HoughLinesP(edge1, 1, np.pi / 180, j, minLineLength=int(0.75 * y),
                                    maxLineGap=int(0.2 * crossline))
            try:
                length = len(lines)
            except:
                j = j
            else:
                if length < 2:
                    break
        try:
            length = len(lines)
        except:
            lineNum = 0
            line = [[0, 0, 0, 0]]
        else:
            lineNum = length
            line = lines[0]
        return lineNum, line[0][0] + edge_l, line[0][1] + edge_t, line[0][2] + edge_l, line[0][3] + edge_t


    def DlgFn_BNclick_Ximg_EndCalculate(self):
        print('EndCalculate')
        self.XPictureAnalyze.label_imgView.End = []
        self.XPictureAnalyze.label_imgView.Endcenter = []
        if self.XPictureAnalyze.label_imgView.ImgCount >= 1:
            # 新建与缩放后图片大小相同的黑色图片
            self.XPictureAnalyze.label_imgView.EndClear()
            img = Xcopy.deepcopy(self.XPictureAnalyze.label_imgView.rsimg)
            box_set = self.XPictureAnalyze.label_imgView.box_set
            vert_set = self.XPictureAnalyze.label_imgView.vert_set
            count = 0
            for box in box_set:
                count = count + 1
                # 计算box范围
                max_x = max(box[:, 0])
                max_y = max(box[:, 1])
                min_x = min(box[:, 0])
                min_y = min(box[:, 1])
                # 右边界
                edge_r = int(max_x + 0.2 * (max_x - min_x))
                if edge_r > img.shape[1]:
                    edge_r = img.shape[0]
                # 下边界
                edge_b = int(max_y + 0.2 * (max_y - min_y))
                if edge_b > img.shape[0]:
                    edge_b = img.shape[1]
                # 左边界
                edge_l = int(min_x - 0.2 * (max_x - min_x))
                if edge_l < 0:
                    edge_l = 0
                # 上边界
                edge_t = int(min_y - 0.2 * (max_y - min_y))
                if edge_t < 0:
                    edge_t = 0
                # 图片裁剪
                pic = img[edge_t:edge_b, edge_l:edge_r]
                lineNum,x0,y0,x1,y1=self.EndAnalyse(pic,edge_t,edge_l)
                if lineNum>=1:
                    self.XPictureAnalyze.label_imgView.End.append([x0,y0,x1,y1])
                    self.XPictureAnalyze.label_imgView.Endcenter.append([int((x0+x1)/2), int((y0+y1)/2)])
            self.XPictureAnalyze.label_imgView.update()
            self.XPictureAnalyze.TE_Hint.setText("提示信息：边缘标记完成")

    def DlgFn_BNclick_Ximg_EndRemove(self):
        if self.XPictureAnalyze.label_imgView.stateflag==35:
            self.XPictureAnalyze.label_imgView.stateflag = 0
            self.BN_MidRemove.setText('边缘移除')
        else:
            self.BN_MidRemove.setText('停止移除')
            self.XPictureAnalyze.label_imgView.RemoveEnd()
            self.XPictureAnalyze.label_imgView.stateflag=35

    def DlgFn_BNclick_Ximg_EndAdd(self):
        if self.XPictureAnalyze.label_imgView.stateflag==36:
            self.XPictureAnalyze.label_imgView.stateflag = 0
            self.BN_MidAdd.setText('边缘添加')
        else:
            self.BN_MidAdd.setText('停止添加')
            self.XPictureAnalyze.label_imgView.stateflag=36

    def DlgFn_BNclick_Ximg_PelvisCalculate(self):
        print('PelvisCalculate')
        if self.XPictureAnalyze.label_imgView.ImgCount>=1:
            self.XPictureAnalyze.TE_Hint.setText("提示信息：已开始骨盆指标计算，请耐心等待计算完成")
            self.XPictureAnalyze.label_imgView.IliumClear()
            filename=self.XPictureAnalyze.label_imgView.filename
            #新建与缩放后图片大小相同的黑色图片
            img = Xcopy.deepcopy(self.XPictureAnalyze.label_imgView.rsimg)
            box_set=self.XPictureAnalyze.label_imgView.box_set
            vert_set = self.XPictureAnalyze.label_imgView.vert_set
            list_center = self.XPictureAnalyze.label_imgView.list_center
            color=(0,0,255)
            count=0
            center_b=list_center[-1]
            box_x = [box_set[-1][0][0], box_set[-1][1][0], box_set[-1][2][0], box_set[-1][3][0]]
            box_y = [box_set[-1][0][1], box_set[-1][1][1], box_set[-1][2][1], box_set[-1][3][1]]
            x_max=max(box_x)
            x_min=min(box_x)
            length=x_max-x_min
            y_min=min(box_y)
            y_max=max(box_y)
            height=y_max-y_min
            edge_l=int(center_b[0]-3*length)
            if edge_l<=0:
                edge_l=0
            edge_r = int(center_b[0] + 3 * length)
            if edge_r>=img.shape[1]:
                edge_r=img.shape[1]
            edge_b=int(center_b[1]+3*length)
            if edge_b>=img.shape[0]:
                edge_b=img.shape[0]
            img_pelvis=img[int(center_b[1]):edge_b,edge_l:edge_r]
            cv2.imwrite("./cut/PelvisResult/"+filename, img_pelvis)
            print(filename)
            self.XPictureAnalyze.label_imgView.update()
            self.XPictureAnalyze.TE_Hint.setText("提示信息：骨盆指标计算完成")

    def DlgFn_BNclick_Ximg_PelvisAdd(self):
        print('Add')

    def DlgFn_BNclick_Ximg_PelvisRemove(self):
        print('Remove')

    def WgtFn_BNclick_Ximg_FusionConfirm(self):
        fusion_num=3
        str1=''
        str2=''
        str3=''
        index1_1=self.SpineFusion.CB_Fusion1_1.currentIndex()
        index1_2=self.SpineFusion.CB_Fusion1_2.currentIndex() 
        index1_3=self.SpineFusion.CB_Fusion1_3.currentIndex()
        index2_1=self.SpineFusion.CB_Fusion2_1.currentIndex()
        index2_2=self.SpineFusion.CB_Fusion2_2.currentIndex()
        index2_3=self.SpineFusion.CB_Fusion2_3.currentIndex()
        index3_1=self.SpineFusion.CB_Fusion3_1.currentIndex()
        index3_2=self.SpineFusion.CB_Fusion3_2.currentIndex()
        index3_3=self.SpineFusion.CB_Fusion3_3.currentIndex()
        #索引对应脊柱截断名称列表
        if index1_1!=0:
            str1=str1+self.XPictureAnalyze.label_imgView.SpineList[index1_1]
        if index1_2!=0:
            str1=str1+'+'+self.XPictureAnalyze.label_imgView.SpineList[index1_2]
        if index1_3!=0:
            str1=str1+'+'+self.XPictureAnalyze.label_imgView.SpineList[index1_3]
        if index2_1!=0:
            str2=str2+self.XPictureAnalyze.label_imgView.SpineList[index2_1]
        if index2_2!=0:
            str2=str2+'+'+self.XPictureAnalyze.label_imgView.SpineList[index2_2]
        if index2_3!=0:
            str2=str2+'+'+self.XPictureAnalyze.label_imgView.SpineList[index2_3]
        if index3_1!=0:
            str3=str3+self.XPictureAnalyze.label_imgView.SpineList[index3_1]
        if index3_2!=0:
            str3=str3+'+'+self.XPictureAnalyze.label_imgView.SpineList[index3_2]
        if index3_3!=0:
            str3=str3+'+'+self.XPictureAnalyze.label_imgView.SpineList[index3_3]
        #默认有三组锥体融合，如果有一组全是“无”，即索引全为0，则减去这一组，同时设置这一组的结果为“无”
        if index1_1+index1_2+index1_3==0:
            fusion_num=fusion_num-1
            str1='无'
        if index2_1+index2_2+index2_3==0:
            fusion_num=fusion_num-1
            str2='无'
        if index3_1+index3_2+index3_3==0:
            fusion_num=fusion_num-1
            str3='无'
        self.XPictureAnalyze.LE_Fusion1Result.setText(str1)
        self.XPictureAnalyze.LE_Fusion2Result.setText(str2)
        self.XPictureAnalyze.LE_Fusion3Result.setText(str3)
        self.XPictureAnalyze.LE_FusionNumResult.setText(str(fusion_num))
        #关闭窗口
        self.SpineFusion.close()
        
    def WgtFn_BNclick_Ximg_FusionCancel(self):
        #关闭窗口
        self.SpineFusion.close()
    #################################
    
    
    def DlgFn_RBclick_InfoInput_ChooseBackFullMarkerMode(self):
        global g_strCheckModeChosenValue
        g_strCheckModeChosenValue="全标记点模式"
		#显示系统信息
        strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
        self.DlgPatientInfoInput.strMessageOut.setPlainText("全标记点模式\n"+strtem)
        print("全标记点模式")  
    def DlgFn_RBclick_InfoInput_ChooseNoMarkerMode(self):
        global g_strCheckModeChosenValue
        g_strCheckModeChosenValue="无标记点模式"
		#显示系统信息
        strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
        self.DlgPatientInfoInput.strMessageOut.setPlainText("无标记点模式\n"+strtem)
        print("无标记点模式")  

    def DlgFn_RBclick_InfoInput_ChooseContrastMode(self):
        global g_strCheckModeChosenValue
        g_strCheckModeChosenValue="疗效对比模式"
		#显示系统信息
        strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
        self.DlgPatientInfoInput.strMessageOut.setPlainText("疗效对比模式\n"+strtem)
        print("疗效对比模式")  

    def DlgFn_RBclick_InfoInput_ChooseBackroundMarkerCorMode(self):
        global g_strCamPosAdjModeChosenValue
        g_strCamPosAdjModeChosenValue="背景标记点坐标相机姿态矫正模式"
		#显示系统信息
        strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
        self.DlgPatientInfoInput.strMessageOut.setPlainText("背景标记点坐标相机姿态矫正模式\n"+strtem)
        print("背景标记点坐标相机姿态矫正模式")  	

    def DlgFn_RBclick_InfoInput_ChooseRotationAngleMode(self):
        global g_strCamPosAdjModeChosenValue
        g_strCamPosAdjModeChosenValue="三个旋转方向角相机姿态矫正模式"
		#显示系统信息
        strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
        self.DlgPatientInfoInput.strMessageOut.setPlainText("三个旋转方向角相机姿态矫正模式\n"+strtem)
        print("三个旋转方向角相机姿态矫正模式")  

    def on_change_func(self, slider):  # 7
        if slider == self.slider_distance:
            self.slider_distance.setValue(self.slider_distance.value())
            print(str(self.slider_distance.value()))
            # self.label.setText(str(self.slider_distance.value()))
        else:
            pass

    def plot_original_points(self):  # 重绘带背景板和标记点的图
        global g_strDataFolderNameOutHospital
        global global_g_strDataFolderNameOutHospital
        print('test成功')

        try:
            sip.delete(self.canvas_contour_map)  # 删除addwidget的窗口
        except BaseException:
            print('无canvas_contour_map')
        try:
            sip.delete(self.canvas_Correct_spine_points)
        except BaseException:
            print('无canvas_Correct_spine_points')
        try:
            sip.delete(self.canvas_Correct_spine_points)
        except BaseException:
            print('无canvas_Correct_spine_points')
        
        self.figure_contour_map = plt.figure()
        self.canvas_contour_map = FigureCanvas(self.figure_contour_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.canvas_contour_map)

        #    from matplotlib import cm      #里面有很多颜色映射表；

        # '数据邹可/back3d邹可.asc'
        points = np.loadtxt(global_g_strDataFolderNameOutHospital + '/back3d.asc',comments=['*'])  #
        points = np.array(points)
        skip = 5  # 对原始点云数据的采样比

        self.axes = self.figure_contour_map.add_subplot(111)
        # plt.hold(True)
        # self.axes.hold(False)
        # self.axes = Axes3D(self.figure_contour_map)
        # self.axes = self.figure_contour_map.gca(projection='3d')
        point_range = range(0, points.shape[0], skip)  # skip points to prevent crash
        y = points[point_range, 0]  # 注意x和y调换
        x = points[point_range, 1]
        z = points[point_range, 2]  # 保留原始数据应该加负号，因为不影响颜色显示所以暂时不加
        # y = np.transpose(y)
        ##########通过标记点来平移点云使得髂后上棘的中点为原点
        # '数据邹可/1Result.txt'
        Unmarked = False
        try:
            fMarkerPoint = np.loadtxt(global_g_strDataFolderNameOutHospital + '/1Result.txt', dtype=str,
                                      comments=['*', 'Su', 'AL'])  # , skiprows=i跳过前i行
        except IOError:  # 无标记点文件
            Unmarked = True
            print('请选择髂后上棘和腋窝')
            # box = QMessageBox(QMessageBox.Question, '退出', '确定退出？')    # 完成提醒窗口
            try:
                fMarkerPoint = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Unmarked_points_correct_fMarkerPoint.txt')
                # fMarkerPoint = plt.ginput(4)  # 右键撤销上一个点
                fMarkerPoint_x = fMarkerPoint[:, 1]
                fMarkerPoint_y = fMarkerPoint[:, 0]
                fMarkerPoint_z = fMarkerPoint[:, 2]
                translation_x = (1 / 2) * (np.array(fMarkerPoint_x[-4:-3]) + np.array(fMarkerPoint_x[-3:-2]))
                translation_y = (1 / 2) * (np.array(fMarkerPoint_y[-4:-3]) + np.array(fMarkerPoint_y[-3:-2]))
                translation_z = (1 / 2) * (np.array(fMarkerPoint_z[-4:-3]) + np.array(fMarkerPoint_z[-3:-2]))
                # 左髂后上棘为倒数第四个点，右髂后上棘为倒数第三个点
                x = x - translation_x
                y = y - translation_y
                z = z - translation_z
                fMarkerPoint_x = fMarkerPoint_x - translation_x
                fMarkerPoint_y = fMarkerPoint_y - translation_y
                fMarkerPoint_z = fMarkerPoint_z - translation_z
                # fMarkerPoint_x, fMarkerPoint_y = fMarkerPoint_y, fMarkerPoint_x
            except BaseException:
                reply = QMessageBox.information(self, '提示', '请先手动选取14点', QMessageBox.Yes)


        else:  # 有标记点文件
            fMarkerPoint = np.loadtxt(global_g_strDataFolderNameOutHospital + '/1Result.txt', dtype=str,
                                      comments=['*', 'Su', 'A'])  # , skiprows=i跳过前i行
            if len(fMarkerPoint) <= 6:
                number_fMarkerPoint = len(fMarkerPoint)  # 六标记点
            else:
                number_fMarkerPoint = len(fMarkerPoint) # 多标记点的最后四个为错误点 - 4 20190618
            fMarkerPoint_x = []
            fMarkerPoint_y = []
            fMarkerPoint_z = []
            for i in range(number_fMarkerPoint):  # range(fMarkerPoint.shape[0])#去除最后四个点-4
                #fMarkerPoint3d = fMarkerPoint[i][1].split('(')[2][:-1]  # 提取出该行的坐标参数
                fMarkerPoint3d = fMarkerPoint[i].split('(')[3][:-1]  # 提取出该行的坐标参数20190618
                fMarkerPoint_x.append(float(fMarkerPoint3d.split(',')[0]))  # 提取出该行的坐标参数x变为数值
                fMarkerPoint_y.append(float(fMarkerPoint3d.split(',')[1]))  # 提取出该行的坐标参数y变为数值
                fMarkerPoint_z.append(float(fMarkerPoint3d.split(',')[2]))  # 提取出该行的坐标参数z变为数值
                # print(fMarkerPoint3d)
            # print(fMarkerPoint[0])
            fMarkerPoint_x, fMarkerPoint_y = fMarkerPoint_y, fMarkerPoint_x

        ##########通过标记点来平移点云使得髂后上棘的中点为原点结束
        try:
            Sampling_points = [x, y, z]  # 采样后的点,x=Sampling_points[0]
            # new_points =
            minx = min(Sampling_points[0])
            minx = math.ceil(minx)
            maxx = max(Sampling_points[0])
            maxx = math.floor(maxx)
            miny = min(Sampling_points[1])
            miny = math.ceil(miny)
            maxy = max(Sampling_points[1])
            maxy = math.floor(maxy)
            minz = min(Sampling_points[2])
            minz = math.ceil(minz)
            maxz = max(Sampling_points[2])
            maxz = math.floor(maxz)
            Number_of_segments = 600  # 将点云分割为800段#值越大颜色变化越快
            Contour_value = np.linspace(minz, maxz, Number_of_segments)  # 一共获取12个数值

            ind = []  # ind为各个段的位置索引，初始化
            colormap = []  # 颜色数组colormap，初始化
            Color_step = 15  # 值越大颜色变化越快
            for i in range(Number_of_segments - 1):  # range(600)为0-599
                ## 画等高线
                Transition_variable = np.where(
                    (Sampling_points[2] >= Contour_value[i]) & (Sampling_points[2] < Contour_value[i + 1]))
                ind.append(Transition_variable)  # ind为各个段的位置索引
                ## 定义颜色数组colormap，随z值变化而变化的颜色

                if (((Color_step * i) // 255) == 0):
                    color_R = 220 / 255
                    color_G = (255 - (Color_step * i) % 255) / 255
                    color_B = 1
                elif (((Color_step * i) // 255) == 1):
                    color_R = 230 / 255
                    color_G = 0
                    color_B = (255 - (Color_step * i) % 255) / 255
                elif (((Color_step * i) // 255) == 2):
                    color_R = 230 / 255
                    color_G = ((Color_step * i) % 255) / 255
                    color_B = 0
                elif (((Color_step * i) // 255) == 3):
                    color_R = (255 - (Color_step * i) % 255) / 255
                    color_G = 1
                    color_B = 0
                elif (((Color_step * i) // 255) == 4):
                    color_R = 0
                    color_G = 1
                    color_B = ((Color_step * i) % 255) / 255
                elif (((Color_step * i) // 255) == 5):
                    color_R = 0
                    color_G = (255 - (Color_step * i) % 255) / 255
                    color_B = 1
                elif (((Color_step * i) // 255) == 6):
                    color_R = 0
                    color_G = 0
                    color_B = (255 - (Color_step * i) % 255) / 255
                else:  # 白色111（不显示）
                    color_R = 0
                    color_G = 0
                    color_B = 0

                color = (color_R, color_G, color_B)
                colormap.append(color)
                # exec(index + '=Transition_variable')#字符串变为变量名

            for i in reversed(range(Number_of_segments - 1)):  # 循环画图，reversed反转数组
                # for i in range(Number_of_segments - 1):  # 循环画图，reversed反转数组
                plt.plot(Sampling_points[0][ind[i][0]], Sampling_points[1][ind[i][0]], '.', marker='.', markersize=2,
                         color=colormap[i])  # , markersize为点的大小

            # 画标记点

            plt.plot(fMarkerPoint_x[:], fMarkerPoint_y[:], '.', marker='*', markersize=5,
                     color='k')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]

            # 坐标显示操作
            # self.axes.view_init(elev=-90, azim=0)
            self.axes.axis('equal')  # 等比例显示
            self.axes.set_xlim((minx, maxx))
            self.axes.set_ylim((miny, maxy))
            # self.axes.set_zlim((minz-200, maxz+200))
            self.axes.set_xticks(np.linspace(minx, maxx, 6))
            self.axes.set_yticks(np.linspace(miny, maxy, 5))
            # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
            self.axes.set_xlabel('X Label')
            self.axes.set_ylabel('Y Label')
            print('plot_contour_map')
        except BaseException:
            pass
    def DlgFn_BNclick_InfoInput_ReadDataAndProcess(self):  # 绘制云纹图
		# #############判断当前检测模式，全标记点模式
		# 0----"全标记点模式")
        # 1----"无标记点模式")
        # 2----"疗效对比模式" 
        #g_strCheckModeChosenValue="全标记点模式"
        if g_strCheckModeChosenValue=="全标记点模式":
            print("当前选择全标记点模式")
			#读取标记点坐标，含脊柱标记点10个，髂嵴左右各一个，两肩各一个，
            # 数据格式：n*3，头朝右，坐标排布xyz，脊柱方向是x，两肩方向是y

            ListfMarkerPoint3d=self.ReadMarkerCor(g_strPatientFullPath)    
			
        elif g_strCheckModeChosenValue=="无标记点模式":
            print("当前选择无标记点模式")
        elif g_strCheckModeChosenValue=="疗效对比模式":
            print("当前选择疗效对比模式")
		# ##################
		# #############判断当前背景模式，4个标记点坐标模式
        global g_fArrayBKPtCorEnd
        global g_fArrayMarkerPtCorEnd
        global fBKCamAngleXYZ
        fBKPtCorXYZ=[[0,0,0,0],[0,0,0,0],[0,0,0,0]]
        if g_strCamPosAdjModeChosenValue=="背景标记点坐标相机姿态矫正模式":

            #如果是用背景标记点坐标计算旋转角，则读取背景标记点，计算旋转角
            print("当前选择背景标记点坐标相机姿态矫正模式")
			#读取背景标记点坐标并拷贝到患者目录
		    #判断当前患者目录是否存在背景矫正点
            # 
            # #读取背景标记点
            fBKPtCorXYZ=self.MainGenFn_ReadBKCor() 
            # 通过四个背景点计算三个方向旋转角RXYZ
            fBKCamAngleXYZ=YH_DataPreprocess.fRotationAngleCalWith4BkMarker(fBKPtCorXYZ)

            
        elif g_strCamPosAdjModeChosenValue=="三个旋转方向角相机姿态矫正模式":
            print("三个旋转方向角相机姿态矫正模式")
            fBKCamAngleXYZ=self.MainGenFn_ReadBKCamAngleXYZ() 
           
        ### 现在已经获得了背景标记点以及脊柱标记点，下面步骤，后背点云无论是否有背景标记点，都预留数组位置，没有时置为0。下面进行点云和标记点姿态矫正
        # 站姿中位姿态矫正        
        fMarkerPointAndBK=np.c_[ListfMarkerPoint3d, fBKPtCorXYZ]			
        iNumMarkerPointAndBK = fMarkerPointAndBK.shape[1]
        onesMatrix=np.ones((1, iNumMarkerPointAndBK))
         #normListfMarkerPointAndBK-->4*(n+numbk)
        normListfMarkerPointAndBK=np.r_[fMarkerPointAndBK, onesMatrix]
  
		#fNormBKCorAfterAdj-->4*4*(n+numbk，脊柱+4标记点+4背景点)

        fNormListMarkerPointAndBKAfterAdj=YH_DataPreprocess.fPtAdjWithCamAngle(normListfMarkerPointAndBK, fBKCamAngleXYZ)
		#将原点移动到髂棘连线fPtMovetoCenterofIliaca(fNormMarkerPtCorXYZ, iHasBkPt)
        
        fNormListMarkerPointAndBKAfterAdj=YH_DataPreprocess.fPtMovetoCenterofIliaca(fNormListMarkerPointAndBKAfterAdj)

        fMarkerPointAndBKAfterAdj = np.delete(fNormListMarkerPointAndBKAfterAdj, 3, 0)  # 删除A的第二行
        #分离出背景点坐标
        fMtrxBKPtCorEnd=fMarkerPointAndBKAfterAdj[:, iNumMarkerPointAndBK-4:iNumMarkerPointAndBK]
        g_fArrayBKPtCorEnd=fMtrxBKPtCorEnd.A
        fMtrxMarkerPtCorEnd=fMarkerPointAndBKAfterAdj[:, 0:iNumMarkerPointAndBK-4]
        g_fArrayMarkerPtCorEnd=fMtrxMarkerPtCorEnd.A
        ######## 

		# ################## 
		#后背标记点分离，分为特殊点和脊柱点
		#g_fArrayBKPtCorEnd--> 3*numbk
        #g_fArrayMarkerPtCorEnd-->3*numMkPt
        global g_fArraySpineMarkerPtCorEnd
        iNumSpineMarkerPt = g_fArrayMarkerPtCorEnd.shape[1]
        g_fArraySpineMarkerPtCorEnd=g_fArrayMarkerPtCorEnd[:, 0:iNumSpineMarkerPt-4]
		# 脊柱节段插值n*3
        global g_fPtArrayAfterInterpolate_n_3
        g_fPtArrayAfterInterpolate_n_3=YH_DataPreprocess.fPtInterpolate(g_fArraySpineMarkerPtCorEnd)
        fPtendT=g_fPtArrayAfterInterpolate_n_3.T
		# 计算曲率半径
        global g_centerx
        global g_centery
        global g_R 
        global g_fMainCStartPtScale
        global g_fMainCEndPtScale
        
		# 获取主C弯起止点间的点列
        iStartPtNo=0#插值后的点列按照比例计算的C弯起点
        iEndPtNo=0#插值后的点列按照比例计算的C弯终点
        fminPt_x=0 #np.min(g_fPtArrayAfterInterpolate_n_3[:, 1])
        #fmaxPt_x=np.max(g_fPtArrayAfterInterpolate_n_3[:, 1])
        fmaxPt_x=np.max(g_fPtArrayAfterInterpolate_n_3)
        fMainCStartPtY=(fmaxPt_x-fminPt_x)*g_fMainCStartPtScale+fminPt_x
        fMainCEndPtY=(fmaxPt_x-fminPt_x)*g_fMainCEndPtScale+fminPt_x

        for i in range(len(g_fPtArrayAfterInterpolate_n_3)-1):
            datatem=g_fPtArrayAfterInterpolate_n_3[i, 0]
            fsdis=abs(fMainCStartPtY-g_fPtArrayAfterInterpolate_n_3[i, 0])
            #print(i, "-", fsdis)
            if fsdis<1:
                iStartPtNo = i #脊柱曲线的拐点初筛
            fedis=abs(fMainCEndPtY-g_fPtArrayAfterInterpolate_n_3[i, 0])
            #print(i, "-", fedis)
            if fedis<1:
                iEndPtNo = i #脊柱曲线的拐点初筛
		#截取的主C弯点列	
        fMainCPt=g_fPtArrayAfterInterpolate_n_3[iStartPtNo:iEndPtNo, :]
        #print(g_fPtArrayAfterInterpolate_n_3[iStartPtNo, :], g_fPtArrayAfterInterpolate_n_3[iEndPtNo, :])
        g_centerx, g_centery, g_R = self.circleLeastFit(fMainCPt)
        print(g_centerx, g_centery, g_R)
        iCirclepointNum=300
        circlex=np.zeros(iCirclepointNum)
        circley=np.zeros(iCirclepointNum)
        cita0=0
        if g_centerx>0:
            cita0=math.pi*9/10
        else:
            cita0=-math.pi/10
        for j in range(iCirclepointNum):
            cita=cita0+math.pi*j/1200
            circlex[j]=g_R*math.cos(cita)+g_centerx
            circley[j]=g_R*math.sin(cita)+g_centery
           

        p1=np.array([g_fPtArrayAfterInterpolate_n_3[iStartPtNo, 0],g_fPtArrayAfterInterpolate_n_3[iStartPtNo, 1]])
        p2=np.array([g_fPtArrayAfterInterpolate_n_3[iEndPtNo, 0],g_fPtArrayAfterInterpolate_n_3[iEndPtNo, 1]])
        p3=p2-p1
        disc=math.hypot(p3[0],p3[1])/2#求起点和终点的距离，除以2用于求cobb角度
        
        cobb=2*math.asin(disc/g_R)*180/math.pi
        cobbMessage=np.array([g_centerx, g_centery, g_R, cobb])
        print('cx=%f,cy=%f,R=%f,cloudCobb=%f')
        print(g_centerx, g_centery, g_R, cobb)
        np.savetxt(g_strPatientFullPath + '/SpinePointListAfterAdj.txt', g_fPtArrayAfterInterpolate_n_3, fmt='%.2f')
        np.savetxt(g_strPatientFullPath + '/SubSpinePointListAfterAdjforMainC.txt', fMainCPt, fmt='%.2f')
        np.savetxt(g_strPatientFullPath + '/cobbMessageCX_CY_R_COBB.txt', cobbMessage, fmt='XC,YC,R,cloudCobb，%.2f')
	#画标记点
   #绘制冠状面背景点
        
		# 冠状面（英文名： coronal plane）
        xCP_BK = g_fArrayBKPtCorEnd[0, :]
        yCP_BK = g_fArrayBKPtCorEnd[1, :]
		
        xCP_MK = g_fArrayMarkerPtCorEnd[0, :]
        yCP_MK = g_fArrayMarkerPtCorEnd[1, :]
		
        xCP_SP = g_fPtArrayAfterInterpolate_n_3[:, 0]
        yCP_SP = g_fPtArrayAfterInterpolate_n_3[:, 1]
		
        xCP_SPc = fMainCPt[:, 0]
        yCP_SPc = fMainCPt[:, 1]
		#self.Ui_DataPreprocessing.graphicsView_CoronalPlane.getsize
        self.figure_contour_map = plt.figure(figsize=(12,15))
        self.canvas_contour_map = FigureCanvas(self.figure_contour_map)
        plt.grid()# 打开图形网格  
        xlim(-300,300)
        ylim(-100, 700)
        #plt.axis('equal')  # 等比例显示
        sValue=90 #点大小
		
        plt.scatter(xCP_BK,yCP_BK,c = 'b',s=sValue,marker = 'o')  
        plt.scatter(xCP_MK,yCP_MK,c = 'r',s=sValue,marker = 'o')  
        plt.scatter(xCP_SP,yCP_SP,c = 'g',s=10,marker = 'o') 
        plt.scatter(xCP_SPc,yCP_SPc,c = 'r',s=10,marker = 'o') 
        plt.scatter(circlex,circley,c = 'r',s=10,marker = 'o') 
        plt.scatter(g_fPtArrayAfterInterpolate_n_3[iStartPtNo, 0],g_fPtArrayAfterInterpolate_n_3[iStartPtNo, 1], c = 'b',s=60,marker = 'x') 
        plt.scatter(g_fPtArrayAfterInterpolate_n_3[iEndPtNo, 0], g_fPtArrayAfterInterpolate_n_3[iEndPtNo, 1],c = 'b',s=60,marker = 'x') 
		
        plt.plot([-200,200],[0,0], color='k')
        plt.plot([0,0],[-200,200], color='k')
		#plt.plot([-200,200],[0,0], color='k')
        graphicscene = QtWidgets.QGraphicsScene()  # 第三步，创建一个QGraphicsScene，因为加载的图形（FigureCanvas）不能直接放到graphicview控件中，必须先放到graphicScene，然后再把graphicscene放到graphicview中
        graphicscene.addWidget(self.canvas_contour_map)  # 第四步，把图形放到QGraphicsScene中，注意：图形是作为一个QWidget放到QGraphicsScene中的
        self.DlgPatientInfoInput.BKPointAfterAdjCoronalView.setScene(graphicscene) # 第五步，把QGraphicsScene放入QGraphicsView
        self.DlgPatientInfoInput.BKPointAfterAdjCoronalView.show()  # 最后，调用show方法呈现图形！Voila!!
         #plt.plot(fBKCorAfterAdj[0, iNumMarkerPointAndBK-3:iNumMarkerPointAndBK],fBKCorAfterAdj[1, iNumMarkerPointAndBK-3:iNumMarkerPointAndBK],  '.', marker='*', markersize=5,
		#	color='b')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]
		#绘制矢状面背景点,矢状面（sagittal plane）
        xSP_BK = g_fArrayBKPtCorEnd[2, :]
        ySP_BK = g_fArrayBKPtCorEnd[1, :]
		
        xSP_MK = g_fArrayMarkerPtCorEnd[2, :]
        ySP_MK = g_fArrayMarkerPtCorEnd[1, :]
		
        xSP_SP = g_fPtArrayAfterInterpolate_n_3[:, 2]
        ySP_SP = g_fPtArrayAfterInterpolate_n_3[:, 1]
		
        xSP_SPc = fMainCPt[:, 2]
        ySP_SPc = fMainCPt[:, 1]
		
        self.figure_contour_map2 = plt.figure(figsize=(12,15))
        self.canvas_contour_map2 = FigureCanvas(self.figure_contour_map2)
        plt.grid()# 打开图形网格  
        xlim(-300,300)
        ylim(-100, 700)
        #plt.axis('equal')  # 等比例显示
        sValue=90 #点大小
		
        plt.scatter(xSP_BK,ySP_BK,c = 'b',s=sValue,marker = 'o')  
        plt.scatter(xSP_MK,ySP_MK,c = 'r',s=sValue,marker = 'o')  	
        plt.scatter(xSP_SP,ySP_SP,c = 'g',s=10,marker = 'o')  	
        plt.scatter(xSP_SPc,ySP_SPc,c = 'r',s=10,marker = 'o') 
		
        plt.plot([-200,200],[0,0], color='k')
        plt.plot([0,0],[-200,200], color='k')
        #plt.plot(fBKCorAfterAdj[2, iNumMarkerPointAndBK-3:iNumMarkerPointAndBK],fBKCorAfterAdj[1, iNumMarkerPointAndBK-3:iNumMarkerPointAndBK],  '.', marker='*', markersize=5,
		#	color='r')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]
        graphicscene2 = QtWidgets.QGraphicsScene()  # 第三步，创建一个QGraphicsScene，因为加载的图形（FigureCanvas）不能直接放到graphicview控件中，必须先放到graphicScene，然后再把graphicscene放到graphicview中
        graphicscene2.addWidget(self.canvas_contour_map2)  # 第四步，把图形放到QGraphicsScene中，注意：图形是作为一个QWidget放到QGraphicsScene中的
        self.DlgPatientInfoInput.BKPointAfterAdjSagittalView.setScene(graphicscene2) # 第五步，把QGraphicsScene放入QGraphicsView
        self.DlgPatientInfoInput.BKPointAfterAdjSagittalView.show()  # 最后，调用show方法呈现图形！Voila!!

         # 坐标显示操作
        #global g_strPatientFullPath
        global fImgScale
        global px
        global py

#        global g_fMainCStartPtScale
 #       global g_fMainCEndPtScale
        global g_fLineAngleForCobb
        pxy=[px, py]
        ScaleStartEndCobb=[g_fMainCStartPtScale, g_fMainCEndPtScale, g_fLineAngleForCobb[0], g_fLineAngleForCobb[1],g_fLineAngleForCobb[2]]
		
        #np.savetxt(g_strPatientFullPath + '/ximg4pointandscaleandcobb.txt', pxy, g_fMainCStartPtScale, g_fMainCEndPtScale, g_fLineAngleForCobb)
        np.savetxt(g_strPatientFullPath + '/ximg4point.txt', pxy, fmt='%.2f')
        np.savetxt(g_strPatientFullPath + '/scaleandcobb.txt', ScaleStartEndCobb, fmt='上比例，下比例，上倾角，下倾角，cobb角，%.2f')

    def Correcting_the_spine_points(self):  # 自动计算脊沟点
        global fMarkerPoint
        global g_strDataFolderNameOutHospital
        global global_g_strDataFolderNameOutHospital
        import sip
        try:
            sip.delete(self.canvas_contour_map)  # 删除addwidget的窗口
        except BaseException:
            sip.delete(self.canvas_Correct_spine_points)
        # else:
        #     sip.delete(self.canvas_contour_map)  # 删除addwidget的窗口
        # sip.delete(self.toolbar)
        try:
            sip.delete(self.canvas_Correct_spine_points)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Correct_spine_points_map = plt.figure()
        self.canvas_Correct_spine_points = FigureCanvas(self.figure_Correct_spine_points_map)
        # # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.canvas_Correct_spine_points)
        self.axes_Correct_spine_points = self.figure_Correct_spine_points_map.add_subplot(111)
        # points = np.loadtxt(
        #     g_strDataFolderNameOutHospital + '/final_points.txt')  # back3d.asc  back3dtest.asc   back3d(无标记点)，back3d邹可，back3d孙锐宁
        points = np.loadtxt(global_g_strDataFolderNameOutHospital + '/back3d.asc',comments=['*'])
        points = np.array(points)
        #大致切除背景
        points_ind = np.where(abs(points[:, 1]) < 50)#20 201906042
        points_temporary = points[points_ind[0],:]
        points_ind = np.where(abs(points_temporary[:, 0]) < 50)#20 201906042
        points_temporary = points_temporary[points_ind[0],:]
        cut_z = np.mean(points_temporary[:,2])
        points_ind = np.where(points[:, 2] < cut_z + 80)#80
        points = points[points_ind[0], :]
        skip = 5  # 对原始点云数据的采样比

        point_range = range(0, points.shape[0], skip)  # skip points to prevent crash
        x = points[point_range, 1]  # 注意x和y调换
        y = points[point_range, 0]
        z = points[point_range, 2]  # 保留原始数据应该加负号，因为不影响颜色显示所以暂时不加

        Sampling_points = [x, y, z]  # 采样后的点,x=Sampling_points[0]
        # new_points =
        minx = min(Sampling_points[0])
        minx = math.ceil(minx)
        maxx = max(Sampling_points[0])
        maxx = math.floor(maxx)
        miny = min(Sampling_points[1])
        miny = math.ceil(miny)
        maxy = max(Sampling_points[1])
        maxy = math.floor(maxy)
        minz = min(Sampling_points[2])
        minz = math.ceil(minz)
        maxz = max(Sampling_points[2])
        maxz = math.floor(maxz)
        Number_of_segments = 180  # 将点云分割为800段#值越大颜色变化越快
        Contour_value = np.linspace(minz, maxz, Number_of_segments)  # 一共获取12个数值

        ind = []  # ind为各个段的位置索引，初始化
        colormap = []  # 颜色数组colormap，初始化
        Color_step = 10  # 值越大颜色变化越快
        for i in range(Number_of_segments - 1):  # range(600)为0-599
            ## 画等高线
            Transition_variable = np.where(
                (Sampling_points[2] >= Contour_value[i]) & (Sampling_points[2] < Contour_value[i + 1]))
            ind.append(Transition_variable)  # ind为各个段的位置索引
            ## 定义颜色数组colormap，随z值变化而变化的颜色
            # 111->101->100->110->010->011->001->000
            #        if (abs((i)%15) <= 1):#绘制黑色等高线
            #            color_R = 0
            #            color_G = 0
            #            color_B = 0

            if (((Color_step * i) // 255) == 0):
                color_R = 220 / 255
                color_G = (255 - (Color_step * i) % 255) / 255
                color_B = 1
            elif (((Color_step * i) // 255) == 1):
                color_R = 230 / 255
                color_G = 0
                color_B = (255 - (Color_step * i) % 255) / 255
            elif (((Color_step * i) // 255) == 2):
                color_R = 230 / 255
                color_G = ((Color_step * i) % 255) / 255
                color_B = 0
            elif (((Color_step * i) // 255) == 3):
                color_R = (255 - (Color_step * i) % 255) / 255
                color_G = 1
                color_B = 0
            elif (((Color_step * i) // 255) == 4):
                color_R = 0
                color_G = 1
                color_B = ((Color_step * i) % 255) / 255
            elif (((Color_step * i) // 255) == 5):
                color_R = 0
                color_G = (255 - (Color_step * i) % 255) / 255
                color_B = 1
            elif (((Color_step * i) // 255) == 6):
                color_R = 0
                color_G = 0
                color_B = (255 - (Color_step * i) % 255) / 255
            else:  # 白色111（不显示）
                color_R = 0
                color_G = 0
                color_B = 0

            color = (color_R, color_G, color_B)
            colormap.append(color)
            # exec(index + '=Transition_variable')#字符串变为变量名

        for i in reversed(range(Number_of_segments - 1)):  # 循环画图，reversed反转数组
            # for i in range(Number_of_segments - 1):  # 循环画图，reversed反转数组
            plt.plot(Sampling_points[0][ind[i][0]], Sampling_points[1][ind[i][0]], '.', marker='.', markersize=2,
                     color=colormap[i])  # , markersize为点的大小
        Contour_value_contour = np.linspace(minz, maxz, 25)
        for i in Contour_value_contour:
            ind_contour = np.where((z >= i) & (z < i + 0.8))  # 非常要注意这个括号 没有括号估计内部执行顺序不对，捣腾不出来的，具体原因可评论留言
            # ind_contour.append(ind_Transition)
            plt.plot(Sampling_points[0][ind_contour[0]], Sampling_points[1][ind_contour[0]], '.', marker='.',
                     markersize=0.8, color='k')
        self.axes_Correct_spine_points.axis('equal')  # 等比例显示
        self.axes.axis('equal')  # 等比例显示
        self.axes.set_xlim((minx, maxx))
        self.axes.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes.set_xticks(np.linspace(minx, maxx, 6))
        self.axes.set_yticks(np.linspace(miny, maxy, 5))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes.set_xlabel('X Label')
        self.axes.set_ylabel('Y Label')
        # plt.plot( [0, 0], [miny, maxy],color='slategray')#竖直方向
        # <editor-fold desc="自动获得脊柱曲线...">

        '''自动获得脊柱曲线'''
    # try:
    #     fMarkerPoint = np.loadtxt(global_g_strDataFolderNameOutHospital + '/Unmarked_points_correct_fMarkerPoint.txt')
    # except IOError:
        dirs = global_g_strDataFolderNameOutHospital + '/相关数据'
        if not os.path.exists(dirs):  # 不存在该路径则新建
            os.makedirs(dirs)
        try:
            fBKPtCor = plt.ginput(14, timeout=-1) # 右键撤销上一个点,中键停止，timeout<=0,表示无定时
            distance_point = []
            sort_distance_point = []
            min_distance_ind_first = []  # 初始化从小到大排列的数在原始数据中的索引位置
            min_distance_ind = []
            for i in range(len(fBKPtCor)):
                distance_point.append((points[:, 1] - fBKPtCor[i][0]) ** 2 + (
                        points[:, 0] - fBKPtCor[i][1]) ** 2)  # 计算第一个点与各个点的距离
                distance_point[i] = list(distance_point[i])  # 将np.array变量转为普通的列表，因为在后面用到list.index，是列表才有的属性
                min_distance_ind.append(np.where(distance_point[i] == min(distance_point[i])))
            min_distance_ind_final = []
            for i in range(len(fBKPtCor)):
                min_distance_ind_final.append(min_distance_ind[i][0][0])
            fMarkerPoint_orl = [points[min_distance_ind_final, 0], points[min_distance_ind_final, 1],
                                points[min_distance_ind_final, 2]]  # 按行上下合并
            fMarkerPoint_orl = np.array(fMarkerPoint_orl)
            fMarkerPoint_orl = np.transpose(fMarkerPoint_orl)  # 转置
            # gauge_3d = [points[:,0][min_distance_ind], points[:,1][min_distance_ind],points[:,2][min_distance_ind]]
            # fBKPtCor_3d = [first_3d, second_3d, third_3d, fourth_3d]
            np.savetxt(global_g_strDataFolderNameOutHospital + '/相关数据' + "/Unmarked_points_correct_fMarkerPoint.txt", fMarkerPoint_orl, fmt='%.6f')
            root = Tk()
            root.withdraw()  # ****实现主窗口隐藏
            messagebox.showinfo("提示", "标记点保存完成！")
            # else:
            fMarkerPoint = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Unmarked_points_correct_fMarkerPoint.txt')
            # self.chonghua()
        except BaseException:
            pass


        # </editor-fold>

        # ##画脊柱曲线
        # Spine_Fitting_Curve = np.loadtxt(g_strDataFolderNameOutHospital + '/Spine_Fitting_Curve.txt')
        # plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
        #          markersize=3, color='b')

        # 画标记点
        # fMarkerPoint = np.loadtxt(g_strDataFolderNameOutHospital + '/correct_fMarkerPoint.txt')
        # import time
        # time.sleep("5")

    def plot_Coronal_and_Sagittal_plane_with_BG(self):  # 显示去除干扰点(含背景板)的冠状面和矢状面
        global g_strDataFolderNameOutHospital
        global global_g_strDataFolderNameOutHospital
        # points = np.array(points)
        skip = 5   # 对原始点云数据的采样比
        # Reserv_background_final_points = YH_DataPreprocess.Reserv_background_final_points
        Reserv_background_final_points = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Reserv_background_final_points.txt')  #
        point_range = range(0, Reserv_background_final_points.shape[0], skip)  # skip points to prevent crash
        x_Reserv_background = Reserv_background_final_points[point_range, 0]#注意x和y调换
        y_Reserv_background = Reserv_background_final_points[point_range, 1]
        z_Reserv_background = Reserv_background_final_points[point_range, 2]

        minx = min(x_Reserv_background)
        minx = math.ceil(minx)
        maxx = max(x_Reserv_background)
        maxx = math.floor(maxx)
        miny = min(y_Reserv_background)
        miny = math.ceil(miny)
        maxy = max(y_Reserv_background)
        maxy = math.floor(maxy)
        minz = min(z_Reserv_background)
        minz = math.ceil(minz)
        maxz = max(z_Reserv_background)
        maxz = math.floor(maxz)

## 画冠状面
        try:
            sip.delete(self.canvas_Coronal_plane_with_background_plate)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Coronal_plane_with_background_plate = plt.figure()
        plt.grid(which='major', axis='x', linewidth=0.75, linestyle='-', color='k')
        plt.grid(which='major', axis='y', linewidth=0.75, linestyle='-', color='k')
        # https://blog.csdn.net/weixin_41789707/article/details/81035997

        self.canvas_Coronal_plane_with_background_plate = FigureCanvas(
            self.figure_Coronal_plane_with_background_plate)
        self.Ui_DataPreprocessing.gridLayout_3.addWidget(self.canvas_Coronal_plane_with_background_plate)
        self.axes_Coronal_plane = self.figure_Coronal_plane_with_background_plate.add_subplot(111)
        #
        # Number_of_segments = 100  # 将点云分割为800段#值越大颜色变化越快
        # Contour_value = np.linspace(minz, maxz, Number_of_segments)  # 一共获取12个数值
        # ind = []
        # colormap_Coronal = []  # 颜色数组colormap，初始化
        # Color_step = 15  # 值越大颜色变化越快15
        # for i in range(Number_of_segments - 1):
        #     ## ind为每个分段在原点云中的索引位置
        #     Transition_variable = np.where((Reserv_background_final_points[:,2] >= Contour_value[i]) & (
        #                 Reserv_background_final_points[:,2] < Contour_value[i + 1]))
        #     ind.append(Transition_variable)  # ind为各个段的位置索引
        #
        #     if (((Color_step * i) // 255) == 0):
        #         color_R = 220 / 255
        #         color_G = (255 - (Color_step * i) % 255) / 255
        #         color_B = 1
        #     elif (((Color_step * i) // 255) == 1):
        #         color_R = 230 / 255
        #         color_G = 0
        #         color_B = (255 - (Color_step * i) % 255) / 255
        #     elif (((Color_step * i) // 255) == 2):
        #         color_R = 230 / 255
        #         color_G = ((Color_step * i) % 255) / 255
        #         color_B = 0
        #     elif (((Color_step * i) // 255) == 3):
        #         color_R = (255 - (Color_step * i) % 255) / 255
        #         color_G = 1
        #         color_B = 0
        #     elif (((Color_step * i) // 255) == 4):
        #         color_R = 0
        #         color_G = 1
        #         color_B = ((Color_step * i) % 255) / 255
        #     elif (((Color_step * i) // 255) == 5):
        #         color_R = 0
        #         color_G = (255 - (Color_step * i) % 255) / 255
        #         color_B = 1
        #     elif (((Color_step * i) // 255) == 6):
        #         color_R = 0
        #         color_G = 0
        #         color_B = (255 - (Color_step * i) % 255) / 255
        #     else:  # 白色111（不显示）
        #         color_R = 0
        #         color_G = 0
        #         color_B = 0
        #
        #     color_Coronal = (color_R, color_G, color_B)
        #     colormap_Coronal.append(color_Coronal)

        # for i in reversed(range(Number_of_segments - 1)):  # 循环画图，reversed反转数组
        #     plt.plot(Reserv_background_final_points[:, 0][ind[i][0]], Reserv_background_final_points[:, 1][ind[i][0]],
        #              '.', marker='.', markersize=2, color=colormap_Coronal[i])  # , markersize为点的大小
        plt.plot(Reserv_background_final_points[:, 0], Reserv_background_final_points[:, 1],
                              '.', marker='.', markersize=2, color='orange')#orange
        # fBKPtCor_3d = np.loadtxt(global_g_strDataFolderNameOutHospital + "/fBKPtCor_3d.txt")
        fBKPtCor_3d = np.loadtxt(r'D:\灵犀脊柱形态评价系统\矫正后的背景点.txt')
        plt.plot(fBKPtCor_3d[:, 0], fBKPtCor_3d[:, 1],
                 '.', marker='+', markersize=5, color='k')  # orange
            # try:
            #     plt.plot(Reserv_background_final_points[:,0][ind[i][0]], Reserv_background_final_points[:,1][ind[i][0]],'.',marker='.', markersize=2, color=colormap_Coronal[i])  # , markersize为点的大小
            # except IndexError:
            #     print('i',i)
            #     continue
        self.axes_Coronal_plane.axis('equal')  # 等比例显示
        self.axes_Coronal_plane.set_xlim((minx, maxx))
        self.axes_Coronal_plane.set_ylim((miny, maxy))
        # self.axes_Coronal_plane.set_zlim((minz-200, maxz+200))
        self.axes_Coronal_plane.set_xticks(np.linspace(minx, maxx, 6))
        self.axes_Coronal_plane.set_yticks(np.linspace(miny, maxy, 5))
        # self.axes_Coronal_plane.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes_Coronal_plane.set_xlabel('X Label')
        self.axes_Coronal_plane.set_ylabel('Y Label')

## 画矢状面
        try:
            sip.delete(self.canvas_Sagittal_plane_with_background_plate)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Sagittal_plane_with_background_plate = plt.figure()
        plt.grid(which='major', axis='x', linewidth=0.75, linestyle='-', color='k')
        # plt.grid(axis="x",linestyle='-.')#c='r',，
        #https://blog.csdn.net/weixin_41789707/article/details/81035997
        self.canvas_Sagittal_plane_with_background_plate = FigureCanvas(self.figure_Sagittal_plane_with_background_plate)
        self.Ui_DataPreprocessing.gridLayout_4.addWidget(self.canvas_Sagittal_plane_with_background_plate)
        self.axes_Sagittal_plane = self.figure_Sagittal_plane_with_background_plate.add_subplot(111)

        # Number_of_segments = 100  # 将点云分割为800段#值越大颜色变化越快
        # Contour_value = np.linspace(minz, maxz, Number_of_segments)  # 一共获取 个数值
        # ind = []
        # colormap_Sagittal = []  # 颜色数组colormap，初始化
        # Color_step = 15  # 值越大颜色变化越快
        # for i in range(Number_of_segments - 1):
        #     ## ind为每个分段在原点云中的索引位置
        #     Transition_variable = np.where((Reserv_background_final_points[:, 2] >= Contour_value[i]) & (
        #             Reserv_background_final_points[:, 2] < Contour_value[i + 1]))
        #     ind.append(Transition_variable)  # ind为各个段的位置索引
        #
        #     if (((Color_step * i) // 255) == 0):
        #         color_R = 220 / 255
        #         color_G = (255 - (Color_step * i) % 255) / 255
        #         color_B = 1
        #     elif (((Color_step * i) // 255) == 1):
        #         color_R = 230 / 255
        #         color_G = 0
        #         color_B = (255 - (Color_step * i) % 255) / 255
        #     elif (((Color_step * i) // 255) == 2):
        #         color_R = 230 / 255
        #         color_G = ((Color_step * i) % 255) / 255
        #         color_B = 0
        #     elif (((Color_step * i) // 255) == 3):
        #         color_R = (255 - (Color_step * i) % 255) / 255
        #         color_G = 1
        #         color_B = 0
        #     elif (((Color_step * i) // 255) == 4):
        #         color_R = 0
        #         color_G = 1
        #         color_B = ((Color_step * i) % 255) / 255
        #     elif (((Color_step * i) // 255) == 5):
        #         color_R = 0
        #         color_G = (255 - (Color_step * i) % 255) / 255
        #         color_B = 1
        #     elif (((Color_step * i) // 255) == 6):
        #         color_R = 0
        #         color_G = 0
        #         color_B = (255 - (Color_step * i) % 255) / 255
        #     else:  # 白色111（不显示）
        #         color_R = 0
        #         color_G = 0
        #         color_B = 0
        #
        #     color_Sagittal = (color_R, color_G, color_B)
        #     colormap_Sagittal.append(color_Sagittal)
        #
        # for i in reversed(range(Number_of_segments - 1)):  # 循环画图，reversed反转数组
        #     plt.plot(Reserv_background_final_points[:, 2][ind[i][0]],
        #              Reserv_background_final_points[:, 1][ind[i][0]], '.', marker='.', markersize=2,
        #              color=colormap_Sagittal[i])  # , markersize为点的大小color=colormap_Sagittal[i]
        plt.plot(Reserv_background_final_points[:, 2],Reserv_background_final_points[:, 1], '.', marker='.', markersize=2,
                              color='orange')#chocolate
        plt.plot(-fBKPtCor_3d[:, 2], fBKPtCor_3d[:, 1],
                 '.', marker='+', markersize=5, color='k')  # orange
            # try:
            #     plt.plot(Reserv_background_final_points[:, 2][ind[i][0]],
            #              Reserv_background_final_points[:, 1][ind[i][0]], '.', marker='.', markersize=2,
            #              color=colormap_Sagittal[i])  # , markersize为点的大小
            # except IndexError:
            #     print('i', i)
            #     continue
        self.axes_Sagittal_plane.axis('equal')  # 等比例显示
        self.axes_Sagittal_plane.set_xlim((minz, maxz))
        self.axes_Sagittal_plane.set_ylim((miny, maxy))
        # self.axes_Sagittal_plane.set_zlim((minz-200, maxz+200))
        self.axes_Sagittal_plane.set_xticks(np.linspace(minz, maxz, 3))
        self.axes_Sagittal_plane.set_yticks(np.linspace(miny, maxy, 5))
        # self.axes_Sagittal_plane.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes_Sagittal_plane.set_xlabel('Z Label')
        self.axes_Sagittal_plane.set_ylabel('Y Label')
        self.plot_original_points()#重绘使得可以重新选四个点

    def circleLeastFit(self,points):  # 获取曲线对应的拟合圆的半径和圆心位置
        N = len(points)
        sum_x = 0
        sum_y = 0
        sum_x2 = 0
        sum_y2 = 0
        sum_x3 = 0
        sum_y3 = 0
        sum_xy = 0
        sum_x1y2 = 0
        sum_x2y1 = 0
        for i in range(N):
            x = float(points[i, 0])
            y = float(points[i, 1])
            x2 = x * x
            y2 = y * y
            sum_x += x
            sum_y += y
            sum_x2 += x2
            sum_y2 += y2
            sum_x3 += x2 * x
            sum_y3 += y2 * y
            sum_xy += x * y
            sum_x1y2 += x * y2
            sum_x2y1 += x2 * y
        C = N * sum_x2 - sum_x * sum_x
        D = N * sum_xy - sum_x * sum_y
        E = N * sum_x3 + N * sum_x1y2 - (sum_x2 + sum_y2) * sum_x
        G = N * sum_y2 - sum_y * sum_y
        H = N * sum_x2y1 + N * sum_y3 - (sum_x2 + sum_y2) * sum_y
        a = (H * D - E * G) / (C * G - D * D + 1e-100)
        b = (H * C - E * D) / (D * D - G * C + 1e-100)
        c = -(a * sum_x + b * sum_y + sum_x2 + sum_y2) / N
        centerx = a / (-2)
        centery = b / (-2)
        rad = math.sqrt(a * a + b * b - 4 * c) / 2
        return centerx, centery, rad

    def plot_Index_chart(self):
        global g_strDataFolderNameOutHospital
        global global_g_strDataFolderNameOutHospital
        global minx, maxx, miny, maxy, Spine_Fitting_Curve
        global Main_bend_center, Main_bend_Boundary_points,Main_bend_rad
        ##for Redraw_clouds_point
        global final_fMarkerPoint_end
        global final_points
        global Sampling_points,fMarkerPoint
        global minz, maxz,z
        ##for Redraw_clouds_point
        ##for YH_Result_dynamic
        global attitude
        ##for YH_Result_dynamic

        # dataprocess_main(Display_prompt_box = False)#20190605删除

        # <editor-fold desc="画后背点云图...">
        '''画后背点云图Ui_ResultShowFirst CResultShowFirst'''
        try:
           sip.delete(self.canvas_back_cloud_points_map)  # 删除addwidget的窗口
        except BaseException:
           pass
        self.figure_back_cloud_points_map = plt.figure()

        self.canvas_back_cloud_points_map = FigureCanvas(self.figure_back_cloud_points_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowFirst.gridLayout_2.addWidget(self.canvas_back_cloud_points_map)

        ###############20190506
        final_fMarkerPoint_end = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/correct_fMarkerPoint.txt')  #矫正以后的标记点
            #  , skiprows=i跳过前i行, dtype=str, comments=['*', 'Su', 'AL']
        #这里的final_fMarkerPoint_end x值和y值是原始数据的顺序，即x，y，z的顺序
        final_points = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/final_points.txt')  # 矫正去除背景后的点云数据
        final_fMarkerPoint_end = np.array(final_fMarkerPoint_end)
        from scipy import interpolate
        Spinal_marker_points_x = np.array(final_fMarkerPoint_end[:-4, 1])  # 只取脊柱上的标记点
        Spinal_marker_points_y = np.array(final_fMarkerPoint_end[:-4, 0])
        Spinal_marker_points_z = np.array(final_fMarkerPoint_end[:-4, 2])

        xnew = np.linspace(min(Spinal_marker_points_x), max(Spinal_marker_points_x), 500)  # 插值后的x(数据的上下方向)值,大约1mm选一个点
        #####################三次b样条计算出函数f
        f_xy = interpolate.interp1d(Spinal_marker_points_x, Spinal_marker_points_y, kind="cubic")
        # ‘slinear’, ‘quadratic’ and ‘cubic’ refer to a spline interpolation of first, second or third order)
        ynew = f_xy(xnew)  # 三次b样条计算出y值
        # 矢状面的拟合曲线
        f_xz = interpolate.interp1d(Spinal_marker_points_x, Spinal_marker_points_z, kind="cubic")
        # ‘slinear’, ‘quadratic’ and ‘cubic’ refer to a spline interpolation of first, second or third order)
        znew = f_xz(xnew)  # 三次b样条计算出y值
        Spine_Fitting_Curve = np.hstack((ynew[:, np.newaxis], xnew[:, np.newaxis], znew[:, np.newaxis]))
        # hstack()在行上合并numpy数组  vstack()在列上合并numpy数组
        # Spine_Fitting_Curve
        np.savetxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Spine_Fitting_Curve.txt', Spine_Fitting_Curve, fmt='%.8f')  # 后面的计算预测C弯需要用到精确值
        number_vertical = 30  # 将脊柱曲线分为30段
        x_vertical = np.linspace(min(np.array(final_fMarkerPoint_end[:-4, 1])),
                                 max(np.array(final_fMarkerPoint_end[:-4, 1])),
                                 number_vertical)  # 选number_vertical个点，准备进行画垂线
        count = 0  # 为区分文件名计数
        all_middle_points = []  # 各个截段的中点的x和y初始化
        Side_spin_angle = []  # 侧旋角初始化
        dirs = global_g_strDataFolderNameOutHospital + r'/各个截段点及中点位置'
        if not os.path.exists(dirs):  # 不存在该路径则新建
            os.makedirs(dirs)
        x_straight_line = []
        y_straight_line = []
        for j in x_vertical:  # 每个截段遍历
            for i in range(len(xnew)):  # 寻找截段附近的脊柱曲线上的点的中点
                if ((j > xnew[i]) and (j < xnew[i + 1])):
                    middle = [(1 / 2) * (xnew[i] + xnew[i + 1]), (1 / 2) * (ynew[i] + ynew[i + 1])]  # 垂线取最临近的俩个点中点
                    Slope = -1 / ((xnew[i] - xnew[i + 1]) / (ynew[i] - ynew[i + 1]))  # 得到垂线的斜率
                    # 画直线y=k(x-x0)+y0,因为要根据y值来划定画直线的范围，所以推导出x=(y-y0)/k+x0
                    x_straight_line_orl = np.linspace(middle[1] - 20, middle[1] + 20, 50)  # 每层垂线的x取值
                    x_straight_line = np.append(x_straight_line, x_straight_line_orl)  # 垂线的x值
                    y_straight_line = np.append(y_straight_line,
                                                (x_straight_line_orl - middle[1]) * Slope + middle[0])  # 垂线y值
                    # plt.plot(middle[1], middle[0], '.', marker='*', markersize=7, color='b')  # 显示线上点
                    # plt.plot(x_straight_line, y_straight_line, '.', marker='.', markersize=3, color='g')  # 显示过点垂线
                    #   点到直线距离公式：D = abs(A*x0+By0+C)/sqrt(A**2+B**2)公式中的直线方程为Ax + By + C = 0，点P的坐标为(x0, y0)
                    distance = []  # 初始化点到垂线的距离
                    for k in range(len(final_points)):
                        x0 = final_points[k][0]
                        y0 = final_points[k][1]
                        distance.append(
                            abs(Slope * x0 - y0 + middle[0] - Slope * middle[1]) / math.sqrt(Slope ** 2 + (-1) ** 2))
                    distance = np.array(distance)
                    Ind_distance = np.where(distance < 2)
                    Segment_points = final_points[Ind_distance, :]  # 找到垂线附近的点的三维坐标
                    middle[0], middle[1] = middle[1], middle[0]  # 调换x和y的位置，使得第一列是x，第二列是y
                    all_middle_points.append(middle)  # 各个截段的中点的x和y
                    # <editor-fold desc="对截段进行矫正...">
                    #########对截段进行矫正（截面为按脊柱曲线的斜率的垂线进行寻找，所以需要矫正到平行于水平面）
                    gamma = -math.atan(Slope)
                    # z轴方向相反，所以角度加负号
                    # 得到最终的点云数据
                    final_four_points_correct_X = (final_points[Ind_distance, 0] - middle[0]) * math.cos(gamma) - (
                            final_points[Ind_distance, 1] - middle[1]) * math.sin(
                        gamma)  # -middle[0]为把中点平移到原点，方便绕z轴旋转
                    final_four_points_correct_Y = (final_points[Ind_distance, 0] - middle[0]) * math.sin(gamma) + (
                            final_points[Ind_distance, 1] - middle[1]) * math.cos(
                        gamma)
                    final_four_points_correct_Z = final_points[Ind_distance, 2]
                    final_points_end = np.hstack(
                        (final_four_points_correct_X[0][:, np.newaxis] + middle[0],  # +middle[0]为把中点从原点平移到原始位置
                         final_four_points_correct_Y[0][:, np.newaxis] + middle[1],
                         final_four_points_correct_Z[0][:, np.newaxis]))  # 最终点云final_points_X[:, np.newaxis](转置操作)

                    # plt.plot(Segment_points[0][:,0], Segment_points[0][:,1], '.', marker='.', markersize=3, color='b')  # 显示过点垂线附近的点（x,y）
                    # plt.plot(Segment_points[0][:, 0], Segment_points[0][:, 1] + Segment_points[0][:, 2], '.',
                    #          marker='.', markersize=3, color='b')  # 显示过点垂线附近的点（x,y+z ）
                    # </editor-fold>

                    save_Segment_points_filename = global_g_strDataFolderNameOutHospital + r'/各个截段点及中点位置' + r'\Segment_points' + str(
                        count) + '.txt'
                    # np.savetxt(save_Segment_points_filename, Segment_points[0], fmt='%.2f')  # 保存各个截段的点坐标
                    np.savetxt(save_Segment_points_filename, final_points_end, fmt='%.2f')  # 保存各个截段的点坐标
                    count = count + 1
        # 保存垂线x和y的值
        x_straight_line = np.array(x_straight_line)
        y_straight_line = np.array(y_straight_line)
        spine_vertical = np.hstack((x_straight_line[:, np.newaxis], y_straight_line[:, np.newaxis]))
        np.savetxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/spine_vertical.txt', spine_vertical, fmt='%.2f')
        save_all_middle_points = global_g_strDataFolderNameOutHospital + '/各个截段点及中点位置' + r'\all_middle_points.txt'
        np.savetxt(save_all_middle_points, all_middle_points, fmt='%.2f')  # 保存各个截段的中点的x和y
        all_middle_points = np.array(all_middle_points)
        middle_and_Segment_filename = global_g_strDataFolderNameOutHospital + '/各个截段点及中点位置'
        # middle_points = np.loadtxt(middle_and_Segment_filename + r'\all_middle_points.txt')  # 导入各个截段的中点的x和y
        save_Side_opening_left_Segment_point = []  # 保存旁左侧开点初始化
        save_Side_opening_right_Segment_point = []  # 保存旁右侧开点初始化
        for i in range(number_vertical - 2):  # 一共有number_vertical - 2个截段（从颈七到髂后上棘中点分了30段，一共产生28个截段）
            Segment_points = np.loadtxt(global_g_strDataFolderNameOutHospital + r'\各个截段点及中点位置' + r'\Segment_points' + str(
                i) + '.txt')  # 获取对应截段的数据middle_and_Segment_filename +
            middle_point = all_middle_points[i, :]  # 获取该截段的中点的x和y值
            middle_point_x = middle_point[0]  # 获取该截段的中点的x值
            Ind_left_Segment_points = np.where(Segment_points[:, 0] < middle_point_x)  # 获取该截段左半部分的数据索引
            Ind_right_Segment_points = np.where(Segment_points[:, 0] >= middle_point_x)  # 获取该截段右半部分的数据索引
            left_Segment_points = Segment_points[Ind_left_Segment_points, :]  # 截段的左半部分数据
            left_Segment_points = left_Segment_points[0]  # 去掉一层[]
            right_Segment_points = Segment_points[Ind_right_Segment_points, :]  # 截段的右半部分数据
            right_Segment_points = right_Segment_points[0]  # 去掉一层[]
            ##################7截面曲线旋转角均值(绝对值)(按右高为正算)
            Side_opening_distance = (final_fMarkerPoint_end[-3:-2, 0] - final_fMarkerPoint_end[-4:-3, 0]) / 2
            # 左髂后上棘为倒数第 四个点，右髂后上棘为倒数第三个点
            distance_left = abs(left_Segment_points[:, 0] - (middle_point_x - Side_opening_distance))  # 左旁开的x（左右方向）位置
            distance_right = abs(right_Segment_points[:, 0] - (middle_point_x + Side_opening_distance))
            Side_opening_left_Segment_point = left_Segment_points[(np.where(distance_left[:] == min(distance_left[:]))),
                                              :]  # 截段左侧旁开点
            Side_opening_right_Segment_point = right_Segment_points[
                                               (np.where(distance_right[:] == min(distance_right[:]))), :]  # 截段右侧旁开点
            Side_opening_left_Segment_point = Side_opening_left_Segment_point[0][0]  # 去掉多余[]
            Side_opening_right_Segment_point = Side_opening_right_Segment_point[0][0]
            slope_Side_spin_angle = (Side_opening_left_Segment_point[2] - Side_opening_right_Segment_point[2]) / (
                    Side_opening_left_Segment_point[0] - Side_opening_right_Segment_point[0])  # 左右旁开处的斜率（相对z和x坐标）
            Side_spin_angle.append(math.atan(slope_Side_spin_angle) * 180 / math.pi)  # 截段侧旋角
            # 保存旁开点
            save_Side_opening_left_Segment_point.append(Side_opening_left_Segment_point)
            save_Side_opening_right_Segment_point.append(Side_opening_right_Segment_point)
            # Side_opening_Segment_points = np.vstack((Max_Z_Left_Two_third_cloud_points[0],Max_Z_Right_Two_third_cloud_points[0]))
        save_left_Side_opening_points_filepath = middle_and_Segment_filename + r'\left_Side_opening_points.txt'
        np.savetxt(save_left_Side_opening_points_filepath, save_Side_opening_left_Segment_point, fmt='%.2f')  # 保存左侧旁开点
        save_right_Side_opening_points_filepath = middle_and_Segment_filename + r'\right_Side_opening_points.txt'
        np.savetxt(save_right_Side_opening_points_filepath, save_Side_opening_right_Segment_point,
                   fmt='%.2f')  # 保存右侧旁开点

        save_Side_spin_angle_filename = middle_and_Segment_filename + r'\Side_spin_angle.txt'
        np.savetxt(save_Side_spin_angle_filename, Side_spin_angle, fmt='%.2f')  ########保存各个截段的侧旋角6，7，9，10，11，19
        ############ 20190506

        #    from matplotlib import cm      #里面有很多颜色映射表；
        points = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/final_points.txt')  # back3d.asc  back3dtest.asc   back3d(无标记点)，back3d邹可，back3d孙锐宁
        points = np.array(points)
        skip = 1  # 对原始点云数据的采样比

        self.axes1 = self.figure_back_cloud_points_map.add_subplot(111)
        point_range = range(0, points.shape[0], skip)  # skip points to prevent crash
        x = points[point_range, 0]  # 注意x和y调换
        y = points[point_range, 1]
        z = -points[point_range, 2]  # 保留原始数据应该加负号，因为不影响颜色显示所以暂时不加

        Sampling_points = [x, y, z]  # 采样后的点,x=Sampling_points[0]
        # new_points =
        minx = min(Sampling_points[0])
        minx = math.ceil(minx)
        maxx = max(Sampling_points[0])
        maxx = math.floor(maxx)
        miny = min(Sampling_points[1])
        miny = math.ceil(miny)
        maxy = max(Sampling_points[1])
        maxy = math.floor(maxy)
        minz = min(Sampling_points[2])
        minz = math.ceil(minz)
        maxz = max(Sampling_points[2])
        maxz = math.floor(maxz)
        Number_of_segments = 180  # 将点云分割为800段#值越大颜色变化越快
        Contour_value = np.linspace(minz, maxz, Number_of_segments)  # 一共获取12个数值

        ind = []  # ind为各个段的位置索引，初始化
        colormap = []  # 颜色数组colormap，初始化
        Color_step = 10  # 值越大颜色变化越快
        for i in range(Number_of_segments - 1):  # range(600)为0-599

            Transition_variable = np.where(
                (Sampling_points[2] >= Contour_value[i]) & (Sampling_points[2] < Contour_value[i + 1]))
            ind.append(Transition_variable)  # ind为各个段的位置索引

            ## 定义颜色数组colormap，随z值变化而变化的颜色
            # 111->101->100->110->010->011->001->000
            #        if (abs((i)%15) <= 1):#绘制黑色等高线
            #            color_R = 0
            #            color_G = 0
            #            color_B = 0

            if (((Color_step * i) // 255) == 0):
                color_R = 230 / 255#220 / 255#1本来应该为1，为了把最开始的白色亮度调低，而减小这个值
                color_G = (255 - (Color_step * i) % 255) / 255
                color_B = 230 / 255#1本来应该为1，为了把最开始的白色亮度调低，而减小这个值
            elif (((Color_step * i) // 255) == 1):
                color_R = 1#220 / 255
                color_G = 0
                color_B = (255 - (Color_step * i) % 255) / 255
            elif (((Color_step * i) // 255) == 2):
                color_R = 1#220 / 255
                color_G = ((Color_step * i) % 255) / 255
                color_B = 0
            elif (((Color_step * i) // 255) == 3):
                color_R = (255 - (Color_step * i) % 255) / 255
                color_G = 1
                color_B = 0
            elif (((Color_step * i) // 255) == 4):
                color_R = 0
                color_G = 1
                color_B = ((Color_step * i) % 255) / 255
            elif (((Color_step * i) // 255) == 5):
                color_R = 0
                color_G = (255 - (Color_step * i) % 255) / 255
                color_B = 1
            elif (((Color_step * i) // 255) == 6):
                color_R = 0
                color_G = 0
                color_B = (255 - (Color_step * i) % 255) / 255
            else:  # 白色111（不显示）
                color_R = 0
                color_G = 0
                color_B = 0

            color = (color_R, color_G, color_B)
            colormap.append(color)
            # exec(index + '=Transition_variable')#字符串变为变量名

        for i in reversed(range(Number_of_segments - 1)):  # 循环画图，reversed反转数组
            # for i in range(Number_of_segments - 1):  # 循环画图，reversed反转数组
            plt.plot(Sampling_points[0][ind[i][0]], Sampling_points[1][ind[i][0]], '.', marker='.', markersize=2,
                     color=colormap[i])  # , markersize为点的大小
        ##画脊柱曲线
        Spine_Fitting_Curve = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Spine_Fitting_Curve.txt')
        plt.plot(Spine_Fitting_Curve[:,0], Spine_Fitting_Curve[:,1], '.', marker='.',
                 markersize=3, color='b')
        ## 画等高线
        Contour_value_contour = np.linspace(-100, 300, 81)  # -60, 60取25个数，使得以5为间隔
        # ind_contour = []
        for i in Contour_value_contour:
            ind_contour = np.where((z >= i ) & (z < i + 0.8))  # 非常要注意这个括号 没有括号估计内部执行顺序不对，捣腾不出来的，具体原因可评论留言
            # ind_contour.append(ind_Transition)
            plt.plot(Sampling_points[0][ind_contour[0]], Sampling_points[1][ind_contour[0]], '.', marker='.',markersize=0.8,color='k')
        #生成网格
        # plt.grid(which='major', axis='x', linewidth=0.75, linestyle='-', color='r')
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画标记点
        fMarkerPoint = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/correct_fMarkerPoint.txt')
        plt.plot(fMarkerPoint[:, 0], fMarkerPoint[:, 1], '.', marker='*', markersize=5,
                 color='b')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]
        #画过零点直线
        plt.plot([minx-100,maxx+100],[0,0], color='slategray')#lavender
        # plt.plot( [0, 0], [miny, maxy],color='slategray')#竖直方向
        # 坐标显示操作
        # self.axes.view_init(elev=-90, azim=0)
        self.axes1.axis('equal')  # 等比例显示
        self.axes1.set_xlim((minx, maxx))
        self.axes1.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes1.set_xticks(np.linspace(minx, maxx, 8))
        self.axes1.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes1.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes1.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        dirs = g_strDataFolderNameOutHospital + "\数据结果图片"
        if not os.path.exists(dirs):  # 不存在该路径则新建
            os.makedirs(dirs)
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\后背点云图.png', format='png')

        print('后背点云图')
        # </editor-fold>

        # <editor-fold desc="脊柱曲线垂线倾角...">
        '''脊柱曲线垂线倾角Ui_ResultShowFirst CResultShowFirst'''
        try:
            sip.delete(self.canvas_Vertical_inclination_of_spine_map)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Vertical_inclination_of_spine_map = plt.figure()
        self.canvas_Vertical_inclination_of_spine_map = FigureCanvas(self.figure_Vertical_inclination_of_spine_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowFirst.gridLayout_3.addWidget(self.canvas_Vertical_inclination_of_spine_map)
        self.axes2 = self.figure_Vertical_inclination_of_spine_map.add_subplot(111)
        spine_vertical = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/spine_vertical.txt')#导入脊柱垂线
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画过零点直线
        plt.plot([minx, maxx], [0, 0], color='slategray')  # lavender
        plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                 markersize=3, color='k')# 画脊柱拟合曲线
        plt.plot(spine_vertical[:,0],spine_vertical[:,1], '.', marker='.',
                 markersize=3, color='g')# 画脊柱垂线
        # 坐标显示操作
        # self.axes.view_init(elev=-90, azim=0)
        self.axes2.axis('equal')  # 等比例显示
        self.axes2.set_xlim((minx, maxx))
        self.axes2.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes2.set_xticks(np.linspace(minx, maxx, 8))
        self.axes2.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes2.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes2.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\脊柱曲线垂线倾角.png', format='png')
        # </editor-fold>

        # <editor-fold desc="冠状面曲线图...">
        '''冠状面曲线图'''
        try:
            sip.delete(self.canvas_Coronal_curve_map)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Coronal_curve_map = plt.figure()
        self.canvas_Coronal_curve_map = FigureCanvas(self.figure_Coronal_curve_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowSecond.gridLayout_2.addWidget(self.canvas_Coronal_curve_map)
        self.axes3 = self.figure_Coronal_curve_map.add_subplot(111)
        plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                 markersize=3, color='b')# 画脊柱拟合曲线
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画过零点直线
        plt.plot([minx - 100, maxx + 100], [0, 0], color='slategray')  # lavender
        # 画标记点
        plt.plot(fMarkerPoint[:, 0], fMarkerPoint[:, 1], '.', marker='.', markersize=5,
                 color='b')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]


        self.axes3.axis('equal')  # 等比例显示
        self.axes3.set_xlim((minx, maxx))
        self.axes3.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes3.set_xticks(np.linspace(minx, maxx, 8))
        self.axes3.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes3.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes3.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\冠状面曲线图.png', format='png')

        # </editor-fold>

        # <editor-fold desc="矢状面曲线图...">
        '''矢状面曲线图'''
        try:
            sip.delete(self.canvas_Sagittal_curve_map)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Sagittal_curve_map = plt.figure()
        self.canvas_Sagittal_curve_map = FigureCanvas(self.figure_Sagittal_curve_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowSecond.gridLayout_3.addWidget(self.canvas_Sagittal_curve_map)
        self.axes4 = self.figure_Sagittal_curve_map.add_subplot(111)
        # plt.plot(Spine_Fitting_Curve[:, 2], Spine_Fitting_Curve[:, 1], '.', marker='.',
        #          markersize=3, color='b')# 画脊柱拟合曲线
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画过零点直线
        plt.plot([minx - 100, maxx + 100], [0, 0], color='slategray')  # lavender
        # 画点云
        plt.plot(-Sampling_points[2], Sampling_points[1], '.', marker='.', markersize=2,color='yellow')
        # 负号为恢复原始的状态，画云纹图时上色将z值取反
        # 画标记点
        # plt.plot(fMarkerPoint[:, 2], fMarkerPoint[:, 1], '.', marker='.', markersize=5,
        #          color='k')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]
        plt.plot(Spine_Fitting_Curve[:, 2], Spine_Fitting_Curve[:, 1], '.', marker='.',
                 markersize=3, color='k')
        self.axes4.axis('equal')  # 等比例显示
        self.axes4.set_xlim((minx, maxx))
        self.axes4.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes4.set_xticks(np.linspace(minx, maxx, 8))
        self.axes4.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes4.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes4.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\矢状面曲线图.png', format='png')
        # </editor-fold>

        # <editor-fold desc="后背截面曲线图...">
        '''后背截面曲线图'''
        try:
            sip.delete(self.canvas_Back_Section_Curve_map)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Back_Section_Curve_map = plt.figure()
        self.canvas_Back_Section_Curve_map = FigureCanvas(self.figure_Back_Section_Curve_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowThird.gridLayout_2.addWidget(self.canvas_Back_Section_Curve_map)
        self.axes5 = self.figure_Back_Section_Curve_map.add_subplot(111)
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画过零点直线
        plt.plot([minx, maxx], [0, 0], color='slategray')  # lavender
        # 画截面曲线
        for count in range(28):
        #os.getcwd()
            Section_Curve = np.loadtxt(global_g_strDataFolderNameOutHospital + r'\各个截段点及中点位置' + r'\Segment_points' + str(count) + '.txt')  # 矫正以后的标记点
            if (count < 5):
                plt.plot(Section_Curve[:, 0], Section_Curve[:, 1] - Section_Curve[:, 2], '.', marker='.',markersize=3, color='yellow')# 画脊柱拟合曲线
            elif (5 <= count < 9):
                plt.plot(Section_Curve[:, 0], Section_Curve[:, 1] - Section_Curve[:, 2], '.', marker='.', markersize=3,
                         color='k')  # 画脊柱拟合曲线
            elif (9 <= count < 19):
                plt.plot(Section_Curve[:, 0], Section_Curve[:, 1] - Section_Curve[:, 2], '.', marker='.', markersize=3,
                         color='lime')  # 画脊柱拟合曲线
            else:
                plt.plot(Section_Curve[:, 0], Section_Curve[:, 1] - Section_Curve[:, 2], '.', marker='.', markersize=3,
                         color='blue')  # 画脊柱拟合曲线

        self.axes5.axis('equal')  # 等比例显示
        self.axes5.set_xlim((minx, maxx))
        self.axes5.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes5.set_xticks(np.linspace(minx, maxx, 8))
        self.axes5.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes5.set_xlim((min(Section_Curve[:, 0]), max(Section_Curve[:, 0])))
        # self.axes5.set_ylim((min(Section_Curve[:, 1]-Section_Curve[:, 2]), max(Section_Curve[:, 1]-Section_Curve[:, 2])))
        # # self.axes.set_zlim((minz-200, maxz+200))
        # self.axes5.set_xticks(np.linspace(min(Section_Curve[:, 0]), max(Section_Curve[:, 0]), 8))
        # self.axes5.set_yticks(np.linspace(min(Section_Curve[:, 1]-Section_Curve[:, 2]), max(Section_Curve[:, 1]-Section_Curve[:, 2]), 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes5.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes5.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\后背截面曲线图.png', format='png')

        # </editor-fold>

        # <editor-fold desc="后背截面倾角（侧旋角）曲线图...">
        '''后背截面倾角（侧旋角）曲线图'''
        try:
            sip.delete(self.canvas_Back_Section_Inclination_Curve_map)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Back_Section_Inclination_Curve_map = plt.figure()
        self.canvas_Back_Section_Inclination_Curve_map = FigureCanvas(self.figure_Back_Section_Inclination_Curve_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowThird.gridLayout_3.addWidget(self.canvas_Back_Section_Inclination_Curve_map)
        self.axes6 = self.figure_Back_Section_Inclination_Curve_map.add_subplot(111)

        #导入各个截段的中点
        all_middle_points = np.loadtxt(global_g_strDataFolderNameOutHospital + r'\各个截段点及中点位置' + r'\all_middle_points.txt')#os.getcwd()
        #导入各个截段的侧旋角
        # os.getcwd()获取当前工作路径
        Side_spin_angle = np.loadtxt(global_g_strDataFolderNameOutHospital + r'\各个截段点及中点位置' + r'\Side_spin_angle.txt')#os.getcwd()
        # 画截面曲线
        # plt.plot(Side_spin_angle[0:5], all_middle_points[0:5, 1], marker='.',markersize=5, color='yellow')
        # plt.plot(Side_spin_angle[5:9], all_middle_points[5:9, 1], marker='.', markersize=5, color='k')
        # plt.plot(Side_spin_angle[9:19], all_middle_points[9:19, 1], marker='.', markersize=5,color='lime')
        # plt.plot(Side_spin_angle[19:28], all_middle_points[19:28, 1], marker='.', markersize=5,color='blue')
        # 画过零点直线
        plt.plot([minx, maxx], [0, 0], color='slategray')  # lavender
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        plt.plot(Side_spin_angle[0:6], all_middle_points[0:6, 1], marker='.', markersize=5, color='yellow')
        plt.plot(Side_spin_angle[5:10], all_middle_points[5:10, 1], marker='.', markersize=5, color='k')
        plt.plot(Side_spin_angle[9:20], all_middle_points[9:20, 1], marker='.', markersize=5, color='lime')
        plt.plot(Side_spin_angle[19:28], all_middle_points[19:28, 1], marker='.', markersize=5, color='blue')

        # self.axes6.axis('equal')  # 等比例显示
        self.axes6.set_xlim((min(Side_spin_angle-0.5), max(Side_spin_angle+0.5)))
        self.axes6.set_ylim((miny, maxy))
        self.axes6.set_xticks(np.linspace(min(Side_spin_angle-0.5),max(Side_spin_angle+0.5), 8))
        self.axes6.set_yticks(np.linspace(miny, maxy, 15))
        self.axes6.set_xlabel('后背截面倾角（侧旋角）(°)', fontproperties="SimHei")
        self.axes6.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\后背截面倾角曲线图.png', format='png')

        # </editor-fold>

        # <editor-fold desc="左右侧屈脊柱曲线图axes8...">
        #20190506
        try:
            YH_Result_dynamic.dynamic_fMarkerPoint(g_strDataFolderNameOutHospital, attitude)
        except BaseException:
            print('动态数据获取时发生错误')
            pass
        else:
            # 20190506
            '''左右侧屈脊柱曲线图...'''
            try:
                sip.delete(self.canvas_flexion_spine_Curve_map)  # 删除addwidget的窗口
            except BaseException:
                pass
            self.figure_flexion_spine_Curve_map = plt.figure()
            self.canvas_flexion_spine_Curve_map = FigureCanvas(self.figure_flexion_spine_Curve_map)
            self.toolbar = NavigationToolbar(self.canvas_flexion_spine_Curve_map, self)  # 窗口操作条
            self.Ui_Result_dynamic.gridLayout_2.addWidget(self.toolbar)
            self.Ui_Result_dynamic.gridLayout_2.addWidget(self.canvas_flexion_spine_Curve_map)
            self.axes8 = self.figure_flexion_spine_Curve_map.add_subplot(111)
            filepath = r'D:\B\zuoye\postgraduate\SecondGradePostgraduateTheFirstHalfOfTheSemester\！数据&任务\27+\直写式20190218'
            all_fMarkerPoint_name = ['Standing_Posture_Right_Limit_fMarkerPoint',
                                     'Standing_Posture_Left_Limit_fMarkerPoint',
                                     'Standing_Posture_Right_fMarkerPoint', 'Standing_Posture_Left_fMarkerPoint']  # 文件夹名称
            color = ['k', 'k', 'lime', 'r']  # lime为绿色
            for k in range(len(all_fMarkerPoint_name)):
                Spine_Fitting_Curve = np.loadtxt(
                    g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Spine_Fitting_Curve_' + all_fMarkerPoint_name[k] + '.txt')  # filepath
                plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                         markersize=3, color=color[k])  # 画脊柱拟合曲线
                # 计算边界值画图用到
                if k==1:
                    left_boundary = np.min(Spine_Fitting_Curve[:, 0])
                if k==0:
                    right_boundary = np.max(Spine_Fitting_Curve[:, 0])
            plt.grid(ls='--', linewidth=0.75)  # 生成网格
            Spine_Fitting_Curve = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + r'\Spine_Fitting_Curve.txt')  # filepath
            plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                     markersize=3, color='b')  # 画脊柱拟合曲线
            # 画过零点直线
            plt.plot([left_boundary-20, right_boundary+20], [0, 0], color='slategray')  # lavender
            self.axes8.axis('equal')  # 等比例显示
            self.axes8.set_xlim((left_boundary-20, right_boundary+20))
            self.axes8.set_ylim((miny, maxy))
            # self.axes.set_zlim((minz-200, maxz+200))
            self.axes8.set_xticks(np.linspace(left_boundary, right_boundary, 6))
            self.axes8.set_yticks(np.linspace(miny, maxy, 15))
            # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
            self.axes8.set_xlabel('水平方向(mm)', fontproperties="SimHei")
            self.axes8.set_ylabel('身高方向(mm)', fontproperties="SimHei")
            plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\左右侧屈脊柱曲线图.png', format='png')

        # </editor-fold>

            # <editor-fold desc="左右侧屈15水平方向差值和（均值）axes9...">
            '''左右侧屈15水平方向差值和（均值）...'''

            #############20190506
            all_fMarkerPoint_name = ['Standing_Posture_Right_fMarkerPoint', 'Standing_Posture_Left_fMarkerPoint']  # 文件夹名称
            Spine_Fitting_Curve_Right = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Spine_Fitting_Curve_' + all_fMarkerPoint_name[0] + '.txt')#右侧屈15脊柱曲线#filepath
            Spine_Fitting_Curve_Left = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Spine_Fitting_Curve_' + all_fMarkerPoint_name[1] + '.txt')#左侧屈15脊柱曲线#filepath
            Pick_points_ind = np.linspace(0,len(Spine_Fitting_Curve_Right)-1,30)
            Pick_points_ind = list(np.round(Pick_points_ind))#四舍五入取整
            for i in range(len(Pick_points_ind)):
                Pick_points_ind[i] = math.floor(Pick_points_ind[i])

            ###20190508之前（利用脊柱曲线上的均匀点进行计算）
            # Pick_points_dif_x = Spine_Fitting_Curve_Right[Pick_points_ind, 0] + Spine_Fitting_Curve_Left[Pick_points_ind, 0]#对称差值，而不是值的差值
            # Pick_points_dif_y = Spine_Fitting_Curve_Right[Pick_points_ind, 1]
            # plot_Left_and_right_difference = np.hstack((Pick_points_dif_x[:, np.newaxis], Pick_points_dif_y[:, np.newaxis]))
            # np.savetxt(g_strDataFolderNameOutHospital + '/动态标记点及脊柱曲线' + '\plot_Left_and_right_difference.txt', plot_Left_and_right_difference,
            #            fmt='%.2f')#filepath
            ###20190508之前（利用脊柱曲线上的均匀点进行计算）
            ###20190508（直接利用脊柱上标记点进行计算）
            Spine_Right_fMarkerPoint = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Correction_Standing_Posture_Right_fMarkerPoint.txt')  # 右侧屈15脊柱曲线#filepath
            Spine_Left_fMarkerPoint = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Correction_Standing_Posture_Left_fMarkerPoint.txt')  # 左侧屈15脊柱曲线#filepath
            Pick_points_dif_x = abs(Spine_Left_fMarkerPoint[:-4, 0]) - abs(Spine_Right_fMarkerPoint[:-4, 0])#对称差值，而不是值的差值
            Pick_points_dif_y = abs(Spine_Left_fMarkerPoint[:-4, 1]) - abs(Spine_Right_fMarkerPoint[:-4, 1])
            Pick_points_y = Spine_Right_fMarkerPoint[:-4, 1]
            plot_Left_and_right_difference = np.hstack((Pick_points_dif_x[:, np.newaxis], Pick_points_dif_y[:, np.newaxis],Pick_points_y[:, np.newaxis]))
            np.savetxt(g_strDataFolderNameOutHospital + '/动态标记点及脊柱曲线' + '\plot_Left_and_right_difference.txt', plot_Left_and_right_difference,
                       fmt='%.2f')#filepath
            ###20190508（直接利用脊柱上标记点进行计算）


            #####保存画图倾角变化数据
            Pick_points_right_x = Spine_Fitting_Curve_Right[Pick_points_ind, 0]
            Pick_points_left_x = Spine_Fitting_Curve_Left[Pick_points_ind, 0]
            Pick_points_right_y = Spine_Fitting_Curve_Right[Pick_points_ind, 1]
            Pick_points_left_y = Spine_Fitting_Curve_Left[Pick_points_ind, 1]

            Lateral_buckling_right_radian = []
            Lateral_buckling_left_radian = []
            Lateral_buckling_right_angle = []
            Lateral_buckling_left_angle = []
            for i in range(len(Pick_points_ind) - 1):  #################################
                # 弧度
                Lateral_buckling_right_radian.append(math.atan((Pick_points_right_x[i + 1] - Pick_points_right_x[i]) / (
                        Pick_points_right_y[i + 1] - Pick_points_right_y[i])))
                Lateral_buckling_left_radian.append(math.atan((Pick_points_left_x[i + 1] - Pick_points_left_x[i]) / (
                        Pick_points_left_y[i + 1] - Pick_points_left_y[i])))
                # 角度
                Lateral_buckling_right_angle.append(Lateral_buckling_right_radian[i] * 180 / np.pi)  # 实际偏离角度
                Lateral_buckling_left_angle.append(Lateral_buckling_left_radian[i] * 180 / np.pi)
            Lateral_buckling_right_angle_diff = [0]  # 偏离角度变化
            Lateral_buckling_left_angle_diff = [0]
            for i in range(len(Lateral_buckling_right_angle) - 1):
                Lateral_buckling_right_angle_diff.append(
                    Lateral_buckling_right_angle[i + 1] - Lateral_buckling_right_angle[i])
                Lateral_buckling_left_angle_diff.append(
                    Lateral_buckling_left_angle[i + 1] - Lateral_buckling_left_angle[i])

            Lateral_buckling_right_angle_diff = np.array(Lateral_buckling_right_angle_diff)
            Lateral_buckling_left_angle_diff = np.array(Lateral_buckling_left_angle_diff)
            plot_Lateral_buckling_angle_change = np.hstack((Lateral_buckling_left_angle_diff[:, np.newaxis],
                                                            Pick_points_left_y[1:][:, np.newaxis],
                                                            Lateral_buckling_right_angle_diff[:, np.newaxis],
                                                            Pick_points_right_y[1:][:, np.newaxis]))
            np.savetxt(g_strDataFolderNameOutHospital + '/动态标记点及脊柱曲线' + '\plot_Lateral_buckling_angle_change.txt',
                       plot_Lateral_buckling_angle_change, fmt='%.2f')  # filepath
            # print('保存画图倾角变化数据有问题')
            #############20190506
            try:
                sip.delete(self.canvas_flexion_spine_Curve_different_map)  # 删除addwidget的窗口
            except BaseException:
                pass
            self.figure_flexion_spine_Curve_different_map = plt.figure()
            self.canvas_flexion_spine_Curve_different_map = FigureCanvas(self.figure_flexion_spine_Curve_different_map)
            self.toolbar = NavigationToolbar(self.canvas_flexion_spine_Curve_different_map, self)#窗口操作条
            self.Ui_Result_dynamic.gridLayout_3.addWidget(self.toolbar)
            self.Ui_Result_dynamic.gridLayout_3.addWidget(self.canvas_flexion_spine_Curve_different_map)
            self.axes9 = self.figure_flexion_spine_Curve_different_map.add_subplot(111)
            filepath = r'D:\B\zuoye\postgraduate\SecondGradePostgraduateTheFirstHalfOfTheSemester\！数据&任务\27+\直写式20190218'
            plt.grid(ls='--', linewidth=0.75)  # 生成网格
            # plot_Left_and_right_difference = np.loadtxt(g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线'+ r'\plot_Left_and_right_difference.txt')#filepath
            # plt.plot(plot_Left_and_right_difference[:, 0], plot_Left_and_right_difference[:, 2], marker='.',
            #          markersize=1, color='r')#x差值
            # plt.plot(plot_Left_and_right_difference[:, 1], plot_Left_and_right_difference[:, 2], marker='.',
            #          markersize=1, color='lime')#y差值
            plt.plot(plot_Left_and_right_difference[:, 0], plot_Left_and_right_difference[:, 2], marker='*',
                     markersize=5, color='r')#x差值
            plt.plot(plot_Left_and_right_difference[:, 1], plot_Left_and_right_difference[:, 2], marker='*',
                     markersize=5, color='lime')#y差值
            # 画过零点直线
            plt.plot([min(min(plot_Left_and_right_difference[:, 0]),min(plot_Left_and_right_difference[:, 1])),
                      max(max(plot_Left_and_right_difference[:, 0]),max(plot_Left_and_right_difference[:, 1]))],
                      [0, 0], color='slategray')  # lavender
            # self.axes9.axis('equal')  # 等比例显示
            self.axes9.set_xlim((min(min(plot_Left_and_right_difference[:, 0]),min(plot_Left_and_right_difference[:, 1]))-3,
                                 max(max(plot_Left_and_right_difference[:, 0]),
                                     max(plot_Left_and_right_difference[:, 1])) + 3))
            self.axes9.set_ylim((miny, maxy))
            # self.axes.set_zlim((minz-200, maxz+200))
            self.axes9.set_xticks(np.linspace(min(plot_Left_and_right_difference[:, 0]), max(plot_Left_and_right_difference[:, 0]), 5))
            self.axes9.set_yticks(np.linspace(miny, maxy, 15))
            # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
            self.axes9.set_xlabel('水平方向(mm)', fontproperties="SimHei")
            self.axes9.set_ylabel('身高方向(mm)', fontproperties="SimHei")
            plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\左右侧屈相同脊柱位置差值.png', format='png')

            # </editor-fold>

            # <editor-fold desc="脊柱曲线相邻点倾角变换axes10...">
            '''脊柱曲线相邻点倾角变换...'''
            try:
                sip.delete(self.canvas_spine_curve_Inclination_change_map)  # 删除addwidget的窗口
            except BaseException:
                pass
            self.figure_spine_curve_Inclination_change_map = plt.figure()
            self.canvas_spine_curve_Inclination_change_map = FigureCanvas(self.figure_spine_curve_Inclination_change_map)
            self.toolbar = NavigationToolbar(self.canvas_spine_curve_Inclination_change_map, self)#窗口操作条
            self.Ui_Result_dynamic.gridLayout_4.addWidget(self.toolbar)
            self.Ui_Result_dynamic.gridLayout_4.addWidget(self.canvas_spine_curve_Inclination_change_map)
            self.axes10 = self.figure_spine_curve_Inclination_change_map.add_subplot(111)
            plt.grid(ls='--', linewidth=0.75)  # 生成网格
            # plot_Lateral_buckling_angle_change = np.loadtxt(g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线'+ r'\plot_Lateral_buckling_angle_change.txt')#filepath
            #导入左右侧屈15°的倾角变换
            plt.plot(plot_Lateral_buckling_angle_change[:, 0], plot_Lateral_buckling_angle_change[:, 1], marker='*',
                     markersize=5, color='r')
            plt.plot(plot_Lateral_buckling_angle_change[:, 2], plot_Lateral_buckling_angle_change[:, 3], marker='*',
                     markersize=5, color='lime')
            # 标识出最大位置的点
            ind_max_abs_x_left = np.where(abs(plot_Lateral_buckling_angle_change[:, 0])==max(abs(plot_Lateral_buckling_angle_change[:, 0])))
            ind_max_abs_x_right = np.where(
                abs(plot_Lateral_buckling_angle_change[:, 2]) == max(abs(plot_Lateral_buckling_angle_change[:, 2])))
            plt.plot(plot_Lateral_buckling_angle_change[ind_max_abs_x_left, 0], plot_Lateral_buckling_angle_change[ind_max_abs_x_left, 1], marker='*',
                     markersize=10, color='k')
            plt.plot(plot_Lateral_buckling_angle_change[ind_max_abs_x_right, 2], plot_Lateral_buckling_angle_change[ind_max_abs_x_right, 3], marker='*',
                     markersize=10, color='k')
            # 画过零点直线
            plt.plot([min(plot_Lateral_buckling_angle_change[:, 0]), max(plot_Lateral_buckling_angle_change[:, 2])], [0, 0], color='slategray')  # lavender
            # self.axes9.axis('equal')  # 等比例显示
            self.axes10.set_xlim((min(plot_Lateral_buckling_angle_change[:, 0]-3), max(plot_Lateral_buckling_angle_change[:, 2]+3)))
            self.axes10.set_ylim((miny, maxy))
            self.axes10.set_xticks(np.linspace(min(plot_Lateral_buckling_angle_change[:, 0]), max(plot_Lateral_buckling_angle_change[:, 2]), 5))
            self.axes10.set_yticks(np.linspace(miny, maxy, 15))
            self.axes10.set_xlabel('水平方向(mm)', fontproperties="SimHei")
            self.axes10.set_ylabel('身高方向(mm)', fontproperties="SimHei")
            plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\脊柱曲线相邻点倾角变化.png', format='png')

        # </editor-fold>

        # <editor-fold desc="冠状面C弯axes7...">
        '''冠状面C弯'''
        try:
            sip.delete(self.canvas_Coronal_C_bend_map)  # 删除addwidget的窗口
        except BaseException:
            pass
        self.figure_Coronal_C_bend_map = plt.figure()
        self.canvas_Coronal_C_bend_map = FigureCanvas(self.figure_Coronal_C_bend_map)
        # self.toolbar = NavigationToolbar(self.canvas_Coronal_C_bend_map, self)#窗口操作条
        # self.Ui_Coronal_c_bend.gridLayout_2.addWidget(self.toolbar)
        self.Ui_Coronal_c_bend.gridLayout_2.addWidget(self.canvas_Coronal_C_bend_map)
        self.axes7 = self.figure_Coronal_C_bend_map.add_subplot(111)
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画过零点直线
        plt.plot([minx - 100, maxx + 100], [0, 0], color='slategray')  # lavender
        plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                 markersize=3, color='b')  # 画脊柱拟合曲线
        xnew = Spine_Fitting_Curve[:,1]
        ynew = Spine_Fitting_Curve[:,0]

        #     预测点！#计算边界值（极值点位置）
        First_derivative = [(ynew[1] - ynew[0])/(xnew[1] - xnew[0])]#为了减少循环，先给一阶导初值赋第一个值
        Second_derivative = []
        for i in range(len(xnew)-2):
            dif_y = ynew[i+2] - ynew[i+1]
            dif_x = xnew[i+2] - xnew[i+1]
            First_derivative.append(dif_y/dif_x)#身体上下方向为x，以上下方向作为自变量计算
            dif_y2 = First_derivative[i+1] - First_derivative[i]
            Second_derivative.append(dif_y2 / dif_x)
        ind_Inflection_point = np.where(abs(np.array(Second_derivative))<0.0002)#脊柱曲线的拐点初筛
        ind_Inflection_point = ind_Inflection_point[0]
        correct_ind_Inflection_point = []#最后的拐点的初始化
        start_ind = 0
        for i in range(len(ind_Inflection_point)-1):#遍历所有可能的拐点
        #     if count_Inflection_point_num == 1 :
            if ((ind_Inflection_point[i+1] - ind_Inflection_point[i]) >= 2):#找到除最后一个拐点的拐点
                correct_ind_Inflection_point.append(math.ceil(np.mean(ind_Inflection_point[start_ind:i+1])))
                start_ind = i+1
            if (i == (len(ind_Inflection_point)-2)):#补最后一个拐点
                correct_ind_Inflection_point.append(math.ceil(np.mean(ind_Inflection_point[start_ind:])))
        # correct_ind_Inflection_point = [0] + correct_ind_Inflection_point  # 增加最下面点为拐点
        correct_ind_Inflection_point = correct_ind_Inflection_point + [len(ynew)-1] # 增加最上面面点为拐点
        print(correct_ind_Inflection_point)
        plt.plot(ynew[correct_ind_Inflection_point], xnew[correct_ind_Inflection_point], '.', marker='.', markersize=8,  color='r')

        # <editor-fold desc="画拟合圆">
        ###########画拟合圆
        all_rad = []#所有拟合圆半径初始化
        all_centerx = []#所有拟合圆圆心x值初始化
        all_centery = []#所有拟合圆圆心y值初始化
        for i in range(len(correct_ind_Inflection_point)-1):#得到并画出每个c弯的拟合圆
            points = Spine_Fitting_Curve[correct_ind_Inflection_point[i]:correct_ind_Inflection_point[i + 1], :]#获取俩个拐点间的脊柱曲线点
            centerx, centery, rad = self.circleLeastFit(points)#调用获取拟合圆半径和圆心坐标的函数
            print(centerx, centery, rad)
            all_rad.append(rad)
            all_centerx.append(centerx)
            all_centery.append(centery)
            # 方法一：参数方程
            theta = np.arange(0, 2 * np.pi, 0.01)#设置参数方程自变量角度
            # range()返回的是range object，而np.nrange()返回的是numpy.adarray()
            # 两者都是均匀地（evenly）等分区间；
            # range尽可用于迭代，而np.arange作用远不止于此，它是一个序列，可被当做向量使用。
            # range()不支持步长为小数，np.arange()支持步长为小数
            x = centerx + rad * np.cos(theta)  # 圆的参数方程
            y = centery + rad * np.sin(theta) # 圆的参数方程
            # plt.plot(x, y, c='green', ls='-', lw=1)#画拟合圆!!
            plt.axis('equal')#等比例显示
        ###########画主弯拟合圆
        ind_Main_bend = np.where(all_rad == np.min(all_rad))#找到曲率半径最小（即曲率最大）的c弯作为主弯
        ind_Main_bend = ind_Main_bend[0][0]
        print('ind_Main_bend ',ind_Main_bend)
        ind_Main_bend_Boundary_points = correct_ind_Inflection_point[ind_Main_bend:ind_Main_bend+2]#主弯俩个拐点的的索引
        plt.plot(ynew[ind_Main_bend_Boundary_points], xnew[ind_Main_bend_Boundary_points], '.', marker='.', markersize=18,
                 color='m')#显示主弯的俩个端点
        Main_bend_rad = all_rad[ind_Main_bend]
        Main_bend_centerx = all_centerx[ind_Main_bend]
        Main_bend_centery = all_centery[ind_Main_bend]
        plt.plot(Main_bend_centerx, Main_bend_centery, '.', marker='.',
                 markersize=18,color='r')  # 显示主弯的圆心
        Main_bend_center = [Main_bend_centerx, Main_bend_centery]  # 全局变量
        Main_bend_Boundary_points = [ynew[ind_Main_bend_Boundary_points], xnew[ind_Main_bend_Boundary_points]]  # 全局变量
        theta = np.arange(0, 2 * np.pi, 0.01)  # 设置参数方程自变量角度
        x = Main_bend_centerx + Main_bend_rad * np.cos(theta)  # 圆的参数方程
        y = Main_bend_centery + Main_bend_rad * np.sin(theta)  # 圆的参数方程
        plt.plot(x, y, c='r', ls='-', lw=1)  # 画主弯拟合圆
        # </editor-fold>


        print('ind_Inflection_point=',ind_Inflection_point)
#######################
        self.axes7.axis('equal')  # 等比例显示
        self.axes7.set_xlim((minx, maxx))
        self.axes7.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes7.set_xticks(np.linspace(minx, maxx, 8))
        self.axes7.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes7.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes7.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\C弯选取示意图.png', format='png')  # 20190622
        # </editor-fold>

#   12冠状面C1弯曲率半径
        Main_bend_rad_to_save = Main_bend_rad
        print('12冠状面C1弯曲率半径=', Main_bend_rad_to_save)
#   13冠状面C1弯cobb角（右弯为正）
#         | a |=√[x1 ^ 2 + y1 ^ 2]
#         | b |=√[x2 ^ 2 + y2 ^ 2]
#         a * b = (x1, y1)(x2, y2) = x1x2 + y1y2
#         cos = a * b / [ | a | * | b |]
#         =(x1x2 + y1y2) / [√[x1 ^ 2 + y1 ^ 2] *√[x2 ^ 2 + y2 ^ 2]]
        vector_1 = [Main_bend_Boundary_points[0][0] - Main_bend_center[0],Main_bend_Boundary_points[1][0] - Main_bend_center[1]]
        vector_2 = [Main_bend_Boundary_points[0][1] - Main_bend_center[0],Main_bend_Boundary_points[1][1] - Main_bend_center[1]]
        cobber_radian = math.acos((vector_1[0] * vector_2[0] + vector_1[1] * vector_2[1])/(Main_bend_rad**2))
        cobber = cobber_radian *180/math.pi
        print('冠状面C1弯cobb角（右弯为正）=', cobber)

        # 20190611
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setFrame(False)  # 边框设置
        # self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setClearButtonEnabled(True)#删除按钮
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setAlignment(QtCore.Qt.AlignCenter)  # 居中
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setStyleSheet("color:green")
        font = QtGui.QFont()
        font.setFamily('微软雅黑')
        font.setBold(True)
        font.setPointSize(15)
        font.setWeight(75)
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setFont(font)

        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setFrame(False)  # 边框设置
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setAlignment(QtCore.Qt.AlignCenter)  # 居中
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setStyleSheet("color:green")
        font = QtGui.QFont()
        font.setFamily('微软雅黑')
        font.setBold(True)
        font.setPointSize(15)
        font.setWeight(75)
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setFont(font)

        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setText(str(Main_bend_rad_to_save))
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setText(str(cobber))

        print('注意正负')
        # Main_bend_center, Main_bend_Boundary_points, Main_bend_rad
    #   14侧弯类型（s,c,多c）
        Side_bend_type_character = ['C弯','C弯','S弯','多C弯']
        Side_bend_number = len(finalCBendfBKPtCor)
        if Side_bend_number > 3:
            Side_bend_number = 3
        Side_bend_type = Side_bend_type_character[Side_bend_number]
        print('14侧弯类型（s,c,多c）=', Side_bend_type)
        # 显示后背点云图完成提醒窗口
        reply = QMessageBox.information(self, '提示', '显示指标示意图完毕', QMessageBox.Yes)

    def Redraw_clouds_point(self):
        global minx, maxx, miny, maxy, Spine_Fitting_Curve
        global final_fMarkerPoint_end
        global final_points
        global Sampling_points,fMarkerPoint
        global minz, maxz, z
        import sip
        try:
            sip.delete(self.canvas_back_cloud_points_map)  # 删除addwidget的窗口
        except BaseException:
            root = Tk()
            root.withdraw()  # ****实现tk主窗口隐藏
            messagebox.showinfo("提示", "没有绘制后背等高线点云图！")
        # <editor-fold desc="画后背点云图...">
        '''画后背点云图Ui_ResultShowFirst CResultShowFirst'''
        self.figure_back_cloud_points_map = plt.figure()

        self.canvas_back_cloud_points_map = FigureCanvas(self.figure_back_cloud_points_map)
        # self.toolbar = NavigationToolbar(self.canvas_contour_map, self)#窗口操作条
        # self.Ui_DataPreprocessing.gridLayout_2.addWidget(self.toolbar)
        self.Ui_ResultShowFirst.gridLayout_2.addWidget(self.canvas_back_cloud_points_map)
        Number_of_segments = 180  # 将点云分割为800段#值越大颜色变化越快
        Contour_value = np.linspace(minz, maxz, Number_of_segments)  # 一共获取12个数值
        ind = []  # ind为各个段的位置索引，初始化
        colormap = []  # 颜色数组colormap，初始化
        Color_step = 10  # 值越大颜色变化越快
        for i in range(Number_of_segments - 1):  # range(600)为0-599颜色映射
            Transition_variable = np.where(
                (Sampling_points[2] >= Contour_value[i]) & (Sampling_points[2] < Contour_value[i + 1]))
            ind.append(Transition_variable)  # ind为各个段的位置索引

            ## 定义颜色数组colormap，随z值变化而变化的颜色
            # 111->101->100->110->010->011->001->000
            #        if (abs((i)%15) <= 1):#绘制黑色等高线
            #            color_R = 0
            #            color_G = 0
            #            color_B = 0

            if (((Color_step * i) // 255) == 0):
                color_R = 230 / 255  # 220 / 255#1本来应该为1，为了把最开始的白色亮度调低，而减小这个值
                color_G = (255 - (Color_step * i) % 255) / 255
                color_B = 230 / 255  # 1本来应该为1，为了把最开始的白色亮度调低，而减小这个值
            elif (((Color_step * i) // 255) == 1):
                color_R = 1  # 220 / 255
                color_G = 0
                color_B = (255 - (Color_step * i) % 255) / 255
            elif (((Color_step * i) // 255) == 2):
                color_R = 1  # 220 / 255
                color_G = ((Color_step * i) % 255) / 255
                color_B = 0
            elif (((Color_step * i) // 255) == 3):
                color_R = (255 - (Color_step * i) % 255) / 255
                color_G = 1
                color_B = 0
            elif (((Color_step * i) // 255) == 4):
                color_R = 0
                color_G = 1
                color_B = ((Color_step * i) % 255) / 255
            elif (((Color_step * i) // 255) == 5):
                color_R = 0
                color_G = (255 - (Color_step * i) % 255) / 255
                color_B = 1
            elif (((Color_step * i) // 255) == 6):
                color_R = 0
                color_G = 0
                color_B = (255 - (Color_step * i) % 255) / 255
            else:  # 白色111（不显示）
                color_R = 0
                color_G = 0
                color_B = 0

            color = (color_R, color_G, color_B)
            colormap.append(color)
            # exec(index + '=Transition_variable')#字符串变为变量名

        for i in reversed(range(Number_of_segments - 1)):  # 循环画图，reversed反转数组
            # for i in range(Number_of_segments - 1):  # 循环画图，reversed反转数组
            plt.plot(Sampling_points[0][ind[i][0]], Sampling_points[1][ind[i][0]], '.', marker='.', markersize=2,
                     color=colormap[i])  # , markersize为点的大小
        ##画脊柱曲线
        # Spine_Fitting_Curve = np.loadtxt(global_g_strDataFolderNameOutHospital + '/Spine_Fitting_Curve.txt')
        plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                 markersize=3, color='b')
        ## 画等高线
        min_Contour_value = -100
        max_Contour_value = 300
        try:
            Contour_spacing = self.Contour_spacing_lineEdit.text()
            Contour_line_width = self.Contour_line_width_lineEdit.text()
            Contour_spacing = float(Contour_spacing)
            Contour_line_width = float(Contour_line_width)
        except BaseException:
            Contour_spacing = 5
            Contour_line_width = 0.8
#        if Contour_spacing == None:# == []:
#            Contour_spacing = 5
#        if Contour_line_width == None:# == []:
#            Contour_line_width = 0.8



        num_Contour_value = (max_Contour_value - min_Contour_value)/Contour_spacing + 1

        Contour_value_contour = np.linspace(min_Contour_value,max_Contour_value,num_Contour_value)  # -60, 60取25个数，使得以5为间隔
        # Contour_value_contour = np.linspace(-100, 300,81)  # -60, 60取25个数，使得以5为间隔
        # ind_contour = []
        for i in Contour_value_contour:
            # ind_contour = np.where((z >= i) & (z < i + 0.8))  # 非常要注意这个括号 没有括号估计内部执行顺序不对，捣腾不出来的，具体原因可评论留言
            ind_contour = np.where((z >= i) & (z < i + Contour_line_width))  # 非常要注意这个括号 没有括号估计内部执行顺序不对，捣腾不出来的，具体原因可评论留言
            plt.plot(Sampling_points[0][ind_contour[0]], Sampling_points[1][ind_contour[0]], '.', marker='.',
                     markersize=0.8, color='k')
        # 生成网格
        # plt.grid(which='major', axis='x', linewidth=0.75, linestyle='-', color='r')
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画标记点
        plt.plot(fMarkerPoint[:, 0], fMarkerPoint[:, 1], '.', marker='*', markersize=5,
                 color='b')  # , markersize为点的大小fMarkerPoint_y[2,:],fMarkerPoint_x[2,:]
        # 画过零点直线
        plt.plot([minx - 100, maxx + 100], [0, 0], color='slategray')  # lavender
        # plt.plot( [0, 0], [miny, maxy],color='slategray')#竖直方向
        # 坐标显示操作
        # self.axes.view_init(elev=-90, azim=0)
        self.axes1.axis('equal')  # 等比例显示
        self.axes1.set_xlim((minx, maxx))
        self.axes1.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes1.set_xticks(np.linspace(minx, maxx, 8))
        self.axes1.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes1.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes1.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        root = Tk()
        root.withdraw()  # ****实现tk主窗口隐藏
        messagebox.showinfo("提示", "重绘后背等高线点云图完毕！")
        # </editor-fold>

    def save_data(self):
        import datetime
        from PyQt5.QtWidgets import QMessageBox
        import numpy as np
        global g_strDataFolderNameOutHospital
        global Case_number
        # print('姓名：', self.DlgPatientInfoInput.strPatientName.text())
        # print('性别：', self.DlgPatientInfoInput.strPatientGender.text())
        # print('年龄：', self.DlgPatientInfoInput.strPatientAge.text())
        # print('身高：', self.DlgPatientInfoInput.strPatientHeight.text())
        # print('体重：', self.DlgPatientInfoInput.strPatientWeight.text())
        # print('职业：', self.DlgPatientInfoInput.strPatientCareer.text())
        # print('身份证号：', self.DlgPatientInfoInput.strPatientPersonalID.text())
        # print('病历号：', self.DlgPatientInfoInput.strPatientHospitalID.text())
        # print('检测时间：', self.DlgPatientInfoInput.strPatientCheckTime.text())
        # print('常住住址：', self.DlgPatientInfoInput.strPatientAddress.text())
        # print('送检医生：', self.DlgPatientInfoInput.strPatientDoctor.text())
        # print('检测医生：', self.DlgPatientInfoInput.strChecker.text())
        # print('初潮年龄：', self.DlgPatientInfoInput.strPatientMenarcheAge.text())
        nowTime = datetime.datetime.now().strftime('%Y-%m-%d-%H-%M-%S')  # 现在
        # name_code_time = self.DlgPatientInfoInput.strPatientName.text() + self.DlgPatientInfoInput.strPatientPersonalID.text() + nowTime
        name_code_time = self.DlgPatientInfoInput.strPatientName.text() + self.DlgPatientInfoInput.strPatientHospitalID.text() + nowTime
        if self.DlgPatientInfoInput.strPatientName.text()=='' or self.DlgPatientInfoInput.strPatientHospitalID.text()=='':
            reply = QMessageBox.information(self, '提示', '请至少输入姓名或病历号', QMessageBox.Yes)
        else:
            # if self.DlgPatientInfoInput.strPatientPersonalID.text()=='':
            #     print('没填姓名')
            # np.savetxt("D:/灵犀脊柱形态评价系统/patientinfornow.txt",[name_code_time],fmt='%s')  #信息录入的文件名
            np.savetxt(r"D:\灵犀脊柱形态评价系统\systemfile/patientinfo.txt", [name_code_time], fmt='%s')  # 信息录入的文件名

            dirs = r'D:\灵犀脊柱形态评价系统\基本信息/'
            if not os.path.exists(dirs):  # 不存在该路径则新建
                os.makedirs(dirs)
            filepathorl = r'D:\灵犀脊柱形态评价系统\基本信息/' + self.DlgPatientInfoInput.strPatientName.text() + self.DlgPatientInfoInput.strPatientHospitalID.text() + nowTime
            filepath = r'D:\灵犀脊柱形态评价系统\基本信息/' + self.DlgPatientInfoInput.strPatientName.text() + self.DlgPatientInfoInput.strPatientHospitalID.text() + nowTime + '/informationsave.txt'
            print('生成的文件的具体位置为：', filepath)

            try:
                # np.savetxt(g_strDataFolderNameOutHospital + r"/被测者基本信息路径.txt", [filepath], fmt='%s')  # 信息录入的文件名
                #20190617
                np.savetxt(r'D:\灵犀脊柱形态评价系统\当前被测者基本信息路径.txt', [filepath], fmt='%s')  # 信息录入的文件名
                # np.savetxt(r'D:\灵犀脊柱形态评价系统\当前被测者基本信息路径.txt', [filepath], fmt='%s')  # 信息录入的文件名20190622
            except BaseException:
                root = Tk()
                root.withdraw()  # ****实现tk主窗口隐藏
                messagebox.showinfo("提示", "请先录入被测者基本信息！")
            else:
                print('filepathorl为：', filepathorl)
                if not os.path.exists(filepathorl):  # 不存在则新建
                    os.makedirs(filepathorl)
                f = open(filepath, "w+")
                           # f = open(r"C:\Users\YH\Desktop\临时\datatest.txt", "w+")
                datasave =  ['姓名：' + self.DlgPatientInfoInput.strPatientName.text() + "\n",          '性别：' + self.DlgPatientInfoInput.strPatientGender.text() + "\n",
                             '年龄：' + self.DlgPatientInfoInput.strPatientAge.text() + "\n",           '身高：' + self.DlgPatientInfoInput.strPatientHeight.text() + "\n",
                             '体重：' + self.DlgPatientInfoInput.strPatientWeight.text() + "\n",        '职业：' + self.DlgPatientInfoInput.strPatientCareer.text() + "\n",
                             '身份证号：' + self.DlgPatientInfoInput.strPatientPersonalID.text() + "\n", '病历号：' + self.DlgPatientInfoInput.strPatientHospitalID.text() + "\n",
                             '检测时间：' + self.DlgPatientInfoInput.strPatientCheckTime.text() + "\n",  '常住住址：' + self.DlgPatientInfoInput.strPatientAddress.text() + "\n",
                             '送检医生：' + self.DlgPatientInfoInput.strPatientDoctor.text() + "\n",     '检测医生：' + self.DlgPatientInfoInput.strChecker.text() + "\n",
                             '初潮年龄：' + self.DlgPatientInfoInput.strPatientMenarcheAge.text() + "\n",'肺活量：' + self.DlgPatientInfoInput.iPatientVital_capacity.text()]  # 初始化
                f.writelines(datasave)
                # np.savetxt(filepath, datasave)  # fmt="%f,%f,%f", fmt="%.2f %.2f %.2f %.0f %.0f %.0f"
                f.close()
                # 保存完成提醒窗口
                # reply = QMessageBox.information(self, '提示', '保存基本信息到txt完毕', QMessageBox.Yes)

                # <editor-fold desc="information存入mysql...">
            # # 使用 execute() 方法执行 SQL，如果表存在则删除
            # cursor.execute("DROP TABLE IF EXISTS patientsinformation")

            # ①查询出有多条数据时：
            # cursor.fetchone()：将只取最上面的第一条结果，返回单个元组如('id', 'name')，然后多次循环使用cursor.fetchone()，依次取得下一条结果，直到为空。
            # cursor.fetchall(): 将返回所有结果，返回二维元组，如(('id', 'name'), ('id', 'name')),
            # ②查询只有一条数据时：
            # cursor.fetchone()：将只返回一条结果，返回单个元组如('id', 'name')。
            # cursor.fetchall(): 也将返回所有结果，返回二维元组，如(('id', 'name'), )

            # 1 新建数据库和数据表

            try:
                # 数据库连接(未指定具体的数据库)
                db = pymysql.connect(host="localhost", user="root", password="lxmwork", charset='utf8',
                                     use_unicode=True)#123456
                # db = MySQLdb.connect(host, user, pw, charset='utf8')
                # 创建游标，通过连接与数据通信
                cursor = db.cursor()
                # 执行sql语句
                cursor.execute('show databases')
                databases_rows = cursor.fetchall()  # 获取所有数据库的名称
                print('所有数据库的名称=', databases_rows)
                # tmp = []
                # for row in databases_rows:
                #     tmp.append("%2s" % row)
                cursor.execute('create database if not exists ' + 'patientsinfordata')  # 不存在数据库则新建
                cursor.execute('show databases')
                databases_rows = cursor.fetchall()  # 获取所有数据库的名称
                print('新建以后所有数据库的名称=', databases_rows)
                # 数据库连接(指定具体的数据库)
                db = pymysql.connect(host="localhost", user="root", password="lxmwork", database="patientsinfordata",
                                     charset='utf8',
                                     use_unicode=True)  #
                cursor = db.cursor()
                cursor.execute('show tables')
                tables_rows = cursor.fetchall()  # 获取所有表的名称
                print('tables_rows=', tables_rows)
                # cursor.execute("""create table if not exists patientsinformation (
                #              姓名  VARCHAR(10) NOT NULL,
                #              年龄 VARCHAR(20),
                #              性别 VARCHAR(20),
                #              身高 VARCHAR(20),
                #              体重 VARCHAR(20),
                #              职业 VARCHAR(20),
                #              身份证号 VARCHAR(20),
                #              病历号 VARCHAR(20),
                #              检测时间 VARCHAR(20),
                #              常住住址 VARCHAR(20),
                #              送检医生 VARCHAR(20),
                #              检测医生 VARCHAR(20),
                #              初潮年龄 VARCHAR(20),
                #              肺活量 VARCHAR(20),
                #              )engine=InnoDB default charset=utf8;""")  # 不存在数据表则新建
                cursor.execute("""create table if not exists patientsinformation (
                             姓名  VARCHAR(20) NOT NULL,
                             年龄 INT NOT NULL,
                             性别 VARCHAR(20),
                             身高 INT,
                             体重 INT,
                             职业 VARCHAR(20),
                             身份证号 INT,
                             病历号 INT NOT NULL,
                             检测时间 VARCHAR(20),
                             常住住址 VARCHAR(20),
                             送检医生 VARCHAR(20),
                             检测医生 VARCHAR(20),
                             初潮年龄 INT,
                             肺活量 INT
                             )engine=InnoDB default charset=utf8;""")  # 不存在数据表则新建,engine=InnoDB数据表引擎,charset=utf8字符集
                cursor.execute('show tables')
                tables_rows = cursor.fetchall()  # 获取所有表的名称
                print('新建表以后的所有表=', tables_rows)
                # <editor-fold desc="弃用...">
                # 判断数据库是否存在
                # Existe_database = False
                # for i in range(len(tmp)):
                #     if tmp[i] == 'patientstest':
                #         Existe_database = True
                # if Existe_database:#存在数据库连接
                #     try:
                #         db = pymysql.connect(host = "localhost",user = "root",password = "123456", database = "patientstest",charset = 'utf8', use_unicode = True)#
                #         cursor = db.cursor()
                #     except:
                #         print('数据库连接失败')
                # else:

                # 判断数据库中是否存在对应的表
                # cursor.execute('show tables')
                # tables_rows = cursor.fetchall()  # 获取所有数据库的名称
                # for row in databases_rows:
                #     tmp.append("%2s" % row)
                # if name == tmp:
                #     cursor.execute('drop database if exists ' + name)
                # cursor.execute('create database if not exists ' + name)
                # 提交到数据库执行
                # </editor-fold>
                db.commit()
                # 2 插入数据
                name = self.DlgPatientInfoInput.strPatientName.text()
                age = self.DlgPatientInfoInput.strPatientAge.text()
                gender = self.DlgPatientInfoInput.strPatientGender.text()
                height = self.DlgPatientInfoInput.strPatientHeight.text()
                Weight = self.DlgPatientInfoInput.strPatientWeight.text()
                Career = self.DlgPatientInfoInput.strPatientCareer.text()
                ID = self.DlgPatientInfoInput.strPatientPersonalID.text()
                Case_number = self.DlgPatientInfoInput.strPatientHospitalID.text()
                CheckTime = self.DlgPatientInfoInput.strPatientCheckTime.text()
                Address = self.DlgPatientInfoInput.strPatientAddress.text()
                Doctor = self.DlgPatientInfoInput.strPatientDoctor.text()
                TestDoctor = self.DlgPatientInfoInput.strChecker.text()
                MenarcheAge = self.DlgPatientInfoInput.strPatientMenarcheAge.text()
                Vital_capacity = self.DlgPatientInfoInput.iPatientVital_capacity.text()


                # <editor-fold desc="解决数据为空的情况...">
                func = lambda x: x if x != '' else 0
                name = func(name)#为空时赋值0
                age = func(age)
                gender = func(gender)
                height = func(height)
                Weight = func(Weight)
                Career = func(Career)
                ID = func(ID)
                Case_number = func(Case_number)
                CheckTime = func(CheckTime)
                Address = func(Address)
                Doctor = func(Doctor)
                TestDoctor = func(TestDoctor)
                MenarcheAge = func(MenarcheAge)
                Vital_capacity = func(Vital_capacity)
                # </editor-fold>
                if name == 0 or Case_number == 0:#无姓名或病历号不进行插入操作
                    reply = QMessageBox.information(self, '提示', '请输入姓名或病历号', QMessageBox.Yes)
                else:
                    # sql = "INSERT INTO 'role' ('role_name', 'role_name_zh') VALUES ('{0}', '{1}')".format('role_name', '测试角色')
                    sql = "INSERT INTO patientsinformation(姓名, 年龄, 性别, 身高, 体重, 职业, 身份证号, 病历号, 检测时间, 常住住址, 送检医生, 检测医生, 初潮年龄, 肺活量)" \
                          " VALUES ('{0}', '{1}', '{2}', '{3}', '{4}', '{5}', '{6}', '{7}', '{8}', '{9}', '{10}', '{11}', '{12}', '{13}')".format(name ,age ,gender ,height ,Weight ,Career ,ID ,Case_number ,CheckTime ,Address ,Doctor ,TestDoctor ,MenarcheAge ,Vital_capacity)
                    # sql = "INSERT INTO patientsinformation(姓名, 年龄, 性别, 身高, 体重, 职业) VALUES ('{0}', '{1}', '{2}', '{3}', '{4}', '{5}')".format( name ,age ,gender ,height ,Weight ,Career)
                    # sql = "INSERT INTO patientsinformation(姓名, 年龄, 性别) VALUES ('{0}', '{1}', '{2}')".format(name, age, gender)
                    # sql = "INSERT INTO patientsinformation(姓名, 年龄, 病历号) VALUES ('{0}', '{1}', '{2}')".format(name, age,Case_number)

                    try:
                        # 执行sql语句
                        cursor.execute(sql)
                        # 提交到数据库执行
                        db.commit()
                        print('基本信息插入数据库成功')
                        # 保存完成提醒窗口
                        reply = QMessageBox.information(self, '提示', '基本信息插入数据库成功', QMessageBox.Yes)
                    except pymysql.err.InternalError:
                        # 如果发生错误则回滚
                        db.rollback()
                        reply = QMessageBox.information(self, '提示', '请注意输入格式', QMessageBox.Yes)

            except:
                print("Error")
            finally:
                # 关闭数据库连接
                db.close()
            # </editor-fold>

    def DlgFn_BNclick_InfoInput_CallCaptureExe(self):# 启动外部exe
        import subprocess

        exe_filepath = r"C:\Users\YH\Desktop\Depth.Quality.Tool.exe"
        avs = os.system(exe_filepath)
        if avs:            # 找不到指定exe提醒窗口
            reply = QMessageBox.information(self, '提示', '请确保可执行文件的路径正确!!', QMessageBox.Yes)
        # subprocess.call([exe_filepath])

    def DlgFn_InfoInput_GetFile(self, dir):#写入list中指定文件夹下的文件名
        import time
        self.DlgPatientInfoInput.List_HistoryDataList.clear()  # 清除list中的数据
        if os.path.isdir(dir):
            fileList = os.listdir(dir)#列出文件夹中的文件名
            # print('1')
            file_Creation_time = []
            file_Creation_time_int = []
            for i in range(len(fileList)):
                filepath = dir + r'/' + fileList[i]


                t = os.path.getctime(filepath)
                timeStruct = time.localtime(t)
                # file_Creation_time.append(time.strftime('%Y-%m-%d %H:%M:%S', timeStruct))
                file_Creation_time.append(time.strftime('%Y%m%d%H%M%S', timeStruct))
                file_Creation_time_int.append(int(file_Creation_time[i]))

        # i += 1
        #     file_Creation_time_int.sort()
            Z = zip(file_Creation_time_int, fileList)
            Z = sorted(Z, reverse=True)
            newfileList = []
            for j in range(len(fileList)):
                newfileList.append(Z[j][1])
            # print('file_Creation_time_int=', file_Creation_time_int)
            # print('newfileList=',newfileList)

            for e in newfileList:
                #print (e)
                self.DlgPatientInfoInput.List_HistoryDataList.addItem(e) #添加一个项
            return fileList
        else:
            reply = QMessageBox.information(self, '提示', '所选文件夹路径（D:\灵犀脊柱形态评价系统）不存在，请更正！',QMessageBox.Yes)

    def refreshfile(self):
        self.HistoryDataList = self.DlgFn_InfoInput_GetFile(r'D:\灵犀脊柱形态评价系统\数据文件')  # 列出文件夹中的文件名。输出鼠标点击所得的文件名#20190605\数据文件

    def DlgFn_BNclick_InfoInput_ChooseDataOutHospital(self):#将鼠标选定的文件名选入到右边第一个list
        #清除list中的数据
        self.DlgPatientInfoInput.List_DataOutHospital.clear()
        #得到选取的位置在listWidget中的行数（行数从0计起）
        iDataNoChooseOutHospital= self.DlgPatientInfoInput.List_HistoryDataList.currentRow()
        #print (row)
		#获得文件夹名字

        global g_strDataFolderNameOutHospital
        global g_strPatientFullPath
        global g_strAllSysDataPath
        global g_strCamPosAdjModeChosenValue
        global g_strCheckModeChosenValue
		
        g_strDataFolderNameOutHospital=self.DlgPatientInfoInput.List_HistoryDataList.item(iDataNoChooseOutHospital).text()
        g_strPatientFullPath=os.path.join(g_strAllSysDataPath, g_strDataFolderNameOutHospital)
		#出院窗口显示已经选择内容
        self.DlgPatientInfoInput.List_DataOutHospital.addItem(g_strDataFolderNameOutHospital)  # 添加一个项               
        

    def DlgFn_BNclick_Ximg_SaveImg(self):
        self.XPictureAnalyze.TE_Hint.setText("保存按钮已按下，请查看保存结果")
        img=self.XPictureAnalyze.label_imgView.bgimg.astype(np.uint8)
        #cv2.resize(img,(self.XPictureAnalyze.label_imgView.height(),self.XPictureAnalyze.label_imgView.width()),cv2.INTER_LINEAR)
        #从label_imgView控件获取数组，这些数组储存了标记的信息
        #脊柱曲线
        curvex=self.XPictureAnalyze.label_imgView.curvex
        curvey=self.XPictureAnalyze.label_imgView.curvey
        inflection=self.XPictureAnalyze.label_imgView.inflection
        #任意位置cobb角
        x0=self.XPictureAnalyze.label_imgView.x0
        y0=self.XPictureAnalyze.label_imgView.y0
        x1=self.XPictureAnalyze.label_imgView.x1
        y1=self.XPictureAnalyze.label_imgView.y1
        #股骨头
        fx0=self.XPictureAnalyze.label_imgView.fx0
        fy0=self.XPictureAnalyze.label_imgView.fy0
        fx1=self.XPictureAnalyze.label_imgView.fx1
        fy1=self.XPictureAnalyze.label_imgView.fy1
        #髂骨
        ilx0=self.XPictureAnalyze.label_imgView.ilx0
        ily0=self.XPictureAnalyze.label_imgView.ily0
        ilx1=self.XPictureAnalyze.label_imgView.ilx1
        ily1=self.XPictureAnalyze.label_imgView.ily1
        #坐骨
        isx0=self.XPictureAnalyze.label_imgView.isx0
        isy0=self.XPictureAnalyze.label_imgView.isy0
        isx1=self.XPictureAnalyze.label_imgView.isx1
        isy1=self.XPictureAnalyze.label_imgView.isy1
        #T1
        t1x0=self.XPictureAnalyze.label_imgView.t1x0
        t1y0=self.XPictureAnalyze.label_imgView.t1y0
        t1x1=self.XPictureAnalyze.label_imgView.t1x1
        t1y1=self.XPictureAnalyze.label_imgView.t1y1
        #锁骨
        cox0=self.XPictureAnalyze.label_imgView.cox0
        coy0=self.XPictureAnalyze.label_imgView.coy0
        cox1=self.XPictureAnalyze.label_imgView.cox1
        coy1=self.XPictureAnalyze.label_imgView.coy1
        #骶骨
        sax0=self.XPictureAnalyze.label_imgView.sax0
        say0=self.XPictureAnalyze.label_imgView.say0
        sax1=self.XPictureAnalyze.label_imgView.sax1
        say1=self.XPictureAnalyze.label_imgView.say1
        #骶骨相对（备用）
        rsax0=self.XPictureAnalyze.label_imgView.rsax0
        rsay0=self.XPictureAnalyze.label_imgView.rsay0
        rsax1=self.XPictureAnalyze.label_imgView.rsax1
        rsay1=self.XPictureAnalyze.label_imgView.rsay1
        #冠1cobb
        cc1x0=self.XPictureAnalyze.label_imgView.cc1x0
        cc1y0=self.XPictureAnalyze.label_imgView.cc1y0
        cc1x1=self.XPictureAnalyze.label_imgView.cc1x1
        cc1y1=self.XPictureAnalyze.label_imgView.cc1y1
        #冠2cobb
        cc2x0=self.XPictureAnalyze.label_imgView.cc2x0
        cc2y0=self.XPictureAnalyze.label_imgView.cc2y0
        cc2x1=self.XPictureAnalyze.label_imgView.cc2x1
        cc2y1=self.XPictureAnalyze.label_imgView.cc2y1
        #冠3cobb
        cc3x0=self.XPictureAnalyze.label_imgView.cc3x0
        cc3y0=self.XPictureAnalyze.label_imgView.cc3y0
        cc3x1=self.XPictureAnalyze.label_imgView.cc3x1
        cc3y1=self.XPictureAnalyze.label_imgView.cc3y1
        #矢1cobb
        sc1x0=self.XPictureAnalyze.label_imgView.sc1x0
        sc1y0=self.XPictureAnalyze.label_imgView.sc1y0
        sc1x1=self.XPictureAnalyze.label_imgView.sc1x1
        sc1y1=self.XPictureAnalyze.label_imgView.sc1y1
        #矢2cobb（备用）
        sc2x0=self.XPictureAnalyze.label_imgView.sc2x0
        sc2y0=self.XPictureAnalyze.label_imgView.sc2y0
        sc2x1=self.XPictureAnalyze.label_imgView.sc2x1
        sc2y1=self.XPictureAnalyze.label_imgView.sc2y1
        #冠状面失衡度
        cr=self.XPictureAnalyze.label_imgView.cr
        #顶椎上弯
        tvmu=self.XPictureAnalyze.label_imgView.tvmu
        #底椎下弯
        tvml=self.XPictureAnalyze.label_imgView.tvml
        #矢状面平衡
        sg=self.XPictureAnalyze.label_imgView.sg
        #腰椎前突角
        lbx0=self.XPictureAnalyze.label_imgView.lbx0
        lby0=self.XPictureAnalyze.label_imgView.lby0
        lbx1=self.XPictureAnalyze.label_imgView.lbx1
        lby1=self.XPictureAnalyze.label_imgView.lby1
        #Risser参数
        rs=self.XPictureAnalyze.label_imgView.rs
        #腰5相对倾角
        l5x0=self.XPictureAnalyze.label_imgView.l5x0
        l5x1=self.XPictureAnalyze.label_imgView.l5x1
        l5y0=self.XPictureAnalyze.label_imgView.l5y0
        l5y1=self.XPictureAnalyze.label_imgView.l5y1
        #脊柱标记点
        px=self.XPictureAnalyze.label_imgView.px
        py=self.XPictureAnalyze.label_imgView.py
        #顶旋
        tvsx0=self.XPictureAnalyze.label_imgView.tvsx0
        tvsy0=self.XPictureAnalyze.label_imgView.tvsy0
        tvsx1=self.XPictureAnalyze.label_imgView.tvsx1
        tvsy1=self.XPictureAnalyze.label_imgView.tvsy1
        ls=self.XPictureAnalyze.label_imgView.ls
        #上旋
        uvsx0=self.XPictureAnalyze.label_imgView.uvsx0
        uvsy0=self.XPictureAnalyze.label_imgView.uvsy0
        uvsx1=self.XPictureAnalyze.label_imgView.uvsx1
        uvsy1=self.XPictureAnalyze.label_imgView.uvsy1
        ls2=self.XPictureAnalyze.label_imgView.ls2
        #下旋
        lvsx0=self.XPictureAnalyze.label_imgView.lvsx0
        lvsy0=self.XPictureAnalyze.label_imgView.lvsy0
        lvsx1=self.XPictureAnalyze.label_imgView.lvsx1
        lvsy1=self.XPictureAnalyze.label_imgView.lvsy1
        ls3=self.XPictureAnalyze.label_imgView.ls3
        #mask轮廓
        vert_set = self.XPictureAnalyze.label_imgView.vert_set
        # 最小外接矩形
        box_set = Xcopy.deepcopy(self.XPictureAnalyze.label_imgView.box_set)
        for k in range(0, len(box_set)):
            box_set[k][0][0] = round(box_set[k][0][0] / fImgScale)
            box_set[k][0][1] = round(box_set[k][0][1] / fImgScale)
            box_set[k][1][0] = round(box_set[k][1][0] / fImgScale)
            box_set[k][1][1] = round(box_set[k][1][1] / fImgScale)
            box_set[k][2][0] = round(box_set[k][2][0] / fImgScale)
            box_set[k][2][1] = round(box_set[k][2][1] / fImgScale)
            box_set[k][3][0] = round(box_set[k][3][0] / fImgScale)
            box_set[k][3][1] = round(box_set[k][3][1] / fImgScale)
        # 储存矩形的宽高数组，暂时没有使用
        rect_set = self.XPictureAnalyze.label_imgView.rect_set
        # 矩形中心点的xy坐标
        list_center = self.XPictureAnalyze.label_imgView.list_center
        # 自动识别拟合出的脊柱曲线
        spacex = self.XPictureAnalyze.label_imgView.spacex
        spacey = self.XPictureAnalyze.label_imgView.spacey
        # 自动识别出的拐点
        inflection_auto = self.XPictureAnalyze.label_imgView.inflection_auto
        #定义颜色，注意这里的顺序位BGR
        red=(0, 0, 255)
        green=(0, 255, 0)
        blue=(255, 0, 0)
        yellow=(0,255,255)
        orange=(0, 165, 255)
        pink=(203, 192, 255)
        violet=(238 , 130, 238)
        brown=(42, 42, 165)
        oliver=(50, 205, 154)
        Moccasin=(181, 228,  255)
        Peru=(63, 133,205)
        Maroon=(96,48,176)
        Bisque=(196,228,255)
        Coral=(86, 114, 255)
        Chartreuse=(0, 255, 127)
        Chocolate=(30,105,210)
        color_list = [orange, pink, violet, brown, oliver, Moccasin, Peru, Maroon, Bisque, Coral]
        #设置文字字体
        font = cv2.FONT_HERSHEY_SIMPLEX
        #图片上画出标记
        #脊柱曲线
        for i in range(0,len(curvex)):
            cv2.circle(img, (round(curvex[i]/fImgScale), round(curvey[i]/fImgScale)), round(2/fImgScale), green, -1)    
        for m in range(0, len(inflection)):
            cv2.circle(img, (round(curvex[inflection[m]]/fImgScale), round(curvey[inflection[m]]/fImgScale)), round(2/fImgScale), red, -1)    
        #任意位置cobb
        for i in range(0,len(x1)):
            cv2.line(img, (round(x0[i]/fImgScale), round(y0[i]/fImgScale)), (round(x1[i]/fImgScale), round(y1[i]/fImgScale)), red, round(2/fImgScale))
        #脊柱标记点，楔形变
        for j in range(len(px)):
            if len(px[j])>=2:
             cv2.line(img, (round(px[j][0]/fImgScale), round(py[j][0]/fImgScale)), (round(px[j][1]/fImgScale), round(py[j][1]/fImgScale)), green, round(2/fImgScale))
            if len(px[j])>=4:
             cv2.line(img, (round(px[j][2]/fImgScale), round(py[j][2]/fImgScale)), (round(px[j][3]/fImgScale), round(py[j][3]/fImgScale)), green, round(2/fImgScale))
             cv2.putText(img, str(j+1), (round(px[j][3]/fImgScale), round(py[j][3]/fImgScale)), font, round(0.6/fImgScale), green, round(1/fImgScale), False)    
        #股骨头
        for i in range(0,len(fx1)):
            cv2.line(img, (round(fx0[i]/fImgScale), round(fy0[i]/fImgScale)), (round(fx1[i]/fImgScale), round(fy1[i]/fImgScale)), blue, round(2/fImgScale))
        #髂骨
        for i in range(0,len(ilx1)):
            cv2.line(img, (round(ilx0[i]/fImgScale), round(ily0[i]/fImgScale)), (round(ilx1[i]/fImgScale), round(ily1[i]/fImgScale)), orange, round(2/fImgScale))
        #坐骨
        for i in range(0,len(isx1)):
            cv2.line(img, (round(isx0[i]/fImgScale), round(isy0[i]/fImgScale)), (round(isx1[i]/fImgScale), round(isy1[i]/fImgScale)), pink, round(2/fImgScale))
        #T1
        for i in range(0,len(t1x1)):
            cv2.line(img, (round(t1x0[i]/fImgScale), round(t1y0[i]/fImgScale)), (round(t1x1[i]/fImgScale), round(t1y1[i]/fImgScale)), Moccasin, round(2/fImgScale))
        #锁骨
        for i in range(0,len(cox1)):
            cv2.line(img, (round(cox0[i]/fImgScale), round(coy0[i]/fImgScale)), (round(cox1[i]/fImgScale), round(coy1[i]/fImgScale)), Moccasin, round(2/fImgScale))
        #骶骨相对（备用）
        for i in range(0,len(rsax1)):
            cv2.line(img, (round(rsax0[i]/fImgScale), round(rsay0[i]/fImgScale)), (round(rsax1[i]/fImgScale), round(rsay1[i]/fImgScale)), red, round(2/fImgScale))
        #骶骨
        for i in range(0,len(sax1)):
            cv2.line(img, (round(sax0[i]/fImgScale), round(say0[i]/fImgScale)), (round(sax1[i]/fImgScale), round(say1[i]/fImgScale)), Moccasin, round(2/fImgScale))
        #冠cobb1-3
        for i in range(0,len(cc1x1)):
            cv2.line(img, (round(cc1x0[i]/fImgScale), round(cc1y0[i]/fImgScale)), (round(cc1x1[i]/fImgScale), round(cc1y1[i]/fImgScale)), red, round(2/fImgScale))
        for i in range(0,len(cc2x1)):
            cv2.line(img, (round(cc2x0[i]/fImgScale), round(cc2y0[i]/fImgScale)), (round(cc2x1[i]/fImgScale), round(cc2y1[i]/fImgScale)), red, round(2/fImgScale))
        for i in range(0,len(cc3x1)):
            cv2.line(img, (round(cc3x0[i]/fImgScale), round(cc3y0[i]/fImgScale)), (round(cc3x1[i]/fImgScale), round(cc3y1[i]/fImgScale)), red, round(2/fImgScale))
        #矢cobb1-2
        for i in range(0,len(sc1x1)):
            cv2.line(img, (round(sc1x0[i]/fImgScale), round(sc1y0[i]/fImgScale)), (round(sc1x1[i]/fImgScale), round(sc1y1[i]/fImgScale)), red, round(2/fImgScale))
        for i in range(0,len(sc2x1)):
            cv2.line(img, (round(sc2x0[i]/fImgScale), round(sc2y0[i]/fImgScale)), (round(sc2x1[i]/fImgScale), round(sc2y1[i]/fImgScale)), red, round(2/fImgScale))
        #冠状面失衡度
        for i in range(0,math.floor(len(cr)/2)):
            cv2.circle(img, (round(cr[2*i]/fImgScale), round(cr[2*i+1]/fImgScale)), round(2/fImgScale), Peru, -1)
        #上弯偏移距离
        for i in range(0,math.floor(len(tvmu)/2)):
            cv2.circle(img, (round(tvmu[2*i]/fImgScale), round(tvmu[2*i+1]/fImgScale)), round(2/fImgScale), Maroon, -1)
        #下弯偏移距离
        for i in range(0,math.floor(len(tvml)/2)):
            cv2.circle(img, (round(tvml[2*i]/fImgScale), round(tvml[2*i+1]/fImgScale)), round(2/fImgScale), Maroon, -1)    
        #矢状面平衡
        for i in range(0,math.floor(len(sg)/2)):
            cv2.circle(img, (round(sg[2*i]/fImgScale), round(sg[2*i+1]/fImgScale)), round(2/fImgScale), Bisque, -1)
        #腰椎前突角
        for i in range(0,len(lbx1)):
            cv2.line(img, (round(lbx0[i]/fImgScale), round(lby0[i]/fImgScale)), (round(lbx1[i]/fImgScale), round(lby1[i]/fImgScale)), Coral, round(2/fImgScale))
        #顶椎旋转
        for i in range(0,len(tvsx1)):
            cv2.line(img, (round(tvsx0[i]/fImgScale), round(tvsy0[i]/fImgScale)), (round(tvsx1[i]/fImgScale), round(tvsy1[i]/fImgScale)), Chartreuse, round(2/fImgScale))
        for i in range(0,math.floor(len(ls)/2)):
            cv2.circle(img, (round(ls[2*i]/fImgScale), round(ls[2*i+1]/fImgScale)), round(2/fImgScale), Chartreuse, -1)
        #上端椎旋转
        for i in range(0,len(uvsx1)):
            cv2.line(img, (round(uvsx0[i]/fImgScale), round(uvsy0[i]/fImgScale)), (round(uvsx1[i]/fImgScale), round(uvsy1[i]/fImgScale)), Chartreuse, round(2/fImgScale))
        for i in range(0,math.floor(len(ls2)/2)):
            cv2.circle(img, (round(ls2[2*i]/fImgScale), round(ls2[2*i+1]/fImgScale)), round(2/fImgScale), Chartreuse, -1)
        #下端椎旋转
        for i in range(0,len(lvsx1)):
            cv2.line(img, (round(lvsx0[i]/fImgScale), round(lvsy0[i]/fImgScale)), (round(lvsx1[i]/fImgScale), round(lvsy1[i]/fImgScale)), Chartreuse, round(2/fImgScale))
        for i in range(0,math.floor(len(ls3)/2)):
            cv2.circle(img, (round(ls3[2*i]/fImgScale), round(ls3[2*i+1]/fImgScale)), round(2/fImgScale), Chartreuse, -1)
        #risser参数
        for i in range(0,math.floor(len(rs)/2)):
            cv2.circle(img, (round(rs[2*i]/fImgScale), round(rs[2*i+1]/fImgScale)), round(2/fImgScale), green, -1)
        #腰5相对倾角
        for i in range(0,len(l5x1)):
            cv2.line(img, (round(l5x0[i]/fImgScale), round(l5y0[i]/fImgScale)), (round(l5x1[i]/fImgScale), round(l5y1[i]/fImgScale)), red, round(2/fImgScale))
        #画mask（自动识别）
        #if len(vert_set)>=1:
            #vert = np.empty((len(vert_set[0]), 2), dtype=np.int32)
            #for j in range(0, len(vert_set[0])):
                #vert[j, 0] = round(vert_set[0][j][0] / fImgScale)
                #vert[j, 1] = round(vert_set[0][j][1] / fImgScale)
            #cv2.fillPoly(img, [vert], red)
            #length = 2 * (
                    #max(box_set[0][0][0], box_set[0][1][0], box_set[0][2][0],
            # box_set[0][3][0]) - min(box_set[0][0][0], box_set[0][1][0],
            # box_set[0][2][0], box_set[0][3][0]))
            # cv2.putText(img, str(1), (
                #int(list_center[0][0] / fImgScale + length), int(list_center[0][1] / fImgScale )), font,
                        #round(0.6 / fImgScale), red, round(1 / fImgScale), False)
        #if len(vert_set)>=13:
            #for i in range(1, 13):
                #vert = np.empty((len(vert_set[i]), 2), dtype=np.int32)
            # for j in range(0, len(vert_set[i])):
            # vert[j, 0] = round(vert_set[i][j][0] / fImgScale)
                    #vert[j, 1] = round(vert_set[i][j][1] / fImgScale)
                #cv2.fillPoly(img, [vert], orange)
                #length=2*(
            # max(box_set[i][0][0], box_set[i][1][0], box_set[i][2][0],
            #  box_set[i][3][0]) - min(box_set[i][0][0], box_set[i][1][0],
        # box_set[i][2][0], box_set[i][3][0]))
            # cv2.putText(img, str(i + 1), (
            # int(list_center[i][0] / fImgScale + length), int(list_center[i][1] / fImgScale )), font,
            # round(0.6 / fImgScale), orange, round(1 / fImgScale), False)
        #if len(vert_set)>=14:
            #for i in range(13, len(vert_set)):
                #vert = np.empty((len(vert_set[i]), 2), dtype=np.int32)
                #for j in range(0, len(vert_set[i])):
            # vert[j, 0] = round(vert_set[i][j][0] / fImgScale)
                    #vert[j, 1] = round(vert_set[i][j][1] / fImgScale)
                #cv2.fillPoly(img, [vert], yellow)
            # length = 2 * (
                        #max(box_set[i][0][0], box_set[i][1][0], box_set[i][2][0],
                            #box_set[i][3][0]) - min(box_set[i][0][0], box_set[i][1][0],
                                                    #box_set[i][2][0], box_set[i][3][0]))
            # cv2.putText(img, str(i + 1), (
                    #int(list_center[i][0] / fImgScale + length), int(list_center[i][1] / fImgScale )), font,
                            #round(0.6 / fImgScale), yellow, round(1 / fImgScale), False)
        #画最小外接矩形（自动识别）
        for k in range(0, len(box_set)):
            cv2.drawContours(img,[box_set[k]],0,green,round(1 / fImgScale))
            length = 2 * (
                    max(box_set[k][0][0], box_set[k][1][0], box_set[k][2][0],
                        box_set[k][3][0]) - min(box_set[k][0][0], box_set[k][1][0],
                                                box_set[k][2][0], box_set[k][3][0]))
            cv2.putText(img, str(k + 1), (
                int(list_center[k][0] / fImgScale - length), int(list_center[k][1] / fImgScale)), font,
                        round(0.6 / fImgScale), green, round(1 / fImgScale), False)
        #外接矩形中心（自动识别）
        for k in range(0, len(list_center)):
            cv2.circle(img, (int(list_center[k][0] / fImgScale), int(list_center[k][1] / fImgScale)), round(2 / fImgScale),
                       green, -1)
        #拟合曲线（自动识别）
        for k in range(0, len(spacex)):
            cv2.circle(img, (int(spacex[k] / fImgScale), int(spacey[k] / fImgScale)), round(2 / fImgScale),
                       green, -1)
        #曲线拐点（自动识别）
        for m in range(0, len(inflection_auto) - 1):
            colorindex = Xcopy.deepcopy(m)
            if colorindex > len(color_list) - 1:
                colorindex = 0
            for k in range(inflection_auto[m], inflection_auto[m + 1]):
                cv2.circle(img,
                           (int(spacex[k] / fImgScale), int(spacey[k] / fImgScale)),
                           round(2 / fImgScale),
                           color_list[colorindex], -1)
        for m in range(0, len(inflection_auto)):
            cv2.circle(img, (int(spacex[inflection_auto[m]] / fImgScale), int(spacey[inflection_auto[m]] / fImgScale)), round(2 / fImgScale),
                       blue, -1)
        #用浏览窗口，选择图片储存的位置，返回路径和类型
        imgName, imgType=QFileDialog.getSaveFileName(self, "save picture", "", "*.jpg;;*.png;;All Files(*)")
        #获取文件的后缀
        BackStr=imgName[-4:]
        #判断文件路径存在且后缀符合要求时，执行存储
        if imgName :
            if imgType=='*.jpg' and BackStr=='.jpg':
                cv2.imencode('.jpg', img)[1].tofile(imgName)
            elif imgType=='*.png' and BackStr=='.png':
                cv2.imencode('.png', img)[1].tofile(imgName)
            else:
                print('文件后缀错误，取消存储')
        else:
            print('取消存储')
        

    def DlgFn_BNclick_Ximg_ReadXImg(self):
        #from PyQt5.QtGui import QImage, QPixmap
 #       self.XPictureAnalyze.label_imgView = myLabel(self)
        #图片和label控件左上角的距离
        x0=0
        y0=25
        global fImgScale
        global XimgName
        fImgScale=self.DSB_fImgScale.value()#图像放大系数
        #浏览窗口打开文件，得到文件的路径和类型
        XimgName, imgType=QFileDialog.getOpenFileName(self, "open picture", "", "*.jpg;;*.png;;All Files(*)")
        #文件路径存在且文件类型符合要求时，进行图片的处理
        if XimgName and (imgType=='*.jpg' or imgType=='*.png'):
          #调用定义的标记清除方法，清除之前的标记
          self.XPictureAnalyze.label_imgView.LineClear()
          self.XPictureAnalyze.LE_RandomCobbResult.clear()
          self.XPictureAnalyze.label_imgView.PointClear()
          self.XPictureAnalyze.label_imgView.MorphClear()
          self.XPictureAnalyze.label_imgView.CoronalCobb1Clear()
          self.XPictureAnalyze.LE_CoronalCobb1Result.clear() 
          self.XPictureAnalyze.label_imgView.CoronalCobb2Clear()
          self.XPictureAnalyze.LE_CoronalCobb2Result.clear() 
          self.XPictureAnalyze.label_imgView.CoronalCobb3Clear()
          self.XPictureAnalyze.label_imgView.IliumClear()
          self.XPictureAnalyze.LE_IliumResult.clear()
          self.XPictureAnalyze.label_imgView.IschiumClear()
          self.XPictureAnalyze.LE_IschiumResult.clear()
          self.XPictureAnalyze.label_imgView.SacrumClear()
          self.XPictureAnalyze.LE_SacrumResult.clear()
          self.XPictureAnalyze.label_imgView.RelSacrumClear()
          self.XPictureAnalyze.LE_RelSacrumResult.clear()
          self.XPictureAnalyze.label_imgView.RsClear()
          self.XPictureAnalyze.LE_RsResult.clear()
          self.XPictureAnalyze.label_imgView.CalRs=1
          self.XPictureAnalyze.label_imgView.T1Clear()
          self.XPictureAnalyze.LE_T1Result.clear()
          self.XPictureAnalyze.label_imgView.CollarboneClear()
          self.XPictureAnalyze.LE_CollarboneResult.clear()
          self.XPictureAnalyze.label_imgView.LumbarClear()
          self.XPictureAnalyze.LE_LumbarResult.clear()
          self.XPictureAnalyze.label_imgView.CoronalClear()
          self.XPictureAnalyze.LE_CoronalResult.clear()
          self.XPictureAnalyze.label_imgView.SagittalClear()
          self.XPictureAnalyze.LE_SagittalResult.clear()
          self.XPictureAnalyze.label_imgView.TVMUperClear()
          self.XPictureAnalyze.LE_TVMUperResult.clear()
          self.XPictureAnalyze.label_imgView.TVMLowerClear()
          self.XPictureAnalyze.LE_TVMLowerResult.clear()
          self.XPictureAnalyze.label_imgView.TVSClear()
          self.XPictureAnalyze.LE_TVSResult.clear()
          self.XPictureAnalyze.label_imgView.UVSClear()
          self.XPictureAnalyze.LE_UVSResult.clear()
          self.XPictureAnalyze.label_imgView.LVSClear()
          self.XPictureAnalyze.LE_LVSResult.clear()
          self.XPictureAnalyze.label_imgView.MaskClear()
          self.XPictureAnalyze.label_imgView.MidClear()
          self.XPictureAnalyze.label_imgView.EndClear()
          
         #读取一张图片，这个方法可以读中文路径
          img0=cv2.imdecode(np.fromfile(XimgName,dtype=np.uint8),1)
          #异常处理
          if img0 is None:
              print('图片读取错误，请更新opencv或检查路径是否有特殊字符')
          else:
             #获取并打印图片的原始尺寸
             size = img0.shape#行数*列数，所以shape[0]是图像高度
             #width = round(size[0]*0.5)#
             #height = round(size[1]*0.5)#
             height = round(size[0]*fImgScale)
             width = round(size[1]*fImgScale)
             print("imgsize")
             print(size)
             #img= cv2.resize(img0,(height,width),cv2.INTER_LINEAR)
            #图片缩放
             img=cv2.resize(img0, (0, 0), fx=fImgScale, fy=fImgScale, interpolation=cv2.INTER_NEAREST)
             #原图存储，方便之后调用
             self.XPictureAnalyze.label_imgView.bgimg=img0
             #储存缩放后的图片，方便以后调用
             self.XPictureAnalyze.label_imgView.rsimg=img
             #设置图片位置，颜色频道，鼠标光标样式
             height, width, bytesPerComponent = img.shape
             self.XPictureAnalyze.label_imgView.setGeometry(QRect(x0, y0,width ,height ))
             bytesPerLine = 3 * width
             cv2.cvtColor(img, cv2.COLOR_BGR2RGB, img)
             QImg = QImage(img.data, width, height, bytesPerLine, QImage.Format_RGB888)
             pixmap = QPixmap.fromImage(QImg)
             self.XPictureAnalyze.label_imgView.setPixmap(pixmap)
             self.XPictureAnalyze.label_imgView.setCursor(Qt.CrossCursor)
             self.show()
             self.XPictureAnalyze.label_imgView.ImgCount=1
             self.XPictureAnalyze.label_imgView.filename=XimgName[-9:]
        else: 
            print('无法读取该文件格式，取消读取')

    def DlgFn_BNclick_Ximg_FlipXImg(self):
        global fImgScale
        x0=0
        y0=25
        if self.XPictureAnalyze.label_imgView.ImgCount==1:
            #利用python的数组切片，使图片的x轴逆序
            img0=self.XPictureAnalyze.label_imgView.bgimg[:, ::-1, :]
            #调整图片大小
            img=cv2.resize(img0, (0, 0), fx=fImgScale, fy=fImgScale, interpolation=cv2.INTER_NEAREST)
            #存储反转后的图片，方便以后调用
            self.XPictureAnalyze.label_imgView.bgimg=img0
            # 存储反转缩放后的图片，方便以后调用
            self.XPictureAnalyze.label_imgView.rsimg = img
            #设置图片的位置，颜色，光标
            height, width, bytesPerComponent = img.shape
            self.XPictureAnalyze.label_imgView.setGeometry(QRect(x0, y0,width ,height ))
            bytesPerLine = 3 * width
            cv2.cvtColor(img, cv2.COLOR_BGR2RGB, img)
            QImg = QImage(img.data, width, height, bytesPerLine, QImage.Format_RGB888)
            pixmap = QPixmap.fromImage(QImg)
            self.XPictureAnalyze.label_imgView.setPixmap(pixmap)
            self.XPictureAnalyze.label_imgView.setCursor(Qt.CrossCursor)
            self.show()


    def MainGenFn_ReadBKCamAngleXYZ(self):
        try:
            strCamAngleXYZFilePath=g_strPatientFullPath+'/AngleResult.txt'
            fCamAngleXYZ = np.loadtxt(strCamAngleXYZFilePath)
            strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
            self.DlgPatientInfoInput.strMessageOut.setPlainText("患者目录相机旋转角已经存在并读取成功！\n"+strtem) 
        except IOError:
            strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
            self.DlgPatientInfoInput.strMessageOut.setPlainText("患者背景点坐标文件不存在！将从系统目录拷贝背景标记点\n"+strtem) 
			
        return fCamAngleXYZ
    def MainGenFn_ReadBKCor(self,strFilePath):
        fBKPtCor=0
        try:
            strBackroundPointCorFilePath=strFilePath+'/背景点.txt'
            fBKPtCor = np.loadtxt(strBackroundPointCorFilePath)
            strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
            self.DlgPatientInfoInput.strMessageOut.setPlainText("患者目录背景点坐标已经存在并读取成功！\n"+strtem) 
        except IOError:
            strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
            self.DlgPatientInfoInput.strMessageOut.setPlainText("患者背景点坐标文件不存在！将从系统目录拷贝背景标记点\n"+strtem) 
			#患者背景点不存在，则需要从系统目录下拷贝一份
            try:
                strBackroundPointCorFilePath=g_strAllSysDataPath+'/背景点.txt'
                fBKPtCor = np.loadtxt(strBackroundPointCorFilePath)
				# 拷贝系统目录下的背景点坐标到患者根目录下，如果有则不拷贝，没有则拷贝
                np.savetxt(g_strPatientFullPath + '/背景点.txt', fBKPtCor)
                strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
                self.DlgPatientInfoInput.strMessageOut.setPlainText("患者目录背景点坐标已经从系统目录拷贝并读取成功！\n"+strtem)   
            except BaseException:
                strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
                self.DlgPatientInfoInput.strMessageOut.setPlainText("系统目录下背景点坐标不存在！\n"+strtem)
        fBKPtCorXYZ=fBKPtCor
        return fBKPtCorXYZ
		# 读取背景点结束
    def ReadMarkerCor(self,strFilePath):
		#读取标记点文件，返回格式为3*n
		#这里的skiprows是指跳过前1行, 如果设置skiprows=2, 就会跳过前两行,  这里的输出为
        strStandMidFilePath=strFilePath+ '\静态扫描\站姿中立位'
		##########通过标记点来平移点云使得髂后上棘的中点为原点
        fMarkerPoint_x = []#赋初值，如果无标记点则为空，有标记点则再赋值给此变量
        fMarkerPoint_y = []
        fMarkerPoint_z = []
        try:
            strMarkerPoint = np.loadtxt(strStandMidFilePath+'/1Result.txt', dtype=str,
                                      comments=['*', 'Su', 'AL'])  # , skiprows=i跳过前i行
        except IOError:#无标记点文件
            #显示系统信息
            strtem=self.DlgPatientInfoInput.strMessageOut.toPlainText()
            self.DlgPatientInfoInput.strMessageOut.setPlainText(strStandMidFilePath+'1Result.txt文件不存在'+str)
        else:#有标记点文件
            strMarkerPoint= np.loadtxt(strStandMidFilePath + '/1Result.txt', 
			                                  dtype=str, comments=['*','Su','A', 'Buttock', 'Shoulder'])#, skiprows=i跳过前i行
            if len(strMarkerPoint) <= 6:
                number_strMarkerPoint = len(strMarkerPoint)#六标记点
            else:
                #number_strMarkerPoint = len(strMarkerPoint)-4#多标记点的最后四个为错误点
                number_strMarkerPoint = len(strMarkerPoint)#多标记点的最后四个为错误点20190618
            for i in range(number_strMarkerPoint):#range(strMarkerPoint.shape[0])#去除最后四个点-4
                strMarkerPoint3d = strMarkerPoint[i][1].split('(')[2][:-1]#提取出该行的坐标参数
                # fMarkerPoint3d = fMarkerPoint[i].split('(')[3][:-1]  # 提取出该行的坐标参数20190618
                fMarkerPoint_x.append(float(strMarkerPoint3d.split(',')[0]))#提取出该行的坐标参数x变为数值
                fMarkerPoint_y.append(float(strMarkerPoint3d.split(',')[1]))#提取出该行的坐标参数y变为数值
                fMarkerPoint_z.append(float(strMarkerPoint3d.split(',')[2]))#提取出该行的坐标参数z变为数值
                # print(fMarkerPoint3d)
            # print(fMarkerPoint[0])
            #fMarkerPoint_x,fMarkerPoint_y = fMarkerPoint_y,fMarkerPoint_x
			#ListfMarkerPoint3d--> 3*n
        ListfMarkerPoint3d=[fMarkerPoint_x, fMarkerPoint_y, fMarkerPoint_z]
        MatrixfMarkerPt3d=np.matrix(ListfMarkerPoint3d)
        return MatrixfMarkerPt3d
    def DlgFn_BNclick_InfoInput_ChooseDataInHospital(self):#将鼠标选定的文件名选入到右边第二个list
         #清除list中的数据
        self.DlgPatientInfoInput.List_DataInHospital.clear()
        #得到选取的位置在listWidget中的行数（行数从0计起）
        iDataNoChooseInHospital= self.DlgPatientInfoInput.List_HistoryDataList.currentRow()
        #print (row)
		#获得文件夹名字
        g_strDataFolderNameInHospital=self.DlgPatientInfoInput.List_HistoryDataList.item(iDataNoChooseInHospital).text()
        #出院窗口显示已经选择内容
        self.DlgPatientInfoInput.List_DataInHospital.addItem(g_strDataFolderNameInHospital)  # 添加一个项               

    def DlgFn_BNclick_InfoInput_IndexHisRecord(self):  # 检索姓名或者编码，并显示再左边的list中
        dir = r'D:\灵犀脊柱形态评价系统\数据文件'#20190605\数据文件
        import time

        # 清空上次显示的内容
        self.DlgPatientInfoInput.List_HistoryDataList.clear()
        fileList = os.listdir(dir)  # 列出文件夹中的文件名

        file_modify_time = []
        file_modify_time_int = []
        i=0
        for full_path, dirl, files in os.walk(dir):

            mtime = os.stat(full_path).st_mtime
            file_modify_time.append(time.strftime('%Y%m%d%H%M%S', time.localtime(mtime)))
            # file_modify_time.append(time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(mtime)))
            # print("{0} 修改时间是: {1}".format(full_path,file_modify_time))
            file_modify_time_int.append(int(file_modify_time[i]))
            i += 1
        fileList.sort(key=lambda file_modify_time_int: file_modify_time_int[-14:])  # 按修改时间排序
        fileList.reverse()  # 新修改的放在前面
        print('fileList',fileList)

        if os.path.isdir(dir):
            if (self.DlgPatientInfoInput.strPatientName.text() !=''):# or (self.DlgPatientInfoInput.strPatientName.text() !='')):
                for List in fileList:
                    if self.DlgPatientInfoInput.strPatientName.text() in List: #or self.DlgPatientInfoInput.strPatientPersonalID.text() in List:
                        print ('由姓名检索到',List)
                        self.DlgPatientInfoInput.List_HistoryDataList.addItem(List) #添加一个项
            elif (self.DlgPatientInfoInput.strPatientPersonalID.text() !=''):# or
                for List in fileList:
                    if self.DlgPatientInfoInput.strPatientPersonalID.text() in List: #or self.DlgPatientInfoInput.strPatientPersonalID.text() in List:
                        print ('由身份证号检索到',List)
                        self.DlgPatientInfoInput.List_HistoryDataList.addItem(List) #添加一个项
            else:
                # reply = QMessageBox.information(self, '提示', '请输入姓名或者身份证号进行检索！', QMessageBox.Yes)
                self.refreshfile()#没有输入显示全部
        else:
            reply = QMessageBox.information(self, '提示', '所选数据文件夹路径不存在，请更正！',QMessageBox.Yes)

        reply = QMessageBox.information(self, '提示', '查询完毕！', QMessageBox.Yes)

    def DlgFn_BNclick_InfoInput_StartAnalyze(self):
        global g_strDataFolderNameOutHospital, global_g_strDataFolderNameOutHospital
        global Main_bend_center, Main_bend_Boundary_points,Main_bend_rad,finalCBendfBKPtCor
        # <editor-fold desc="指标全局变量">
        global Coronal_trunk_imbalance#1冠状面躯干失平衡度
        global Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y#   2左右最高点水平倾角（左高为正）
        global Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z#   3左右最高点深度方向倾角（左高(凸)为正）
        global Height_dif_of_posterior_superior_iliac_spine#   4左右髂后上棘高度差（左高为正）
        global Horizontal_obliquity_of_pelvis#   5骨盆水平倾斜角（左高为正）
        global Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve    #   6截面曲线旋转角均值(代数和)(按右高为正算)
        global abs_Mean_Side_spin_angle_of_Section_Curve#   7截面曲线旋转角均值(绝对值)
        global Height_difference_between_shoulders#   8两肩高度差（左高为正）
        global Thoracic_rotation_angle#   9胸段旋转角（右高为正）（选前9段）
        global Thoracolumbar_rotation_angle#   10胸腰段旋转角（右高为正）(选接下来的10段)
        global Lumbar_rotation_angle       #   11腰段旋转角（右高为正）(选接下来的5段)
        global Main_bend_rad_to_save       #   12冠状面C1弯曲率半径
        global cobber#   13冠状面C1弯cobb角（右弯为正）注意正负
        global Side_bend_type       #   14侧弯类型（s,c,多c）
        global mean_of_abs_differences,Sum_of_abs_differences#   15冠状面脊柱曲线围成的面积（绝对值）
        global mean_of_differences, Sum_of_differences      #   16冠状面脊柱曲线围成的面积（代数值）
        global linear_inclination_of_coronal_spine      #   17冠状面脊柱平均直线倾角(右侧为正)
        global Hunchback_inclination#   18驼背倾角
        global rotation_angle_Variance#   19截面曲线旋转角离散度
        global Lateral_buckling_limit_right_angle,Lateral_buckling_limit_left_angle#   20侧屈极限角(°)
        global mean_Left_and_right_difference#   21侧屈15左右围成面积
        global dif_between_left_and_right_lateral_flexion_horizontal_direction#   22左右侧屈15水平方向差值和（均值）
        # </editor-fold>

        print('启动数据分析计算指标')
        # final_fMarkerPoint_end = YH_DataPreprocess.final_fMarkerPoint_end #矫正以后的标记点
        # print('final_fMarkerPoint_end',final_fMarkerPoint_end)
        final_fMarkerPoint_end = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/correct_fMarkerPoint.txt')  #矫正以后的标记点
            #  , skiprows=i跳过前i行, dtype=str, comments=['*', 'Su', 'AL']
        #这里的final_fMarkerPoint_end x值和y值是原始数据的顺序，即x，y，z的顺序
        final_points = np.loadtxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/final_points.txt')  # 矫正去除背景后的点云数据
        # <editor-fold desc="1-5...">
        #   1冠状面躯干失平衡度（过颈七的垂线和过髂棘连线的中点的垂线的水平距离）
        final_fMarkerPoint_end = np.array(final_fMarkerPoint_end)
        spina_iliaca_middle_y = (final_fMarkerPoint_end[-4:-3,0] + final_fMarkerPoint_end[-3:-2,0])/2#髂棘连线中点的x值（图中的y值（数据中的第一列））20190622
        Neck_seven_y = final_fMarkerPoint_end[-5:-4,0]
        Coronal_trunk_imbalance = abs(spina_iliaca_middle_y - Neck_seven_y) #1冠状面躯干失平衡度20190622
        # Coronal_trunk_imbalance = Neck_seven_y#20190622
        print('1冠状面躯干失平衡度=', Coronal_trunk_imbalance)
#   2左右最高点水平倾角（左高为正）
        fMarkerPoint_spina_mean_X = np.mean(final_fMarkerPoint_end[:6,0])#脊柱前六个标记点y的均值
        Ind_Left_cloud_points = np.where(final_points[:,0] < fMarkerPoint_spina_mean_X)#final_points中第一列是图中的x坐标
        Ind_Right_cloud_points = np.where(final_points[:,0] >= fMarkerPoint_spina_mean_X)#final_points中第一列是图中的x坐标

        # 把(X, Y)以原点为旋转中心按逆时针方向旋转90度后对应点坐标是(-Y, X),所以x的正方向为左
        Left_cloud_points = final_points[Ind_Left_cloud_points]  # 左侧点云图
        Right_cloud_points = final_points[Ind_Right_cloud_points]#右侧点云图
        Min_Y_cloud_points = np.min(final_points[:,1])
        Max_Y_cloud_points = np.max(final_points[:,1])
        Two_third_Y = Min_Y_cloud_points + 7/12*(Max_Y_cloud_points - Min_Y_cloud_points)#去掉下面的7/12（留1/3太少，1/2太多）
        Ind_Left_Two_third_Y_cloud_points = np.where(Left_cloud_points[:, 1] >= Two_third_Y)
        Ind_Right_Two_third_Y_cloud_points = np.where(Right_cloud_points[:, 1] >= Two_third_Y)
        Left_Two_third_Y_cloud_points = Left_cloud_points[Ind_Left_Two_third_Y_cloud_points] #点云左侧上三分之的点
        Right_Two_third_Y_cloud_points = Right_cloud_points[Ind_Right_Two_third_Y_cloud_points]#点右左侧上三分之的点

        Ind_Max_Z_Left_Two_third_cloud_points = np.where(Left_Two_third_Y_cloud_points[:, 2] == np.max(Left_Two_third_Y_cloud_points[:, 2]))
        Ind_Max_Z_Right_Two_third_cloud_points = np.where(Right_Two_third_Y_cloud_points[:, 2] == np.max(Right_Two_third_Y_cloud_points[:, 2]))
        Max_Z_Left_Two_third_cloud_points = Left_Two_third_Y_cloud_points[Ind_Max_Z_Left_Two_third_cloud_points]
        Max_Z_Right_Two_third_cloud_points = Right_Two_third_Y_cloud_points[Ind_Max_Z_Right_Two_third_cloud_points]
        # Max_Z_Left_One_third_X_cloud_points = list(Max_Z_Left_One_third_X_cloud_points[0])
        # Max_Z_Right_One_third_X_cloud_points =list(Max_Z_Right_One_third_X_cloud_points[0])
        Max_Z_Left_and_Right = np.vstack((Max_Z_Left_Two_third_cloud_points[0],Max_Z_Right_Two_third_cloud_points[0]))
        #vstack在竖直方向上堆叠;np.hstack():在水平方向上平铺

        Left_and_Right_Highest_Points_Slope_Y = -(Max_Z_Left_Two_third_cloud_points[0][1] - Max_Z_Right_Two_third_cloud_points[0][1])/(Max_Z_Left_Two_third_cloud_points[0][0] - Max_Z_Right_Two_third_cloud_points[0][0])
        Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y = math.atan(Left_and_Right_Highest_Points_Slope_Y) * 180 / math.pi  #
            #左右最高点水平倾角（左高为正）
        np.savetxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Left_Two_third_Y_cloud_points.txt', Left_Two_third_Y_cloud_points, fmt='%.2f')
        np.savetxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Right_Two_third_Y_cloud_points.txt', Right_Two_third_Y_cloud_points , fmt='%.2f')
        np.savetxt(global_g_strDataFolderNameOutHospital + '/相关数据' + '/Max_Z_Left_and_Right.txt', Max_Z_Left_and_Right, fmt='%.2f')
        print('左最高点=', Max_Z_Left_and_Right[0])
        print('右最高点=', Max_Z_Left_and_Right[0])
        print('2左右最高点水平倾角（左高为正）=',Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y)
#   3左右最高点深度方向倾角（左高(凸)为正）
        Left_and_Right_Highest_Points_Slope_Z = -(Max_Z_Left_Two_third_cloud_points[0][2] - Max_Z_Right_Two_third_cloud_points[0][2])/(Max_Z_Left_Two_third_cloud_points[0][0] - Max_Z_Right_Two_third_cloud_points[0][0])
        Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z = math.atan(Left_and_Right_Highest_Points_Slope_Z) * 180 / math.pi  #

        print('左右最高点深度方向倾角（左高(凸)为正）=',Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z)

#   4左右髂后上棘高度差（左高为正）
        Height_dif_of_posterior_superior_iliac_spine = final_fMarkerPoint_end[-4:-3, 1] - final_fMarkerPoint_end[-3:-2, 1]
        #左髂后上棘为倒数第四个点，右髂后上棘为倒数第三个点
        print('4左右髂后上棘高度差（左高为正）=', Height_dif_of_posterior_superior_iliac_spine)
#   5骨盆水平倾斜角（左高为正）
        Horizontal_obliquity_of_pelvis_Slope = -(final_fMarkerPoint_end[-4:-3, 1] - final_fMarkerPoint_end[-3:-2, 1])/(final_fMarkerPoint_end[-4:-3, 0] - final_fMarkerPoint_end[-3:-2, 0])
        Horizontal_obliquity_of_pelvis = math.atan(Horizontal_obliquity_of_pelvis_Slope) * 180 / math.pi  # 5骨盆水平倾斜角
        print('5骨盆水平倾斜角（左高为正）=', Horizontal_obliquity_of_pelvis)
        # </editor-fold>
        # <editor-fold desc="6-10...">
#   6截面曲线旋转角均值(代数和)(按右高为正算)
        Side_spin_angle = np.loadtxt(global_g_strDataFolderNameOutHospital + '/各个截段点及中点位置' + r'\Side_spin_angle.txt')
        Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve = np.mean(Side_spin_angle)#6截面曲线旋转角均值(代数和)
        print('6截面曲线旋转角均值(代数和)=',Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve)

#   7截面曲线旋转角均值(绝对值)
        abs_Mean_Side_spin_angle_of_Section_Curve = np.mean(abs(np.array(Side_spin_angle)))  #7截面曲线旋转角均值(绝对值)
        print('7截面曲线旋转角均值(绝对值)=', abs_Mean_Side_spin_angle_of_Section_Curve)

#   8两肩高度差（左高为正）
        Height_difference_between_shoulders = final_fMarkerPoint_end[-2:-1, 1] - final_fMarkerPoint_end[-1:, 1]
        print('8两肩高度差（左高为正）=',Height_difference_between_shoulders)
#   9胸段旋转角（右高为正）（选前9段）
        Thoracic_rotation_angle = np.mean(Side_spin_angle[:9])
        print('9胸段旋转角（右高为正）（选前9段）=', Thoracic_rotation_angle)
#   10胸腰段旋转角（右高为正）(选接下来的10段)
        Thoracolumbar_rotation_angle = np.mean(Side_spin_angle[9:19])
        print('10胸腰段旋转角（右高为正）=', Thoracolumbar_rotation_angle)
        # </editor-fold>
        # <editor-fold desc="11-15">
        #   11腰段旋转角（右高为正）(选接下来的5段)
        Lumbar_rotation_angle = np.mean(Side_spin_angle[19:24])
        print('11腰段旋转角（右高为正）=', Lumbar_rotation_angle)
##############201890429:12到14移至plot_index和refresh
#   12冠状面C1弯曲率半径
        Main_bend_rad_to_save = Main_bend_rad
        print('12冠状面C1弯曲率半径=', Main_bend_rad_to_save)
#   13冠状面C1弯cobb角（右弯为正）
#         | a |=√[x1 ^ 2 + y1 ^ 2]
#         | b |=√[x2 ^ 2 + y2 ^ 2]
#         a * b = (x1, y1)(x2, y2) = x1x2 + y1y2
#         cos = a * b / [ | a | * | b |]
#         =(x1x2 + y1y2) / [√[x1 ^ 2 + y1 ^ 2] *√[x2 ^ 2 + y2 ^ 2]]
        vector_1 = [Main_bend_Boundary_points[0][0] - Main_bend_center[0],Main_bend_Boundary_points[1][0] - Main_bend_center[1]]
        vector_2 = [Main_bend_Boundary_points[0][1] - Main_bend_center[0],Main_bend_Boundary_points[1][1] - Main_bend_center[1]]
        cobber_radian = math.acos((vector_1[0] * vector_2[0] + vector_1[1] * vector_2[1])/(Main_bend_rad**2))
        cobber = cobber_radian *180/math.pi
        print('冠状面C1弯cobb角（右弯为正）=', cobber)
        print('注意正负')

        #20190611
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setFrame(False)#边框设置
        # self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setClearButtonEnabled(True)#删除按钮
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setAlignment(QtCore.Qt.AlignCenter)#居中
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setStyleSheet("color:green")
        font = QtGui.QFont()
        font.setFamily('微软雅黑')
        font.setBold(True)
        font.setPointSize(15)
        font.setWeight(75)
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setFont(font)

        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setFrame(False)  # 边框设置
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setAlignment(QtCore.Qt.AlignCenter)#居中
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setStyleSheet("color:green")
        font = QtGui.QFont()
        font.setFamily('微软雅黑')
        font.setBold(True)
        font.setPointSize(15)
        font.setWeight(75)
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setFont(font)

        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setText(str(Main_bend_rad_to_save))
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setText(str(cobber))
        # Main_bend_center, Main_bend_Boundary_points, Main_bend_rad
    #   14侧弯类型（s,c,多c）
        Side_bend_type_character = ['C弯','C弯','S弯','多C弯']
        Side_bend_number = len(finalCBendfBKPtCor)
        if Side_bend_number > 3:
            Side_bend_number = 3
        Side_bend_type = Side_bend_type_character[Side_bend_number]
        print('14侧弯类型（s,c,多c）=', Side_bend_type)
#   15冠状面脊柱曲线围成的面积（绝对值）
        #################   15冠状面脊柱曲线围成的面积（绝对值）（髂后上棘中点坐标为（0，0，0））
        #################   16冠状面脊柱曲线围成的面积（代数值）
        Neck_seven_x = final_fMarkerPoint_end[-5:-4, 1]
        Neck_seven_y = final_fMarkerPoint_end[-5:-4, 0]
        Neck_seven_z = final_fMarkerPoint_end[-5:-4, 2]
        # 获得颈七与髂后上棘中点的连线方程
        # y = (Neck_seven_y/Neck_seven_x)*x为原方程
        global  Spine_Fitting_Curve
        xnew = Spine_Fitting_Curve[:,1]
        ynew = Spine_Fitting_Curve[:,0]
        new_straight_line_y = (Neck_seven_y / Neck_seven_x) * xnew  # 以原来的y值作为自变量（脊柱直线方程（颈七和髂后上棘中点连线））

        mean_of_abs_differences = np.mean(abs(ynew - new_straight_line_y))  # 绝对值均值
        Sum_of_abs_differences = mean_of_abs_differences * Neck_seven_x  # 绝对值和
        print('mean_of_abs_differences=', mean_of_abs_differences)
        print('Sum_of_abs_differences=', Sum_of_abs_differences)
        # </editor-fold>
        # <editor-fold desc="16-20">
        #   16冠状面脊柱曲线围成的面积（代数值）
        mean_of_differences = np.mean(ynew - new_straight_line_y)  # 代数和均值
        Sum_of_differences = mean_of_differences * Neck_seven_x  # 冠状面脊柱曲线围成的面积
        print('mean_of_differences=', mean_of_differences)
        print('Sum_of_differences=', Sum_of_differences)

#   17冠状面脊柱平均直线倾角(右侧为正)
#         y = (Neck_seven_y/Neck_seven_x)*x #为原方程
        new_straight_line_slope = Neck_seven_y/Neck_seven_x  #冠状面脊柱平均直线斜率
        linear_inclination_of_coronal_spine = math.atan(new_straight_line_slope) * 180 / math.pi
        print('17冠状面脊柱平均直线倾角(右侧为正)=', linear_inclination_of_coronal_spine)
#   18驼背倾角
        if Max_Z_Left_and_Right[0][2] > Max_Z_Left_and_Right[1][2]:
            Max_Z_Thoracic_segment = Max_Z_Left_and_Right[0][2]
            Max_Y_Thoracic_segment = Max_Z_Left_and_Right[0][1]
        else:
            Max_Z_Thoracic_segment = Max_Z_Left_and_Right[1][2]
            Max_Y_Thoracic_segment = Max_Z_Left_and_Right[1][1]
        Hunchback_inclination_slope = abs((Max_Y_Thoracic_segment - Neck_seven_y)/(Max_Z_Thoracic_segment - Neck_seven_z))
        Hunchback_inclination = math.atan(Hunchback_inclination_slope) * 180 / math.pi
        print('18驼背倾角=', Hunchback_inclination)
#   19截面曲线旋转角离散度
        rotation_angle_Variance = np.var(Side_spin_angle)
        print('19截面曲线旋转角离散度=', rotation_angle_Variance)
#   20侧屈极限角(°)

        filepath = r'D:\B\zuoye\postgraduate\SecondGradePostgraduateTheFirstHalfOfTheSemester\！数据&任务\27+\直写式20190218'
        # filepath = os.getcwd()
        # 'Standing_Posture_Right_Limit_fMarkerPoint','Standing_Posture_Left_Limit_fMarkerPoint',
        try:
            Limit_Left_fMarkerPoint = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Correction_Standing_Posture_Left_Limit_fMarkerPoint.txt')  # filepath
            Limit_Right_fMarkerPoint = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Correction_Standing_Posture_Right_Limit_fMarkerPoint.txt')  # filepath
        except BaseException:
            pass
        else:
            #弧度
            Lateral_buckling_limit_right_radian = math.atan(-Limit_Left_fMarkerPoint[-5:-4,0]/Limit_Left_fMarkerPoint[-5:-4,1])
            Lateral_buckling_limit_left_radian = math.atan(Limit_Right_fMarkerPoint[-5:-4,0]/Limit_Right_fMarkerPoint[-5:-4,1])
            #角度
            Lateral_buckling_limit_right_angle =Lateral_buckling_limit_right_radian * 180/np.pi
            Lateral_buckling_limit_left_angle =Lateral_buckling_limit_left_radian * 180/np.pi
            # print('Limit_Left_fMarkerPoint =', Limit_Left_fMarkerPoint[-5:-4,0])
            # print('Limit_Left_fMarkerPoint=', Limit_Left_fMarkerPoint[-5:-4,1])
            print('20侧屈极限角(°) =',Lateral_buckling_limit_right_angle,Lateral_buckling_limit_left_angle)
        # </editor-fold>

#   21侧屈15左右围成面积
        '''对整个脊柱曲线进行计算'''
        all_fMarkerPoint_name = ['Standing_Posture_Right_fMarkerPoint', 'Standing_Posture_Left_fMarkerPoint']  # 文件夹名称
        try:
            Spine_Fitting_Curve_Right = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Spine_Fitting_Curve_' + all_fMarkerPoint_name[
                    0] + '.txt')  # 右侧屈15脊柱曲线#filepath
            Spine_Fitting_Curve_Left = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\Spine_Fitting_Curve_' + all_fMarkerPoint_name[
                    1] + '.txt')  # 左侧屈15脊柱曲线#filepath
        except BaseException:
            pass
        else:
            Left_and_right_difference = Spine_Fitting_Curve_Right[:,0] + Spine_Fitting_Curve_Left[:,0]#本身为一正一负
            mean_Left_and_right_difference = np.mean(Left_and_right_difference)
            print('21侧屈15左右围成面积 =',mean_Left_and_right_difference)
            # plt.grid(ls='--', linewidth=0.75)  # 生成网格
#   22左右侧屈15水平方向差值和（均值）
        '''对脊柱曲线上抽样点进行计算'''
        # Pick_points_ind = np.linspace(0,len(Spine_Fitting_Curve_Right)-1,30)
        # Pick_points_ind = list(np.round(Pick_points_ind))#四舍五入取整
        # for i in range(len(Pick_points_ind)):
        #     Pick_points_ind[i] = math.floor(Pick_points_ind[i])
        # Pick_points_dif_x = Spine_Fitting_Curve_Right[Pick_points_ind, 0] + Spine_Fitting_Curve_Left[Pick_points_ind, 0]#对称差值，而不是值的差值
        # Pick_points_dif_y = Spine_Fitting_Curve_Right[Pick_points_ind, 1]
        # # plt.plot(Pick_points_dif_x,Pick_points_dif_y)
        try:
            plot_Left_and_right_difference = np.loadtxt(
                g_strDataFolderNameOutHospital + r'\动态标记点及脊柱曲线\plot_Left_and_right_difference.txt')  # filepath
        except BaseException:
            pass
        else:
            dif_between_left_and_right_lateral_flexion_horizontal_direction = np.mean(plot_Left_and_right_difference[:,0])
            print('22左右侧屈15水平方向差值均值 =',dif_between_left_and_right_lateral_flexion_horizontal_direction)

        root = Tk()
        root.withdraw()  # ****实现tk主窗口隐藏
        messagebox.showinfo("提示", "指标计算完成！")
        # box = QMessageBox .information(self, '提 示' '指标计算完成！', QMessageBox.Yes)  # 完成提醒窗口

        print('指标计算完成')
############################抽屉链接窗口处理区
    def MainWndFn_MenuLinktoSubWnd(self):
#患者信息录入
         if self.toolBox.currentIndex() == 0:
             # 把QSplitter的指定位置的窗体从QSplitter中剥离
             self.splitter.widget(1).setParent(None)                    #在QSplitter的指定位置载入新窗体，但要先用上面的“self.splitter.widget(1).setParent(None)”命令。
             self.splitter.insertWidget(1, self.DlgPatientInfoInput)
             # 此函数用于设定：控件是否可伸缩。第一个参数用于指定控件的序号。第二个函数大于0时，表示控件可伸缩，小于0时，表示控件不可伸缩。
             self.splitter.setStretchFactor(0, 0) 
             self.splitter.setStretchFactor(1, 1)
             #  设置 QSplitter 分割器各部分最小化时的情况，设置为“False”意味着左右拉动分隔栏时各部分不会消失；此设置也可以在界面设计时在 QtDesigner 里设置。
             self.splitter.setChildrenCollapsible(False)
             #  设置 QSplitter 分割器随主窗口自适应大小变化。此设置也可以在界面设计时在 QtDesigner 里设置。
             self.splitter.setAutoFillBackground(True)
#数据预处理             
         elif self.toolBox.currentIndex() == 1:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_DataPreprocessing)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)             
#云纹图及垂线倾角
         elif self.toolBox.currentIndex() == 2:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_ResultShowFirst)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#冠、矢状面曲线图
         elif self.toolBox.currentIndex() == 3:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_ResultShowSecond)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#后背截面曲线和倾角
         elif self.toolBox.currentIndex() == 4:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_ResultShowThird)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#坐姿左右侧屈分析
         elif self.toolBox.currentIndex() == 5:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_Result_dynamic)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#冠状面C弯
         elif self.toolBox.currentIndex() == 6:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_Coronal_c_bend)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#足底压力分析
         elif self.toolBox.currentIndex() == 7:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.Ui_FootPressAnalyze)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#x光片分析
         elif self.toolBox.currentIndex() == 8:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.XPictureAnalyze)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#量表分析
         elif self.toolBox.currentIndex() == 9:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.DlgSystemSet)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#坐姿调整系统
         elif self.toolBox.currentIndex() == 10:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.DlgSystemSet)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#足部辅具设计
         elif self.toolBox.currentIndex() == 11:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.DlgSystemSet)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#统计分析
         elif self.toolBox.currentIndex() == 12:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.DlgStatisticsAnalyzing)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)
#系统设置
         elif self.toolBox.currentIndex() == 13:
             self.splitter.widget(1).setParent(None)
             self.splitter.insertWidget(1, self.DlgSystemSet)
             self.splitter.setStretchFactor(0, 0)
             self.splitter.setStretchFactor(1, 1)
             self.splitter.setChildrenCollapsible(False)
             self.splitter.setAutoFillBackground(True)

    def CBendStart_EndPoint(self):#获得冠状面C弯起止点
        global finalCBendfBKPtCor
        # 获得手动标的拐点
        CBendfBKPtCor = plt.ginput(2, timeout=-1)  # 右键撤销上一个点
        finalCBendfBKPtCor.append(CBendfBKPtCor)
        print('选择的两个点x,y值为：', CBendfBKPtCor)

    def CBendRefresh(self):
        global minx, maxx, miny, maxy, Spine_Fitting_Curve, finalCBendfBKPtCor
        global Main_bend_center,Main_bend_Boundary_points,Main_bend_rad

        global Main_bend_rad_to_save       #   12冠状面C1弯曲率半径
        global cobber#   13冠状面C1弯cobb角（右弯为正）注意正负
        global Side_bend_type       #   14侧弯类型（s,c,多c）

        # plt.cla()
        # self.Ui_Coronal_c_bend.gridLayout_2.removeWidget(self.figure_Coronal_C_bend_map)
        # self.Ui_Coronal_c_bend.gridLayout_2.remove(self.figure_Coronal_C_bend_map)
        import sip
        sip.delete(self.canvas_Coronal_C_bend_map)  # 删除addwidget的窗口
        # sip.delete(self.toolbar)
        '''手动标冠状面C弯'''
        self.figure_Coronal_C_bend_map = plt.figure()
        self.canvas_Coronal_C_bend_map = FigureCanvas(self.figure_Coronal_C_bend_map)
        # self.toolbar = NavigationToolbar(self.canvas_Coronal_C_bend_map, self)  # 窗口操作条
        # self.Ui_Coronal_c_bend.gridLayout_2.addWidget(self.toolbar)
        self.Ui_Coronal_c_bend.gridLayout_2.addWidget(self.canvas_Coronal_C_bend_map)
        self.axes7 = self.figure_Coronal_C_bend_map.add_subplot(111)
        self.axes7.axis('equal')  # 等比例显示
        self.axes7.set_xlim((minx, maxx))
        self.axes7.set_ylim((miny, maxy))
        # self.axes.set_zlim((minz-200, maxz+200))
        self.axes7.set_xticks(np.linspace(minx, maxx, 8))
        self.axes7.set_yticks(np.linspace(miny, maxy, 15))
        # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
        self.axes7.set_xlabel('水平方向(mm)', fontproperties="SimHei")
        self.axes7.set_ylabel('身高方向(mm)', fontproperties="SimHei")
        plt.grid(ls='--', linewidth=0.75)  # 生成网格
        # 画过零点直线
        # plt.plot([minx - 100, maxx + 100], [0, 0], color='slategray')  # lavender
        plt.plot(Spine_Fitting_Curve[:, 0], Spine_Fitting_Curve[:, 1], '.', marker='.',
                 markersize=3, color='b')  # 画脊柱拟合曲线
        xnew = Spine_Fitting_Curve[:, 1]
        ynew = Spine_Fitting_Curve[:, 0]

        #     预测点！#计算边界值（极值点位置）
        First_derivative = [(ynew[1] - ynew[0]) / (xnew[1] - xnew[0])]  # 为了减少循环，先给一阶导初值赋第一个值
        Second_derivative = []
        for i in range(len(xnew) - 2):
            dif_y = ynew[i + 2] - ynew[i + 1]
            dif_x = xnew[i + 2] - xnew[i + 1]
            First_derivative.append(dif_y / dif_x)  # 身体上下方向为x，以上下方向作为自变量计算
            dif_y2 = First_derivative[i + 1] - First_derivative[i]
            Second_derivative.append(dif_y2 / dif_x)
        ind_Inflection_point = np.where(abs(np.array(Second_derivative)) < 0.0002)  # 脊柱曲线的拐点初筛
        ind_Inflection_point = ind_Inflection_point[0]
        correct_ind_Inflection_point = []  # 最后的拐点的初始化
        start_ind = 0
        for i in range(len(ind_Inflection_point) - 1):  # 遍历所有可能的拐点
            #     if count_Inflection_point_num == 1 :
            if ((ind_Inflection_point[i + 1] - ind_Inflection_point[i]) >= 2):  # 找到除最后一个拐点的拐点（通过和下一个点的距离来判定是否属于同一个拐点）
                correct_ind_Inflection_point.append(math.ceil(np.mean(ind_Inflection_point[start_ind:i+1])))#+1  20190611
                start_ind = i + 1
            if (i == (len(ind_Inflection_point) - 2)):  # 补最后一个拐点
                correct_ind_Inflection_point.append(math.ceil(np.mean(ind_Inflection_point[start_ind:])))
        # correct_ind_Inflection_point = [0] + correct_ind_Inflection_point  # 增加最下面点为拐点
        correct_ind_Inflection_point = correct_ind_Inflection_point + [len(ynew) - 1]  # 增加最上面面点为拐点
        print(correct_ind_Inflection_point)
        plt.plot(ynew[correct_ind_Inflection_point], xnew[correct_ind_Inflection_point], '.', marker='.', markersize=8,
                 color='r')
        all_rad = []  # 所有拟合圆半径初始化
        all_centerx = []  # 所有拟合圆圆心x值初始化
        all_centery = []  # 所有拟合圆圆心y值初始化
        all_min_distance_ind = []#所有离获取鼠标点最近的索引初始化
        if len(finalCBendfBKPtCor) != 0:
            for i in range(len(finalCBendfBKPtCor)):#遍历每组点
                # plt.plot(finalCBendfBKPtCor[i][0][0], finalCBendfBKPtCor[i][0][1], '.', marker='.', markersize=8, color='r')#画第i组点的第一个点（鼠标点的位置）
                # plt.plot(finalCBendfBKPtCor[i][1][0], finalCBendfBKPtCor[i][1][1], '.', marker='.',
                #          markersize=8, color='r')#画第i组点的第二个点（鼠标点的位置）
                distance_first_point = (ynew - finalCBendfBKPtCor[i][0][0]) ** 2 + (
                        xnew - finalCBendfBKPtCor[i][0][1]) ** 2  # 计算第一个点与各个点的距离
                distance_second_point = (ynew - finalCBendfBKPtCor[i][1][0]) ** 2 + (
                        xnew - finalCBendfBKPtCor[i][1][1]) ** 2  # 计算第一个点与各个点的距离
                distance_first_point = list(distance_first_point)  # 将np.array变量转为普通的列表，因为在后面用到list.index，是列表才有的属性
                distance_second_point = list(distance_second_point)
                #
                sort_distance_first_point = np.sort(distance_first_point)  # 将距离数组按从小到大排列的顺序
                sort_distance_second_point = np.sort(distance_second_point)
                # min_distance_ind_first = []  # 初始化从小到大排列的数在原始数据中的索引位置
                # min_distance_ind_second = []
                # # for i in range(30):  ### 找到距离三个点最近的30个点在原始数据中的索引
                min_distance_ind_first = distance_first_point.index(sort_distance_first_point[0])
                min_distance_ind_second = distance_second_point.index(sort_distance_second_point[0])
                min_distance_ind = [min_distance_ind_first , min_distance_ind_second]
                plt.plot(ynew[min_distance_ind], xnew[min_distance_ind], '.', marker='.', markersize=10,color='k')
                all_min_distance_ind.append(min_distance_ind)
                # # <editor-fold desc="画拟合圆">
                # ###########画拟合圆
                if min_distance_ind_first > min_distance_ind_second:
                    begin_ind = min_distance_ind_second
                    end__ind = min_distance_ind_first
                else:
                    begin_ind = min_distance_ind_first
                    end__ind = min_distance_ind_second
                points = Spine_Fitting_Curve[begin_ind : end__ind,:]  # 获取俩个拐点间的脊柱曲线点
                centerx, centery, rad = self.circleLeastFit(points)  # 调用获取拟合圆半径和圆心坐标的函数
                print(centerx, centery, rad)
                all_rad.append(rad)
                all_centerx.append(centerx)
                all_centery.append(centery)
                # 方法一：参数方程
                theta = np.arange(0, 2 * np.pi, 0.01)  # 设置参数方程自变量角度
                # range()返回的是range object，而np.nrange()返回的是numpy.adarray()
                # 两者都是均匀地（evenly）等分区间；
                # range尽可用于迭代，而np.arange作用远不止于此，它是一个序列，可被当做向量使用。
                # range()不支持步长为小数，np.arange()支持步长为小数
                x = centerx + rad * np.cos(theta)  # 圆的参数方程
                y = centery + rad * np.sin(theta)  # 圆的参数方程
                plt.plot(x, y, c='green', ls='-', lw=1)#画拟合圆!!
                plt.axis('equal')  # 等比例显示
            ###########画主弯拟合圆
            ind_Main_bend = np.where(all_rad == np.min(all_rad))
            ind_Main_bend = ind_Main_bend[0][0]
            print('ind_Main_bend ', ind_Main_bend)
            Main_bend_rad = all_rad[ind_Main_bend]#圆半径
            Main_bend_centerx = all_centerx[ind_Main_bend]#圆心x值
            Main_bend_centery = all_centery[ind_Main_bend]#圆心y值
            ind_Main_bend_Boundary_points = all_min_distance_ind[ind_Main_bend]#手动选的主弯的位置索引
            plt.plot(ynew[ind_Main_bend_Boundary_points], xnew[ind_Main_bend_Boundary_points], '.', marker='.', markersize=18,color='m')# 显示主弯的俩个端点
            Main_bend_center = [Main_bend_centerx,Main_bend_centery]#全局变量
            Main_bend_Boundary_points = [ynew[ind_Main_bend_Boundary_points], xnew[ind_Main_bend_Boundary_points]]#全局变量
            theta = np.arange(0, 2 * np.pi, 0.01)  # 设置参数方程自变量角度
            x = Main_bend_centerx + Main_bend_rad * np.cos(theta)  # 圆的参数方程
            y = Main_bend_centery + Main_bend_rad * np.sin(theta)  # 圆的参数方程
            plt.plot(x, y, c='r', ls='-', lw=1)  # 画主弯拟合圆
            # # </editor-fold>
            #
            # print('ind_Inflection_point=', ind_Inflection_point)
            #######################
            self.axes7.axis('equal')  # 等比例显示
            self.axes7.set_xlim((minx, maxx))
            self.axes7.set_ylim((miny, maxy))
            # self.axes.set_zlim((minz-200, maxz+200))
            self.axes7.set_xticks(np.linspace(minx, maxx, 8))
            self.axes7.set_yticks(np.linspace(miny, maxy, 15))
            # self.axes.set_zticks(np.linspace(minz-50, maxz+50, 1))
            self.axes7.set_xlabel('水平方向(mm)', fontproperties="SimHei")
            self.axes7.set_ylabel('身高方向(mm)', fontproperties="SimHei")
            plt.savefig(g_strDataFolderNameOutHospital + "\数据结果图片" + r'\C弯选取示意图.png', format='png')#20190622
        #   12冠状面C1弯曲率半径
        Main_bend_rad_to_save = Main_bend_rad
        print('12冠状面C1弯曲率半径=', Main_bend_rad_to_save)
        #   13冠状面C1弯cobb角（右弯为正）
        #         | a |=√[x1 ^ 2 + y1 ^ 2]
        #         | b |=√[x2 ^ 2 + y2 ^ 2]
        #         a * b = (x1, y1)(x2, y2) = x1x2 + y1y2
        #         cos = a * b / [ | a | * | b |]
        #         =(x1x2 + y1y2) / [√[x1 ^ 2 + y1 ^ 2] *√[x2 ^ 2 + y2 ^ 2]]
        vector_1 = [Main_bend_Boundary_points[0][0] - Main_bend_center[0],
                    Main_bend_Boundary_points[1][0] - Main_bend_center[1]]
        vector_2 = [Main_bend_Boundary_points[0][1] - Main_bend_center[0],
                    Main_bend_Boundary_points[1][1] - Main_bend_center[1]]
        cobber_radian = math.acos((vector_1[0] * vector_2[0] + vector_1[1] * vector_2[1]) / (Main_bend_rad ** 2))
        cobber = cobber_radian * 180 / math.pi
        print('13冠状面C1弯cobb角（右弯为正）=', cobber)
        print('注意正负')
        # Main_bend_center, Main_bend_Boundary_points, Main_bend_rad
        #   14侧弯类型（s,c,多c）
        Side_bend_type_character = ['C弯', 'C弯', 'S弯', '多C弯']
        Side_bend_number = len(finalCBendfBKPtCor)
        if Side_bend_number > 3:
            Side_bend_number = 3
        Side_bend_type = Side_bend_type_character[Side_bend_number]
        print('14侧弯类型（s,c,多c）=', Side_bend_type)

        #20190611
        self.Ui_Coronal_c_bend.CoronalCBendingRatio_lineEdit.setText(str(Main_bend_rad_to_save))
        self.Ui_Coronal_c_bend.CoronalCBendingCobberAngle_lineEdit.setText(str(cobber))

        print('刷新完毕')

    def CBenddelet(self):
        global finalCBendfBKPtCor
        finalCBendfBKPtCor = []
        print('清除完毕')

    # def mysql_creat_insert(self):

    def save_to_word(self):
        # <editor-fold desc="指标全局变量...">
        print('开始生成word')
        global Coronal_trunk_imbalance#1冠状面躯干失平衡度
        global Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y#   2左右最高点水平倾角（左高为正）
        global Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z#   3左右最高点深度方向倾角（左高(凸)为正）
        global Height_dif_of_posterior_superior_iliac_spine#   4左右髂后上棘高度差（左高为正）
        global Horizontal_obliquity_of_pelvis#   5骨盆水平倾斜角（左高为正）
        global Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve    #   6截面曲线旋转角均值(代数和)(按右高为正算)
        global abs_Mean_Side_spin_angle_of_Section_Curve#   7截面曲线旋转角均值(绝对值)
        global Height_difference_between_shoulders#   8两肩高度差（左高为正）
        global Thoracic_rotation_angle#   9胸段旋转角（右高为正）（选前9段）
        global Thoracolumbar_rotation_angle#   10胸腰段旋转角（右高为正）(选接下来的10段)
        global Lumbar_rotation_angle       #   11腰段旋转角（右高为正）(选接下来的5段)
        global Main_bend_rad_to_save       #   12冠状面C1弯曲率半径
        global cobber#   13冠状面C1弯cobb角（右弯为正）注意正负
        global Side_bend_type       #   14侧弯类型（s,c,多c）
        global mean_of_abs_differences,Sum_of_abs_differences#   15冠状面脊柱曲线围成的面积（绝对值）##############
        global mean_of_differences, Sum_of_differences      #   16冠状面脊柱曲线围成的面积（代数值）#############
        global linear_inclination_of_coronal_spine      #   17冠状面脊柱平均直线倾角(右侧为正)
        global Hunchback_inclination#   18驼背倾角
        global rotation_angle_Variance#   19截面曲线旋转角离散度
        global Lateral_buckling_limit_right_angle,Lateral_buckling_limit_left_angle#   20侧屈极限角(°)##########
        global mean_Left_and_right_difference#   21侧屈15左右围成面积
        global dif_between_left_and_right_lateral_flexion_horizontal_direction#   22左右侧屈15水平方向差值和（均值）
        global g_strDataFolderNameOutHospital
        global item2
        try:
            # infofilepath = np.loadtxt(g_strDataFolderNameOutHospital + r"/被测者基本信息路径.txt", dtype=str)  # 字符型加dtype=str
            #20190617
            infodata = np.loadtxt(r'D:\灵犀脊柱形态评价系统\基本信息' + r'/' + item2 + '/informationsave.txt',dtype=str)  # 20180622没有该文件名对应的基本信息
            #infofilepath = np.loadtxt(r'D:\灵犀脊柱形态评价系统\当前被测者基本信息路径.txt', dtype=str)  # 字符型加dtype=str20190622
            ##########20190622

            ############
        except BaseException:
            root = Tk()
            root.withdraw()  # ****实现主窗口隐藏
            # messagebox.showinfo("提示", "被测者基本信息没有填写！")
            messagebox.showinfo("提示", "没有该文件名对应的基本信息！") # 20180622
        else:

            # '%.3f' % a
            Coronal_trunk_imbalance[0] = '%.3f' % Coronal_trunk_imbalance[0]#1冠状面躯干失平衡度
            Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y = '%.3f' % Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y#   2左右最高点水平倾角（左高为正）
            Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z = '%.3f' % Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z#   3左右最高点深度方向倾角（左高(凸)为正）
            Height_dif_of_posterior_superior_iliac_spine[0] = '%.3f' % Height_dif_of_posterior_superior_iliac_spine[0]#   4左右髂后上棘高度差（左高为正）
            Horizontal_obliquity_of_pelvis = '%.3f' % Horizontal_obliquity_of_pelvis#   5骨盆水平倾斜角（左高为正）
            Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve = '%.3f' % Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve    #   6截面曲线旋转角均值(代数和)(按右高为正算)
            abs_Mean_Side_spin_angle_of_Section_Curve = '%.3f' % abs_Mean_Side_spin_angle_of_Section_Curve#   7截面曲线旋转角均值(绝对值)
            Height_difference_between_shoulders[0] = '%.3f' % Height_difference_between_shoulders[0]#   8两肩高度差（左高为正）
            Thoracic_rotation_angle = '%.3f' % Thoracic_rotation_angle#   9胸段旋转角（右高为正）（选前9段）
            Thoracolumbar_rotation_angle = '%.3f' % Thoracolumbar_rotation_angle#   10胸腰段旋转角（右高为正）(选接下来的10段)
            Lumbar_rotation_angle   = '%.3f' % Lumbar_rotation_angle     #   11腰段旋转角（右高为正）(选接下来的5段)
            Main_bend_rad_to_save    = '%.3f' % Main_bend_rad_to_save    #   12冠状面C1弯曲率半径
            cobber = '%.3f' % cobber#   13冠状面C1弯cobb角（右弯为正）注意正负
            #Side_bend_type  =   '%.3f' % Side_bend_type    #   14侧弯类型（s,c,多c）
            mean_of_abs_differences,Sum_of_abs_differences = '%.3f' % mean_of_abs_differences,Sum_of_abs_differences#   15冠状面脊柱曲线围成的面积（绝对值）##############
            mean_of_differences, Sum_of_differences = '%.3f' % mean_of_differences, Sum_of_differences      #   16冠状面脊柱曲线围成的面积（代数值）#############
            linear_inclination_of_coronal_spine  = '%.3f' % linear_inclination_of_coronal_spine    #   17冠状面脊柱平均直线倾角(右侧为正)
            Hunchback_inclination = '%.3f' % Hunchback_inclination#   18驼背倾角
            rotation_angle_Variance = '%.3f' % rotation_angle_Variance#   19截面曲线旋转角离散度

            try:
                Lateral_buckling_limit_right_angle = '%.3f' % Lateral_buckling_limit_right_angle
                Lateral_buckling_limit_left_angle = '%.3f' % Lateral_buckling_limit_left_angle  # 20侧屈极限角(°)##########
                mean_Left_and_right_difference = '%.3f' % mean_Left_and_right_difference  # 21侧屈15左右围成面积
                dif_between_left_and_right_lateral_flexion_horizontal_direction = '%.3f' % dif_between_left_and_right_lateral_flexion_horizontal_direction  # 22左右侧屈15水平方向差值和（均值）
            except NameError:
                pass
            # else:
            #     Lateral_buckling_limit_right_angle = '%.3f' % Lateral_buckling_limit_right_angle
            #     Lateral_buckling_limit_left_angle = '%.3f'%Lateral_buckling_limit_left_angle#   20侧屈极限角(°)##########
            #     mean_Left_and_right_difference = '%.3f' %  mean_Left_and_right_difference#   21侧屈15左右围成面积
            #     dif_between_left_and_right_lateral_flexion_horizontal_direction = '%.3f' % dif_between_left_and_right_lateral_flexion_horizontal_direction#   22左右侧屈15水平方向差值和（均值）

            # </editor-fold>

            # <editor-fold desc="写入word...">
            # ######20190513
            from docx import Document#pip install python-docx
            import win32com
            from win32com.client import Dispatch, constants
            root = os.getcwd()  # 获得当前路径 /home/dir1
            print(root)
    #        os.chdir(g_strDataFolderNameOutHospital + "\数据结果图片")
            images1 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\后背点云图.png"  # r 'D:\灵犀脊柱形态评价系统\数据结果图片\figure.png'# 保存在本地的图片
            images2 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\脊柱曲线垂线倾角.png"
            images3 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\冠状面曲线图.png"
            images4 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\矢状面曲线图.png"
            images5 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\后背截面曲线图.png"
            images6 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\后背截面倾角曲线图.png"
            try:
                images7 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\左右侧屈脊柱曲线图.png"
                images8 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\左右侧屈相同脊柱位置差值.png"
                images9 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\脊柱曲线相邻点倾角变化.png"
            except BaseException:
                pass
            # else:
            #     images7 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\左右侧屈脊柱曲线图.png"
            #     images8 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\左右侧屈相同脊柱位置差值.png"
            #     images9 = g_strDataFolderNameOutHospital + "\数据结果图片" + r"\脊柱曲线相邻点倾角变化.png"

            doc = Document()  # doc对象
            doc.add_picture(images1)  # 添加图, 设置宽度
            doc.add_picture(images2)
            doc.add_picture(images3)
            doc.add_picture(images4)
            doc.add_picture(images5)
            doc.add_picture(images6)
            try:
                doc.add_picture(images7)
                doc.add_picture(images8)
                doc.add_picture(images9)
            except BaseException:
                #20190611
                #        doc.save(r'D:\灵犀脊柱形态评价系统\数据结果图片\生成文档.docx')     # 保存路径
                os.chdir(root)
                word = win32com.client.Dispatch("word.application")
                # doc = word.Documents.Open(root + r'\脊柱健康检测报告（单次版）.doc')
                doc = word.Documents.Open(root + r'\脊柱健康检测报告（单次版无动态版）.doc')
            else:
                # doc.add_picture(images7)
                # doc.add_picture(images8)
                # doc.add_picture(images9)
                # 20190611
                #        doc.save(r'D:\灵犀脊柱形态评价系统\数据结果图片\生成文档.docx')     # 保存路径
                os.chdir(root)
                word = win32com.client.Dispatch("word.application")
                doc = word.Documents.Open(root + r'\脊柱健康检测报告（单次版）.doc')
                # doc = word.Documents.Open(root + r'\脊柱健康检测报告（单次版无动态版）.doc')
            print('word完')

            t = doc.Tables[0]
            t1 = doc.Tables[1]
            try:
                t2 = doc.Tables[2]
            except BaseException:
                pass
            # else:
            #     t2 = doc.Tables[2]
            d1 = t.Rows[2].Cells[1].Range.Text
            # infodata = np.loadtxt(str(infofilepath), dtype=str)#20180622
            infoindex = []
            for i in range(len(infodata)):
                infoindex.append(infodata[i].split('：')[1])  # 提取出该行的坐标参数

            t.Rows[0].Cells[1].Range.Text = infoindex[0]
            t.Rows[0].Cells[3].Range.Text = infoindex[1]
            t.Rows[0].Cells[5].Range.Text = infoindex[6]
            t.Rows[1].Cells[1].Range.Text = infoindex[7]
            t.Rows[1].Cells[3].Range.Text = infoindex[2]
            t.Rows[1].Cells[5].Range.Text = infoindex[5]#职业
            t.Rows[2].Cells[1].Range.Text = infoindex[3]
            t.Rows[2].Cells[3].Range.Text = infoindex[12]
            t.Rows[2].Cells[5].Range.Text = infoindex[9]
            t.Rows[3].Cells[1].Range.Text = infoindex[4]
            t.Rows[3].Cells[3].Range.Text = infoindex[13]#肺活量
            t.Rows[3].Cells[5].Range.Text = infoindex[8]
            t.Rows[6].Cells[1].Range.Text = Coronal_trunk_imbalance[0]
            t.Rows[6].Cells[4].Range.Text = Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y
            t.Rows[7].Cells[1].Range.Text = Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z
            t.Rows[7].Cells[4].Range.Text = Height_dif_of_posterior_superior_iliac_spine[0]
            t.Rows[8].Cells[1].Range.Text = Horizontal_obliquity_of_pelvis#5
            t.Rows[8].Cells[4].Range.Text = Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve
            t.Rows[9].Cells[1].Range.Text = abs_Mean_Side_spin_angle_of_Section_Curve
            t.Rows[9].Cells[4].Range.Text = Height_difference_between_shoulders[0]
            t.Rows[10].Cells[1].Range.Text = Thoracic_rotation_angle
            t.Rows[10].Cells[4].Range.Text = Thoracolumbar_rotation_angle#10
            t.Rows[11].Cells[1].Range.Text = Lumbar_rotation_angle
            t.Rows[11].Cells[4].Range.Text = Main_bend_rad_to_save
            t.Rows[12].Cells[1].Range.Text = cobber
            t.Rows[12].Cells[4].Range.Text = Side_bend_type
            t.Rows[13].Cells[1].Range.Text = mean_of_abs_differences#15
            t.Rows[13].Cells[4].Range.Text = mean_of_differences
            t.Rows[14].Cells[1].Range.Text = linear_inclination_of_coronal_spine
            t.Rows[14].Cells[4].Range.Text = Hunchback_inclination
            t.Rows[15].Cells[1].Range.Text = rotation_angle_Variance
            try:
                t.Rows[15].Cells[4].Range.Text = Lateral_buckling_limit_left_angle  # 20
                t.Rows[16].Cells[1].Range.Text = Lateral_buckling_limit_right_angle
                t.Rows[16].Cells[4].Range.Text = mean_Left_and_right_difference
                t.Rows[17].Cells[1].Range.Text = dif_between_left_and_right_lateral_flexion_horizontal_direction
            except BaseException:
                pass
            # else:
            #     t.Rows[15].Cells[4].Range.Text = Lateral_buckling_limit_left_angle  # 20
            #     t.Rows[16].Cells[1].Range.Text = Lateral_buckling_limit_right_angle
            #     t.Rows[16].Cells[4].Range.Text = mean_Left_and_right_difference
            #     t.Rows[17].Cells[1].Range.Text = dif_between_left_and_right_lateral_flexion_horizontal_direction


            t1.Rows[1].Cells[0].Range.InLineShapes.AddPicture(images1)
            t1.Rows[1].Cells[1].Range.InLineShapes.AddPicture(images2)
            t1.Rows[3].Cells[0].Range.InLineShapes.AddPicture(images3)
            t1.Rows[3].Cells[1].Range.InLineShapes.AddPicture(images4)
            t1.Rows[5].Cells[0].Range.InLineShapes.AddPicture(images5)
            t1.Rows[5].Cells[1].Range.InLineShapes.AddPicture(images6)
            try:
                t2.Rows[1].Cells[0].Range.InLineShapes.AddPicture(images7)
                t2.Rows[1].Cells[1].Range.InLineShapes.AddPicture(images8)
                t2.Rows[3].Cells[0].Range.InLineShapes.AddPicture(images9)
            except BaseException:
                pass
            # else:
            #     t2.Rows[1].Cells[0].Range.InLineShapes.AddPicture(images7)
            #     t2.Rows[1].Cells[1].Range.InLineShapes.AddPicture(images8)
            #     t2.Rows[3].Cells[0].Range.InLineShapes.AddPicture(images9)

            # t2.Rows[1].Cells[0].Range.InLineShapes.AddPicture(r'C:\Users\lxm\Desktop\picture\2.png')
            # t2.Rows[1].Cells[1].Range.InLineShapes.AddPicture(r'C:\Users\lxm\Desktop\picture\2.png')
            # t2.Rows[3].Cells[0].Range.InLineShapes.AddPicture(r'C:\Users\lxm\Desktop\picture\2.png')
            # t2.Rows[3].Cells[1].Range.InLineShapes.AddPicture(r'C:\Users\lxm\Desktop\picture\2.png')
            # global item2#20190622
            # path = g_strDataFolderNameOutHospital + "\数据结果图片" + r'/' + item2 +'.docx'
            global g_strDataFolderNameOutHospital_orl
            word_path = g_strDataFolderNameOutHospital_orl + r'/' + item2 +'.docx'#20190618
            doc.SaveAs(word_path)
            #doc.Close#20190618
            root = Tk()
            root.withdraw()  # ****实现主窗口隐藏
            messagebox.showinfo("提示", "生成word完成！")
            print('word完成')
            # </editor-fold>

            # <editor-fold desc="写入excel...">
            name = infoindex[0]
            try:
                Gender = infoindex[1]  #
                ID_number = infoindex[6]  #
                data = infoindex[8]  #
            except BaseException:
                Gender = 0  #
                ID_number = 0  #
                data = 0  #

            # index_1=Coronal_trunk_imbalance[0]#1冠状面躯干失平衡度
            index_1 = Coronal_trunk_imbalance[0]  # 1冠状面躯干失平衡度
            index_2 = Horizontal_Inclination_of_Left_and_Right_Highest_Points_Y  # 2左右最高点水平倾角（左高为正）
            index_3 = Horizontal_Inclination_of_Left_and_Right_Highest_Points_Z  # 3左右最高点深度方向倾角（左高(凸)为正）
            # index_4=Height_dif_of_posterior_superior_iliac_spine[0]#4左右髂后上棘高度差（左高为正）
            index_4 = Height_dif_of_posterior_superior_iliac_spine[0]  # 4左右髂后上棘高度差（左高为正）
            index_5 = Horizontal_obliquity_of_pelvis  # 5骨盆水平倾斜角（左高为正）
            index_6 = Algebra_Sum_Mean_Side_spin_angle_of_Section_Curve  # 6截面曲线旋转角均值(代数和)(按右高为正算)
            index_7 = abs_Mean_Side_spin_angle_of_Section_Curve  # 7截面曲线旋转角均值(绝对值)
            # index_8=Height_difference_between_shoulders[0]#8两肩高度差（左高为正）
            index_8 = Height_difference_between_shoulders[0]  # 8两肩高度差（左高为正）
            index_9 = Thoracic_rotation_angle  # 9胸段旋转角（右高为正）
            index_10 = Thoracolumbar_rotation_angle  # 10#10胸腰段旋转角（右高为正）
            index_11 = Lumbar_rotation_angle  # 11腰段旋转角（右高为正）
            index_12 = Main_bend_rad_to_save  # 12冠状面C1弯曲率半径
            index_13 = cobber  # 13冠状面C1弯cobb角（右弯为正）
            index_14 = Side_bend_type  # 14侧弯类型（s,c,多c）
            index_15 = mean_of_abs_differences  # 15#15冠状面脊柱曲线围成的面积（绝对值）
            index_16 = mean_of_differences  # 16冠状面脊柱曲线围成的面积（代数值）
            index_17 = linear_inclination_of_coronal_spine  # 17冠状面脊柱平均直线倾角(右侧为正)
            index_18 = Hunchback_inclination  # 18驼背倾角
            index_19 = rotation_angle_Variance  # 19截面曲线旋转角离散度
            try:
                index_20 = Lateral_buckling_limit_left_angle  # 20#20侧屈右极限角(°)
                index_21 = Lateral_buckling_limit_right_angle  # 20侧屈左极限角(°)
                index_22 = mean_Left_and_right_difference  # 21侧屈15左右围成面积
                index_23 = dif_between_left_and_right_lateral_flexion_horizontal_direction  # 22左右侧屈15水平方向差值和（均值）
            except BaseException:
                index_20 = []  # 20#20侧屈右极限角(°)
                index_21 = []  # 20侧屈左极限角(°)
                index_22 = []  # 21侧屈15左右围成面积
                index_23 = []  # 22左右侧屈15水平方向差值和（均值）
            # else:
            #     index_20 = Lateral_buckling_limit_left_angle  # 20#20侧屈右极限角(°)
            #     index_21 = Lateral_buckling_limit_right_angle  # 20侧屈左极限角(°)
            #     index_22 = mean_Left_and_right_difference  # 21侧屈15左右围成面积
            #     index_23 = dif_between_left_and_right_lateral_flexion_horizontal_direction  # 22左右侧屈15水平方向差值和（均值）


            index = [name, Gender, ID_number, data, index_1, index_2, index_3, index_4, index_5, index_6, index_7, index_8,
                     index_9, index_10,
                     index_11, index_12, index_13, index_14, index_15, index_16, index_17, index_18, index_19, index_20,
                     index_21, index_22,
                     index_23]
            # for i in range(len(index)):
            #     index[i] = index[i][:-2]
            # print(index)

            from xlutils import copy
            import xlrd
            excel_path = r'D:\灵犀脊柱形态评价系统\test_in_excel.xls'  # 文件路径
            # excel_path=unicode('D:\\测试.xls','utf-8')#识别中文路径
            rbook = xlrd.open_workbook(excel_path)  # 打开文件

            worksheet1 = rbook.sheet_by_name('Sheet1')  # 获取工作簿中所有表格中的的第一个表格
            rows_old = worksheet1.nrows  # 获取表格中已存在的数据的行数
            print(rows_old)
            position = 'A' + str(rows_old + 1)
            # rbook = xlrd.open_workbook(excel_path,formatting_info=True)#打开文件
            wbook = copy.copy(rbook)  # 复制文件并保留格式
            w_sheet = wbook.get_sheet(0)  # 索引sheet表
            row = rows_old  # 从0开始
            # col = 0
            # value = 20180803
            # w_sheet.write_row('F1',index)
            for col in range(len(index)):
                w_sheet.write(row, col, index[col])
            wbook.save(excel_path)  # 保存文件
            print('excel完成')
            # </editor-fold>

            # <editor-fold desc="data存入mysql...">
            global Case_number
            Case_number = '231'
            # # 使用 execute() 方法执行 SQL，如果表存在则删除
            # cursor.execute("DROP TABLE IF EXISTS patientsinformation")

            # ①查询出有多条数据时：
            # cursor.fetchone()：将只取最上面的第一条结果，返回单个元组如('id', 'name')，然后多次循环使用cursor.fetchone()，依次取得下一条结果，直到为空。
            # cursor.fetchall(): 将返回所有结果，返回二维元组，如(('id', 'name'), ('id', 'name')),
            # ②查询只有一条数据时：
            # cursor.fetchone()：将只返回一条结果，返回单个元组如('id', 'name')。
            # cursor.fetchall(): 也将返回所有结果，返回二维元组，如(('id', 'name'), )

            # 1 新建数据库和数据表

            try:
                # 数据库连接(未指定具体的数据库)
                db = pymysql.connect(host="localhost", user="root", password="lxmwork", charset='utf8',
                                     use_unicode=True)#123456
                # db = MySQLdb.connect(host, user, pw, charset='utf8')
                # 创建游标，通过连接与数据通信
                cursor = db.cursor()
                # 执行sql语句
                cursor.execute('show databases')
                databases_rows = cursor.fetchall()  # 获取所有数据库的名称
                print('所有数据库的名称=', databases_rows)
                # tmp = []
                # for row in databases_rows:
                #     tmp.append("%2s" % row)
                cursor.execute('create database if not exists ' + 'patientsinfordata')  # 不存在数据库则新建
                cursor.execute('show databases')
                databases_rows = cursor.fetchall()  # 获取所有数据库的名称
                print('新建以后所有数据库的名称=', databases_rows)
                # 数据库连接(指定具体的数据库)
                db = pymysql.connect(host="localhost", user="root", password="lxmwork", database="patientsinfordata",
                                     charset='utf8',
                                     use_unicode=True)  #
                cursor = db.cursor()
                cursor.execute('show tables')
                tables_rows = cursor.fetchall()  # 获取所有表的名称
                print('tables_rows=', tables_rows)
                cursor.execute("""create table if not exists patientsdata (

                                     病历号 VARCHAR(20) NOT NULL,
                                     1冠状面躯干失平横度  VARCHAR(20),
                                     2左右最高点水平倾角左高为正 VARCHAR(20),
                                     3左右最高点深度方向倾角左高凸为正 VARCHAR(20),
                                     4左右髂后上棘高度差左高为正 VARCHAR(20),
                                     5骨盆水平倾斜角左高为正 VARCHAR(20),
                                     6截面曲线旋转角均值代数和右高为正 VARCHAR(20),
                                     7截面曲线旋转角均值绝对值 VARCHAR(20),
                                     8两肩高度差左高为正 VARCHAR(20),
                                     9胸段旋转角右高为正 VARCHAR(20),
                                     10胸腰段旋转角右高为正 VARCHAR(20),
                                     11腰段旋转角右高为正 VARCHAR(20),
                                     12冠状面C1弯曲率 VARCHAR(20),
                                     13冠状面C1弯cobb角右弯为正 VARCHAR(20),
                                     14侧弯类型sc多弯 VARCHAR(20),
                                     15冠状面脊柱曲线围成的面积绝对值 VARCHAR(20),
                                     16冠状面脊柱曲线围成的面积代数值 VARCHAR(20),
                                     17冠状面脊柱平均直线倾角右侧为正 VARCHAR(20),
                                     18驼背倾角 VARCHAR(20),
                                     19截面曲线旋转角离散度 VARCHAR(20),
                                     20侧屈右极限角左 VARCHAR(20),
                                     20侧屈左极限角右 VARCHAR(20),
                                     21侧屈15左右围成面积 VARCHAR(20),
                                     22左右侧屈15水平方向差值和均值 VARCHAR(20)

                                     )engine=InnoDB default charset=utf8;""")  # 不存在数据表则新建,engine=InnoDB数据表引擎,charset=utf8字符集
                cursor.execute('show tables')
                tables_rows = cursor.fetchall()  # 获取所有表的名称
                print('新建表以后的所有表=', tables_rows)
                # 提交到数据库执行
                # </editor-fold>
                # db.commit()
                # 2 插入数据

                # <editor-fold desc="解决数据为空的情况...">
                # func = lambda x: x if x != '' else 0
                # name = func(name)  # 为空时赋值0
                # age = func(age)
                # </editor-fold>
                if Case_number == 0:  # 无姓名或病历号不进行插入操作
                    reply = QMessageBox.information(self, '提示', '无病历号', QMessageBox.Yes)
                else:
                    # sql = "INSERT INTO 'role' ('role_name', 'role_name_zh') VALUES ('{0}', '{1}')".format('role_name', '测试角色')
                    # sql = "INSERT INTO patientsdata(病历号,1冠状面躯干失平横度,2左右最高点水平倾角左高为正,3左右最高点深度方向倾角左高凸为正" \
                    #       " VALUES ('{0}', '{1}', '{2}', '{3}')".format(
                    #     Case_number, index_1, index_2, index_3)
                    # sql = "INSERT INTO patientsdata(病历号,1冠状面躯干失平横度) VALUES ('{0}', '{1}')".format(Case_number, index_1)
                    sql = "INSERT INTO patientsdata(病历号,1冠状面躯干失平横度,2左右最高点水平倾角左高为正,3左右最高点深度方向倾角左高凸为正,4左右髂后上棘高度差左高为正,5骨盆水平倾斜角左高为正,6截面曲线旋转角均值代数和右高为正,7截面曲线旋转角均值绝对值,8两肩高度差左高为正,9胸段旋转角右高为正,10胸腰段旋转角右高为正,11腰段旋转角右高为正,12冠状面C1弯曲率,13冠状面C1弯cobb角右弯为正,14侧弯类型sc多弯,15冠状面脊柱曲线围成的面积绝对值,16冠状面脊柱曲线围成的面积代数值,17冠状面脊柱平均直线倾角右侧为正,18驼背倾角, 19截面曲线旋转角离散度,20侧屈右极限角左,20侧屈左极限角右,21侧屈15左右围成面积 ,22左右侧屈15水平方向差值和均值)" \
                          " VALUES ('{0}', '{1}', '{2}', '{3}', '{4}', '{5}', '{6}', '{7}', '{8}', '{9}', '{10}', '{11}', '{12}', '{13}','{14}', '{15}', '{16}', '{17}', '{18}', '{19}', '{20}', '{21}', '{22}', '{23}')".format(
                        Case_number, index_1, index_2, index_3, index_4, index_5, index_6, index_7, index_8,index_9, index_10,index_11, index_12, index_13, index_14, index_15, index_16, index_17, index_18, index_19, index_20,index_21, index_22,index_23)
                    try:
                        # 执行sql语句
                        print('123')
                        cursor.execute(sql)
                        # 提交到数据库执行
                        db.commit()
                        print('指标插入数据库成功')
                        # 保存完成提醒窗口
                        reply = QMessageBox.information(self, '提示', '指标插入数据库成功', QMessageBox.Yes)
                    except pymysql.err.InternalError:
                        # 如果发生错误则回滚
                        db.rollback()
                        reply = QMessageBox.information(self, '提示', '请注意输入格式', QMessageBox.Yes)

            except BaseException:
                print("Error")
            finally:
                # 关闭数据库连接
                db.close()
            # </editor-fold>

    def Changed_attitude(self):
        #站姿坐姿选择
        global attitude
        global g_strDataFolderNameOutHospital
        global global_g_strDataFolderNameOutHospital

        print(self.DlgPatientInfoInput.checkBox_sit.isChecked())
        if self.DlgPatientInfoInput.checkBox_sit.isChecked()==False:
            attitude = '站'
        elif self.DlgPatientInfoInput.checkBox_sit.isChecked()==True:
            attitude = '坐'
        print(attitude)
        global_g_strDataFolderNameOutHospital = g_strDataFolderNameOutHospital + r'/' + attitude + '姿中立位'
        plot_contour_map(myshow, g_strDataFolderNameOutHospital, global_g_strDataFolderNameOutHospital, attitude)  # 画点云图

# #########################################窗口类定义

class Ui_DlgPatientInfoInput(QWidget, Ui_DlgPatientInfoInput):
    def __init__(self):
        super(Ui_DlgPatientInfoInput, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)
class Ui_XPictureAnalyze(QWidget, Ui_XPictureAnalyze):
    def __init__(self):
        super(Ui_XPictureAnalyze, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)
class Ui_SpineFusion(QWidget, Ui_SpineFusion):
    def __init__(self):
        super(Ui_SpineFusion, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)
class Ui_DlgStatisticsAnalyzing(QWidget, Ui_DlgStatisticsAnalyzing):
    def __init__(self):
        super(Ui_DlgStatisticsAnalyzing, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)



class myLabel(QLabel):
    filename=''
    clicked = pyqtSignal()
    #比例尺相关变量
    sc=[0, 0, 0, 0]
    scale=1
    scaleCount=0
    #显示的图片数量
    ImgCount=0
    #鼠标拖动直线相关变量
    xpress = 0
    ypress = 0
    xmove = 0
    ymove = 0
    pressed= False
    #状态与计数变量
    stateflag=0
    count=np.zeros(40)
    #任意位置cobb线相关变量
    x0 = []
    y0 = []
    x1 = []
    y1 = []
    #冠状面Cobb角1相关变量
    cc1x0 = []
    cc1y0 = []
    cc1x1 = []
    cc1y1 = []
    #冠状面Cobb角2相关变量
    cc2x0 = []
    cc2y0 = []
    cc2x1 = []
    cc2y1 = []
    #冠状面Cobb角3相关变量
    cc3x0 = []
    cc3y0 = []
    cc3x1 = []
    cc3y1 = []
    #矢状面Cobb1角相关变量
    sc1x0 = []
    sc1y0 = []
    sc1x1 = []
    sc1y1 = []
    #矢状面Cobb角2相关变量
    sc2x0 = []
    sc2y0 = []
    sc2x1 = []
    sc2y1 = []
    #股骨头相关变量
    fx0 = []
    fy0 = []
    fx1 = []
    fy1 = []
    #髂骨相关变量
    ilx0 = []
    ily0 = []
    ilx1 = []
    ily1 = []
    #坐骨相关变量
    isx0 = []
    isy0 = []
    isx1 = []
    isy1 = []
    #骶骨倾斜相关变量
    sax0 = []
    say0 = []
    sax1 = []
    say1 = []
    #骶骨相对倾斜角相关变量
    rsax0 = []
    rsay0 = []
    rsax1 = []
    rsay1 = []
    #Risser参数相关变量
    rs=[]
    #腰骶角相关变量
    l5x0=[]
    l5x1=[]
    l5y0=[]
    l5y1=[]
    #锁骨角度相关变量
    cox0 = []
    coy0 = []
    cox1 = []
    coy1 = []
    #T1倾斜度相关变量
    t1x0 = []
    t1y0 = []
    t1x1 = []
    t1y1 = []
    #腰椎前凸角相关变量
    lbx0=[]
    lby0=[]
    lbx1=[]
    lby1=[]
    #顶椎上弯偏移距离相关变量
    tvmu=[]
    #底椎下弯偏移距离相关变量
    tvml=[]
    #冠状面平衡相关变量
    cr=[]
    #矢状面平衡相关变量
    sg=[]
    #顶椎旋转相关变量
    tvsx0=[]
    tvsy0=[]
    tvsx1=[]
    tvsy1=[]
    ls=[]
    #上端椎旋转相关变量
    uvsx0=[]
    uvsy0=[]
    uvsx1=[]
    uvsy1=[]
    ls2=[]
    #下端椎旋转相关变量
    lvsx0=[]
    lvsy0=[]
    lvsx1=[]
    lvsy1=[]
    ls3=[]
    #锥体楔形变相关变量
    morphx=[]
    morphy=[]
    #锥体名称列表（锥体融合）
    SpineList=['','颈7', '胸1', '胸2', '胸3', '胸4', '胸5',  '胸6', '胸7', '胸8', '胸9', '胸10', '胸11', '胸12', '腰1', '腰2', '腰3', '腰4', '腰5']
    #画点相关变量
    centerx=[]
    centery=[]
    curvex=[]
    curvey=[]
    inflection=[]
    pointflag = False
    pointcount = 0
    groupcount =0
    #楔形变相关变量
    px = [[]]
    py = [[]]
    culpx = [[]]
    culpy = [[]]
    morphError=0
    morphResult=[]
    morphQues=[]
    morphNum=0
    #判断rsl与rsr是否进行了计算的变量
    CalRs=1
    #用于储存识别结果的数组
    #mask轮廓
    vert_set=[]
    #最小外接矩形
    box_set=[]
    #储存矩形的宽高数组，暂时没有使用
    rect_set=[]
    #矩形中心点的xy坐标
    list_center=[]
    #自动识别拟合出的脊柱曲线
    spacex=[]
    spacey=[]
    #自动识别出的拐点
    inflection_auto=[]
    #临时储存添加mask的标记
    addmaskx=[]
    addmasky=[]
    addmaskcount=0
    #中线相关变量
    mid=[]
    midcenter=[]
    # 边缘相关变量
    End = []
    Endcenter = []
    #result用于储存计算结果
    result=np.zeros(28)
    #bgimg用来储存读进来的图片
    bgimg=cv2.imread('x.jpg')

    # 函数功能：返回点集中第一个点的y坐标
    def takeFirstY(self, elem):
        return elem[0][1]

    # 函数功能：返回y坐标
    def takeY(self, elem):
        return elem[1]

    def mousePressEvent(self,event):
        global fImgScale
        global px
        global py
        global iNumPChose
        global g_fMainCStartPtScale
        global g_fMainCEndPtScale
        
        
        if iNumPChose>1:
            px.append(event.x()/fImgScale)		
            py.append(event.y()/fImgScale)
            print(iNumPChose)
        
        if iNumPChose==5:
            g_fMainCStartPtScale=abs(py[3]-py[2])/abs(py[3]-py[0])
            g_fMainCEndPtScale=abs(py[3]-py[1])/abs(py[3]-py[0])
            #print("ks=%f,kend=%f", g_fMainCStartPtScale, g_fMainCEndPtScale)
            print("ks="+str(g_fMainCStartPtScale)+"  kend="+str(g_fMainCEndPtScale) )
        iNumPChose=iNumPChose+1	

        print(str(event.x()/fImgScale)+','+str(event.y()/fImgScale))
        if self.stateflag==1 and self.scaleCount< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.sc[0]=event.x()
         self.sc[1]=event.y()
        if self.stateflag==2 and self.count[2]<4 and self.groupcount< 18:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.px[self.groupcount].append(event.x())
         self.py[self.groupcount].append(event.y())
         self.culpx[self.groupcount].append(event.x())  
         self.culpy[self.groupcount].append(event.y())
         self.count[2]=self.count[2]+1
        if self.stateflag==3 and self.count[1]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.x0.append(event.x())
         self.y0.append(event.y())
        if self.stateflag==4 and self.count[4]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.cc1x0.append(event.x())
         self.cc1y0.append(event.y())
        if self.stateflag==5 and self.count[5]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.cc2x0.append(event.x())
         self.cc2y0.append(event.y())
        if self.stateflag==6 and self.count[6]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.cc3x0.append(event.x())
         self.cc3y0.append(event.y())
        if self.stateflag==9 and self.count[9]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.sc1x0.append(event.x())
         self.sc1y0.append(event.y())
        if self.stateflag==10 and self.count[10]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.sc2x0.append(event.x())
         self.sc2y0.append(event.y())
        if self.stateflag==13 and self.count[13]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.fx0.append(event.x())
         self.fy0.append(event.y())
        if self.stateflag==14 and self.count[14]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.ilx0.append(event.x())
         self.ily0.append(event.y())
        if self.stateflag==15 and self.count[15]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.isx0.append(event.x())
         self.isy0.append(event.y())
        if self.stateflag==16 and self.count[16]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.sax0.append(event.x())
         self.say0.append(event.y())
        if self.stateflag==17 and self.count[17]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         if self.count[13]<1:
             self.fx0.append(event.x())
             self.fy0.append(event.y())
         elif self.count[16]<1:
             self.sax0.append(event.x())
             self.say0.append(event.y())
         else:
             print("已完成标记，不需要再次标定")
        if self.stateflag==18 and self.count[18]< 6:
         self.rs.append(event.x())
         self.rs.append(event.y())
         self.count[18]=self.count[18]+2
         if self.count[18]>= 6:
              self.stateflag = 0
        if self.stateflag==19 and self.count[19]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.cox0.append(event.x())
         self.coy0.append(event.y())
        if self.stateflag==20 and self.count[20]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.t1x0.append(event.x())
         self.t1y0.append(event.y())
        if self.stateflag==21 and self.count[21]< 1:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.lbx0.append(event.x())
         self.lby0.append(event.y())
        if self.stateflag==22 and self.count[22]< 8:
         self.tvmu.append(event.x())
         self.tvmu.append(event.y())
         self.count[22]=self.count[22]+2
         if self.count[22]>= 4:
              self.stateflag = 0
        if self.stateflag==23 and self.count[23]< 8:
         self.tvml.append(event.x())
         self.tvml.append(event.y())
         self.count[23]=self.count[23]+2
         if self.count[23]>= 4:
              self.stateflag = 0
        if self.stateflag==24 and self.count[24]< 4:
         self.cr.append(event.x())
         self.cr.append(event.y())
         self.count[24]=self.count[24]+2
         if self.count[24]>= 4:
              self.stateflag = 0
        if self.stateflag==25 and self.count[25]< 4:
         self.sg.append(event.x())
         self.sg.append(event.y())
         self.count[25]=self.count[25]+2
         if self.count[25]>= 4:
              self.stateflag = 0
        if self.stateflag==26 and self.count[26]< 3:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.tvsx0.append(event.x())
         self.tvsy0.append(event.y())
        if self.stateflag==26 and self.count[26]>= 3:
         self.ls.append(event.x())
         self.ls.append(event.y())
         self.count[26]=self.count[26]+1
         if self.count[26]>= 5:
              self.stateflag = 0
        if self.stateflag==27 and self.count[27]< 3:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.uvsx0.append(event.x())
         self.uvsy0.append(event.y())
        if self.stateflag==27 and self.count[27]>= 3:
         self.ls2.append(event.x())
         self.ls2.append(event.y())
         self.count[27]=self.count[27]+1
         if self.count[27]>= 5:
              self.stateflag = 0
        if self.stateflag==28 and self.count[28]< 3:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.lvsx0.append(event.x())
         self.lvsy0.append(event.y())
        if self.stateflag==28 and self.count[28]>= 3:
         self.ls3.append(event.x())
         self.ls3.append(event.y())
         self.count[28]=self.count[28]+1
         if self.count[28]>= 5:
            self.stateflag = 0
        if self.stateflag==29:
         self.centerx.append(event.x())
         self.centery.append(event.y())
        if self.stateflag==30 and self.count[30]< 2:
         self.pressed= True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.l5x0.append(event.x())
         self.l5y0.append(event.y())
        if self.stateflag==31:
         x = event.x()
         y = event.y()
         if len(self.list_center)>1:
             self.RemoveMask()
             min_dists, min_dist_idy = self.inflectiontree.query([x, y])
             self.list_center.pop(min_dist_idy)
             self.box_set.pop(min_dist_idy)
         elif len(self.list_center)==1:
             self.list_center.pop(0)
             self.box_set.pop(0)
         if len(self.list_center)>4:
             self.RecalculateCurve()
         else:
             self.spacex=[]
             self.spacey=[]
             self.inflection=[]
             self.inflection_auto=[]
        if self.stateflag==32:
         self.pressed = True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
         self.addmaskx.append(event.x())
         self.addmasky.append(event.y())
         self.addmaskcount=self.addmaskcount+1
        if self.stateflag==33:
         x = event.x()
         y = event.y()
         if len(self.mid)>1:
             self.RemoveMid()
             min_dists, min_dist_idy = self.inflectiontree.query([x,y])
             self.mid.pop(min_dist_idy)
             self.midcenter.pop(min_dist_idy)
         elif len(self.midcenter)==1:
             self.mid.pop(0)
             self.midcenters.pop(0)
        if self.stateflag==34:
         self.pressed = True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()
        if self.stateflag==35:
         x = event.x()
         y = event.y()
         if len(self.End)>1:
             self.RemoveEnd()
             min_dists, min_dist_idy = self.inflectiontree.query([x,y])
             self.End.pop(min_dist_idy)
             self.Endcenter.pop(min_dist_idy)
         elif len(self.midcenter)==1:
             self.End.pop(0)
             self.Endcenters.pop(0)
        if self.stateflag==36:
         self.pressed = True
         self.xpress = event.x()
         self.ypress = event.y()
         self.xmove = event.x()
         self.ymove = event.y()



        self.update()
    def mouseReleaseEvent(self,event):
        
        #if event.button() == Qt.LeftButton:
        self.clicked.emit()
        
        if self.stateflag==1 and self.scaleCount< 1:
         self.sc[2]=event.x()
         self.sc[3]=event.y()
         self.scaleCount=self.scaleCount+1
         self.pressed= False
         if self.scaleCount>= 1:
              self.stateflag = 0
         self.scale=np.sqrt(pow(self.sc[0]-self.sc[2], 2)+pow(self.sc[1]-self.sc[3], 2))/50
        self.update()
        if self.stateflag==2 and self.pointcount<4 and self.groupcount <18:
          self.px[self.groupcount].append(event.x())  
          self.py[self.groupcount].append(event.y())
          self.culpx[self.groupcount].append(event.x())  
          self.culpy[self.groupcount].append(event.y())
          self.count[2]=self.count[2]+1
          self.pressed= False
        if self.count[2]>=4:
          self.px.append([])
          self.py.append([])
          self.culpx.append([])
          self.culpy.append([])    
          self.groupcount=self.groupcount+1
          self.count[2]=0
        if self.stateflag==3 and self.count[1]< 2:
         self.x1.append(event.x())
         self.y1.append(event.y())
         self.count[3]=self.count[3]+1
         self.pressed= False
         if self.count[3]>= 2:
              self.stateflag = 0
        if self.stateflag==4 and self.count[4]< 2:
         self.cc1x1.append(event.x())
         self.cc1y1.append(event.y())
         self.count[4]=self.count[4]+1
         self.pressed= False
         if self.count[4]>= 2:
              self.stateflag = 0
        if self.stateflag==5 and self.count[5]< 2:
         self.cc2x1.append(event.x())
         self.cc2y1.append(event.y())
         self.count[5]=self.count[5]+1
         self.pressed= False
         if self.count[5]>= 2:
              self.stateflag = 0
        if self.stateflag==6 and self.count[6]< 2:
         self.cc3x1.append(event.x())
         self.cc3y1.append(event.y())
         self.count[6]=self.count[6]+1
         self.pressed= False
         if self.count[6]>= 2:
              self.stateflag = 0
        if self.stateflag==9 and self.count[9]< 2:
         self.sc1x1.append(event.x())
         self.sc1y1.append(event.y())
         self.count[9]=self.count[9]+1
         self.pressed= False
         if self.count[9]>= 2:
              self.stateflag = 0
        if self.stateflag==10 and self.count[10]< 2:
         self.sc2x1.append(event.x())
         self.sc2y1.append(event.y())
         self.count[10]=self.count[10]+1
         self.pressed= False
         if self.count[10]>= 2:
              self.stateflag = 0   
        if self.stateflag==13 and self.count[13]< 1:
         self.fx1.append(event.x())
         self.fy1.append(event.y())
         self.count[13]=self.count[13]+1
         self.pressed= False
         for i in range(0,len(self.fx1)):
            if self.fx0[i]>self.fx1[i]:
                temp=self.fx0[0]
                self.fx0[0]=self.fx1[0]
                self.fx1[0]=temp
                temp=self.fy0[0]
                self.fy0[0]=self.fy1[0]
                self.fy1[0]=temp
         if self.count[13]>= 1:
              self.stateflag = 0
        self.update()     
        if self.stateflag==14 and self.count[14]< 1:
         self.ilx1.append(event.x())
         self.ily1.append(event.y())
         self.count[14]=self.count[14]+1
         self.pressed= False
         for i in range(0,len(self.ilx1)):
            if self.ilx0[i]>self.ilx1[i]:
                temp=self.ilx0[0]
                self.ilx0[0]=self.ilx1[0]
                self.ilx1[0]=temp
                temp=self.ily0[0]
                self.ily0[0]=self.ily1[0]
                self.ily1[0]=temp
         if self.count[14]>= 1:
              self.stateflag = 0
        if self.stateflag==15 and self.count[15]< 1:
         self.isx1.append(event.x())
         self.isy1.append(event.y())
         self.count[15]=self.count[15]+1
         self.pressed= False
         for i in range(0,len(self.isx1)):
            if self.isx0[i]>self.isx1[i]:
                temp=self.isx0[0]
                self.isx0[0]=self.isx1[0]
                self.isx1[0]=temp
                temp=self.isy0[0]
                self.isy0[0]=self.isy1[0]
                self.isy1[0]=temp
         if self.count[15]>= 1:
              self.stateflag = 0
        if self.stateflag==16 and self.count[16]< 1:
         self.sax1.append(event.x())
         self.say1.append(event.y())
         self.count[16]=self.count[16]+1
         self.pressed= False
         for i in range(0,len(self.sax1)):
            if self.sax0[i]>self.sax1[i]:
                temp=self.sax0[0]
                self.sax0[0]=self.sax1[0]
                self.sax1[0]=temp
                temp=self.say0[0]
                self.say0[0]=self.say1[0]
                self.say1[0]=temp
         if self.count[16]>= 1:
              self.stateflag = 0
        self.update()
        if self.stateflag==17 :
         if self.count[13]<1:
             self.fx1.append(event.x())
             self.fy1.append(event.y())
             self.count[13]=self.count[11]+1
         elif self.count[16]<1:
             self.sax1.append(event.x())
             self.say1.append(event.y())
             self.count[16]=self.count[16]+1
         else:
             print("已完成标记，不需要再次标记")
         self.pressed= False
         if self.count[13]>= 1 and self.count[16]>= 1:
              self.stateflag = 0   
        if self.stateflag==19 and self.count[19]< 1:
         self.cox1.append(event.x())
         self.coy1.append(event.y())
         self.count[19]=self.count[19]+1
         self.pressed= False
         for i in range(0,len(self.cox1)):
            if self.cox0[i]>self.cox1[i]:
                temp=self.cox0[0]
                self.cox0[0]=self.cox1[0]
                self.cox1[0]=temp
                temp=self.coy0[0]
                self.coy0[0]=self.coy1[0]
                self.coy1[0]=temp
         if self.count[19]>= 1:
              self.stateflag = 0
        self.update()       
        if self.stateflag==20 and self.count[20]< 1:
         self.t1x1.append(event.x())
         self.t1y1.append(event.y())
         self.count[20]=self.count[20]+1
         self.pressed= False
         for i in range(0,len(self.t1x1)):
            if self.t1x0[i]>self.t1x1[i]:
                temp=self.t1x0[0]
                self.t1x0[0]=self.t1x1[0]
                self.t1x1[0]=temp
                temp=self.t1y0[0]
                self.t1y0[0]=self.t1y1[0]
                self.t1y1[0]=temp
         if self.count[20]>= 1:
              self.stateflag = 0
        self.update()   
        if self.stateflag==21 and self.count[21]< 1:
         self.lbx1.append(event.x())
         self.lby1.append(event.y())
         self.count[21]=self.count[21]+1
         self.pressed= False
         if self.count[21]>= 1:
              self.stateflag = 0
        if self.stateflag==26 and self.count[26]< 3:
         self.tvsx1.append(event.x())
         self.tvsy1.append(event.y())
         self.count[26]=self.count[26]+1
         self.pressed= False
        if self.stateflag==27 and self.count[27]< 3:
         self.uvsx1.append(event.x())
         self.uvsy1.append(event.y())
         self.count[27]=self.count[27]+1
         self.pressed= False
        if self.stateflag==28 and self.count[28]< 3:
         self.lvsx1.append(event.x())
         self.lvsy1.append(event.y())
         self.count[28]=self.count[28]+1
         self.pressed= False
        if self.stateflag==30 and self.count[30]< 2:
         self.l5x1.append(event.x())
         self.l5y1.append(event.y())
         self.count[30]=self.count[30]+1
         self.pressed= False
         if self.count[30]>= 2:
              self.stateflag = 0
        if self.stateflag==32:
         self.addmaskx.append(event.x())
         self.addmasky.append(event.y())
         self.addmaskcount=self.addmaskcount+1
         self.pressed= False
         if self.addmaskcount >= 4:
              list=[]
              self.addmaskx[1], self.addmaskx[2] = self.addmaskx[2], self.addmaskx[1]
              self.addmasky[1], self.addmasky[2] = self.addmasky[2], self.addmasky[1]
              for i in range(0,4):
                  list.append((self.addmaskx[i],self.addmasky[i]))
              list = np.array(list, dtype=np.int0)
              rect = cv2.minAreaRect(list)
              rotate_box = cv2.boxPoints(rect)
              rotate_box = np.int0(rotate_box)
              self.box_set.append(rotate_box)
              self.list_center.append((np.mean(rotate_box[:, 0]), np.mean(rotate_box[:, 1])))
              # 将box_set按takeFirstY的返回值排序（
              self.box_set.sort(key=self.takeFirstY)
              # 将list_center按takeY的返回值排序
              self.list_center.sort(key=self.takeY)
              if len(self.box_set)>3:
                  self.RecalculateCurve()
              self.addmaskx=[]
              self.addmasky=[]
              self.addmaskcount = 0
        if self.stateflag==34:
         self.mid.append([self.xpress,self.ypress,event.x(),event.y()])
         self.midcenter.append([int((self.xpress+event.x())/2), int((self.ypress+event.y())/2)])
         self.pressed= False
        if self.stateflag==36:
         self.End.append([self.xpress,self.ypress,event.x(),event.y()])
         self.Endcenter.append([int((self.xpress+event.x())/2), int((self.ypress+event.y())/2)])
         self.pressed= False
    def mouseMoveEvent(self,event):
        if self.stateflag==1 and self.scaleCount< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==2 and self.count[2]<4 and self.groupcount<18:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==3 and self.count[3]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==4 and self.count[4]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==5 and self.count[5]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==6 and self.count[6]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==9 and self.count[9]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()  
        if self.stateflag==10 and self.count[10]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==13 and self.count[13]< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==14 and self.count[14]< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==15 and self.count[15]< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==16 and self.count[16]< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()    
        if self.stateflag==17 and self.count[17]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()  
        if self.stateflag==19 and self.count[19]< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==20 and self.count[20]< 1:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update() 
        if self.stateflag==21 and self.count[21]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==22 and self.count[22]< 3:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==23 and self.count[23]< 3:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==26 and self.count[26]< 3:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==27 and self.count[27]< 3:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==28 and self.count[28]< 3:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==30 and self.count[30]< 2:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==32:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==34:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()
        if self.stateflag==36:
            self.xmove = event.x()
            self.ymove = event.y()
            self.update()

    def ScaleClear(self):
        self.sc=[0, 0, 0, 0]
        self.scale=1
        self.scaleCount = 0
        self.stateflag = 0
        self.update()
    def LineClear(self):
        self.x0 = []
        self.y0 = []
        self.x1 = []
        self.y1 = []
        self.count[3] = 0
        self.stateflag = 0
        self.update()
    def PointClear(self):
         self.centerx=[]
         self.centery=[]
         self.curvex=[]
         self.curvey=[]
         self.inflection=[]
         self.stateflag = 0
         self.update()
    def MorphClear(self):
        self.px = [[]]
        self.py = [[]]
        self.culpx = [[]]
        self.culpy = [[]]
        self.stateflag = 0
        self.count[2] = 0
        self.groupcount = 0
        self.morphError=0
        self.morphResult=[]
        self.morphQues=[]
        self.morphNum=0
        self.update()
    def FemurClear(self):
        self.fx0 = []
        self.fy0 = []
        self.fx1 = []
        self.fy1 = []
        self.count[13] = 0
        self.stateflag = 0
        self.update()
    def IliumClear(self):
        self.ilx0 = []
        self.ily0 = []
        self.ilx1 = []
        self.ily1 = []
        self.count[14] = 0
        self.stateflag = 0
        self.update()
    def IschiumClear(self):
        self.isx0 = []
        self.isy0 = []
        self.isx1 = []
        self.isy1 = []
        self.count[15] = 0
        self.stateflag = 0
        self.update()
    def T1Clear(self):
        self.t1x0 = []
        self.t1y0 = []
        self.t1x1 = []
        self.t1y1 = []
        self.count[20] = 0
        self.stateflag = 0
        self.update()
    def CollarboneClear(self):
        self.cox0 = []
        self.coy0 = []
        self.cox1 = []
        self.coy1 = []
        self.count[19] = 0
        self.stateflag = 0
        self.update()
    def SacrumClear(self):
        self.sax0 = []
        self.say0 = []
        self.sax1 = []
        self.say1 = []
        self.count[16] = 0
        self.stateflag = 0
        self.update()
    def RelSacrumClear(self):
        self.fx0 = []
        self.fy0 = []
        self.fx1 = []
        self.fy1 = []
        self.count[13] = 0
        self.sax0 = []
        self.say0 = []
        self.sax1 = []
        self.say1 = []
        self.count[16] = 0
        self.stateflag = 0
        self.update()
    def RelL5Clear(self):
        self.l5x0 = []
        self.l5y0 = []
        self.l5x1 = []
        self.l5y1 = []
        self.count[30] = 0
        self.stateflag = 0
        self.update()
    def CoronalCobb1Clear(self):
        self.cc1x0 = []
        self.cc1y0 = []
        self.cc1x1 = []
        self.cc1y1 = []
        self.count[4] = 0
        self.stateflag = 0
        self.update()
    def CoronalCobb2Clear(self):
        self.cc2x0 = []
        self.cc2y0 = []
        self.cc2x1 = []
        self.cc2y1 = []
        self.count[5] = 0
        self.stateflag = 0
        self.update()
    def CoronalCobb3Clear(self):
        self.cc3x0 = []
        self.cc3y0 = []
        self.cc3x1 = []
        self.cc3y1 = []
        self.count[6] = 0
        self.stateflag = 0
        self.update()
    def SagittalCobb1Clear(self):
        self.sc1x0 = []
        self.sc1y0 = []
        self.sc1x1 = []
        self.sc1y1 = []
        self.count[9] = 0
        self.stateflag = 0
        self.update()
    def SagittalCobb2Clear(self):
        self.sc2x0 = []
        self.sc2y0 = []
        self.sc2x1 = []
        self.sc2y1 = []
        self.count[10] = 0
        self.stateflag = 0
        self.update()
    def CoronalClear(self):
        self.cr = []
        self.count[24] = 0
        self.stateflag = 0
        self.update()
    def TVMUperClear(self):
        self.tvmu = []
        self.count[22] = 0
        self.stateflag = 0
        self.update()
    def TVMLowerClear(self):
        self.tvml = []
        self.count[23] = 0
        self.stateflag = 0
        self.update()
    def SagittalClear(self):
        self.sg = []
        self.count[25] = 0
        self.stateflag = 0
        self.update()
    def LumbarClear(self):
        self.lbx0 = []
        self.lby0 = []
        self.lbx1 = []
        self.lby1 = []
        self.count[21] = 0
        self.stateflag = 0
        self.update()
    def RsClear(self):
        self.rs = []
        self.count[18] = 0
        self.stateflag = 0
        self.update()
    def TVSClear(self):
        self.tvsx0 = []
        self.tvsy0 = []
        self.tvsx1 = []
        self.tvsy1 = []
        self.ls=[]
        self.count[26] = 0
        self.stateflag = 0
        self.update()
    def UVSClear(self):
        self.uvsx0 = []
        self.uvsy0 = []
        self.uvsx1 = []
        self.uvsy1 = []
        self.ls2=[]
        self.count[27] = 0
        self.stateflag = 0
        self.update()
    def LVSClear(self):
        self.lvsx0 = []
        self.lvsy0 = []
        self.lvsx1 = []
        self.lvsy1 = []
        self.ls3=[]
        self.count[28] = 0
        self.stateflag = 0
        self.update()
    def CalculateAngle(self):
        global g_fLineAngleForCobb
        if self.count[3] >= 2:
            if (self.x1[0]-self.x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.y1[0]+self.y0[0])/(self.x1[0]-self.x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 4)
            text1='第一条直线的角度是：'+str(a1)
            print(text1)
            if (self.x1[1]-self.x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.y1[1]+self.y0[1])/(self.x1[1]-self.x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 4)
            text2='第二条直线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='两条直线的转角是：'+str(angle)
            self.result[0]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]    
         ##################################
    def CalculatePoint(self):
        if len(self.centery)>=3 and len(self.centery) == len(set(self.centery)):
            #拟合脊柱曲线
            aryPx=np.zeros(len(self.centerx), dtype=int)
            aryPy=np.zeros(len(self.centery), dtype=int)
            for m in range(0, len(self.centerx)):
               aryPx[m]=int(self.centerx[m])
               aryPy[m]=int(self.centery[m])
            ynew = np.linspace(min(aryPy), max(aryPy), 500)
            f_yx =scipy.interpolate.interp1d(aryPy, aryPx, kind="cubic")#为了满足唯一性要求，这里使用Y做自变量
            xnew=f_yx(ynew)
            print("曲线计算完成")
            for n in range(0, len(xnew)):
               self.curvex.append(int(xnew[n]))
               self.curvey.append(int(ynew[n]))
            #计算拐点      
            iflc=[]   
            para=np.polyfit(ynew, xnew, 7)
            for n in range(1, len(ynew)-1):
               fdrv2=42*para[0]*pow(ynew[n-1], 5)+30*para[1]*pow(ynew[n-1], 4)+20*para[2]*pow(ynew[n-1], 3)+12*para[3]*pow(ynew[n-1], 2)+6*para[4]*ynew[n-1]+2*para[5]
               drv2=42*para[0]*pow(ynew[n], 5)+30*para[1]*pow(ynew[n], 4)+20*para[2]*pow(ynew[n], 3)+12*para[3]*pow(ynew[n], 2)+6*para[4]*ynew[n]+2*para[5]
               bdrv2=42*para[0]*pow(ynew[n+1], 5)+30*para[1]*pow(ynew[n+1], 4)+20*para[2]*pow(ynew[n+1], 3)+12*para[3]*pow(ynew[n+1], 2)+6*para[4]*ynew[n+1]+2*para[5]
               if abs(drv2)<=0.02 and (bool(fdrv2>=0) != bool(bdrv2>=0)):
                   iflc.append(n)
            for o in range(0, len(iflc)):
                iflcIndex=iflc[o]
                self.inflection.append(iflcIndex)
        self.update()
    
        
        
    def CalculateFemur(self):
        if self.count[13] >= 1:
            if (self.fx1[0]-self.fx0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.fy1[0]+self.fy0[0])/(self.fx1[0]-self.fx0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            l1=abs(self.fy0[0]-self.fy1[0])
            l1=l1/self.scale
            l1=round(l1, 2)
            text1='股骨头第一个标记点坐标是：'+str(self.fx0[0])+','+str(self.fy0[0])
            text2='股骨头第二个标记点坐标是：'+str(self.fx1[0])+','+str(self.fy1[0])
            text3='股骨头标记连线的角度是：'+str(a1)
            self.result[10]=a1
            print(text1)
            print(text2)
            print(text3)
    def CalculateIlium(self):
        if self.count[14] >= 1:
            if (self.ilx1[0]-self.ilx0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.ily1[0]+self.ily0[0])/(self.ilx1[0]-self.ilx0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='髂骨顶点第一个标记点坐标是：'+str(self.ilx0[0])+','+str(self.ily0[0])
            text2='髂骨顶点第二个标记点坐标是：'+str(self.ilx1[0])+','+str(self.ily1[0])
            text3='髂骨顶点连线的角度是：'+str(a1)
            self.result[11]=a1
            print(text1)
            print(text2)
            print(text3)
    def CalculateIschium(self):
        if self.count[15] >= 1:
            if (self.isx1[0]-self.isx0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.isy1[0]+self.isy0[0])/(self.isx1[0]-self.isx0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='坐骨底点第一个标记点坐标是：'+str(self.isx0[0])+','+str(self.isy0[0])
            text2='坐骨底点第二个标记点坐标是：'+str(self.isx1[0])+','+str(self.isy1[0])
            text3='坐骨底点连线的角度是：'+str(a1)
            self.result[12]=a1
            print(text1)
            print(text2)
            print(text3)
    def CalculateT1(self):
        if self.count[20] >= 1:
            if (self.t1x1[0]-self.t1x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.t1y1[0]+self.t1y0[0])/(self.t1x1[0]-self.t1x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='T1第一个标记点坐标是：'+str(self.t1x0[0])+','+str(self.t1y0[0])
            text2='T1第二个标记点坐标是：'+str(self.t1x1[0])+','+str(self.t1y1[0])
            text3='T1标记连线的角度是：'+str(a1)
            self.result[17]=a1
            print(text1)
            print(text2)
            print(text3)  
    def CalculateCollarbone(self):
        if self.count[19] >= 1:
            if (self.cox1[0]-self.cox0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.coy1[0]+self.coy0[0])/(self.cox1[0]-self.cox0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='锁骨第一个标记点坐标是：'+str(self.cox0[0])+','+str(self.coy0[0])
            text2='锁骨第二个标记点坐标是：'+str(self.cox1[0])+','+str(self.coy1[0])
            text3='锁骨标记连线的角度是：'+str(a1)
            self.result[16]=a1
            print(text1)
            print(text2)
            print(text3)        
    def CalculateSacrum(self):
        if self.count[16] >= 1:
            if (self.sax1[0]-self.sax0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.say1[0]+self.say0[0])/(self.sax1[0]-self.sax0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='骶骨第一个标记点坐标是：'+str(self.sax0[0])+','+str(self.say0[0])
            text2='骶骨第二个标记点坐标是：'+str(self.sax1[0])+','+str(self.say1[0])
            text3='骶骨连线的角度是：'+str(a1)
            self.result[13]=a1
            print(text1)
            print(text2)
            print(text3) 
    def CalculateRelSacrumCobb(self):
        global g_fLineAngleForCobb
        if self.count[13] >= 1:
            if (self.fx1[0]-self.fx0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.fy1[0]+self.fy0[0])/(self.fx1[0]-self.fx0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            self.result[10]=a1
            text1='股骨头标记连线的角度是：'+str(a1)
            print(text1)
        if self.count[16] >= 1:
            if (self.sax1[0]-self.sax0[0]) == 0:
                a2=-90.00
            else:
                k2=(-self.say1[0]+self.say0[0])/(self.sax1[0]-self.sax0[0])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            self.result[13]=a2
            text2='骶骨连线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='骶骨相对倾斜角是：'+str(angle)
            self.result[14]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]       
    def CalculateCoronalCobb1(self):
        global g_fLineAngleForCobb
        if self.count[4] >= 2:
            if (self.cc1x1[0]-self.cc1x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.cc1y1[0]+self.cc1y0[0])/(self.cc1x1[0]-self.cc1x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='第一条直线的角度是：'+str(a1)
            print(text1)
            if (self.cc1x1[1]-self.cc1x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.cc1y1[1]+self.cc1y0[1])/(self.cc1x1[1]-self.cc1x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            text2='第二条直线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='两条直线的转角是：'+str(angle)
            self.result[1]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]   
    def CalculateCoronalCobb2(self):
        global g_fLineAngleForCobb
        if self.count[5] >= 2:
            if (self.cc2x1[0]-self.cc2x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.cc2y1[0]+self.cc2y0[0])/(self.cc2x1[0]-self.cc2x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='第一条直线的角度是：'+str(a1)
            print(text1)
            if (self.cc2x1[1]-self.cc2x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.cc2y1[1]+self.cc2y0[1])/(self.cc2x1[1]-self.cc2x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            text2='第二条直线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='两条直线的转角是：'+str(angle)
            self.result[2]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]         
    def CalculateCoronalCobb3(self):
        global g_fLineAngleForCobb
        if self.count[6] >= 2:
            if (self.cc3x1[0]-self.cc3x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.cc3y1[0]+self.cc3y0[0])/(self.cc3x1[0]-self.cc3x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='第一条直线的角度是：'+str(a1)
            print(text1)
            if (self.cc3x1[1]-self.cc3x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.cc3y1[1]+self.cc3y0[1])/(self.cc3x1[1]-self.cc3x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            text2='第二条直线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='两条直线的转角是：'+str(angle)
            self.result[3]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]       
    def CalculateSagittalCobb1(self):
        global g_fLineAngleForCobb
        if self.count[9] >= 2:
            if (self.sc1x1[0]-self.sc1x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.sc1y1[0]+self.sc1y0[0])/(self.sc1x1[0]-self.sc1x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='第一条直线的角度是：'+str(a1)
            print(text1)
            if (self.sc1x1[1]-self.sc1x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.sc1y1[1]+self.sc1y0[1])/(self.sc1x1[1]-self.sc1x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            text2='第二条直线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='两条直线的转角是：'+str(angle)
            self.result[6]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]    
    def CalculateSagittalCobb2(self):
        global g_fLineAngleForCobb
        if self.count[10] >= 2:
            if (self.sc2x1[0]-self.sc2x0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.sc2y1[0]+self.sc2y0[0])/(self.sc2x1[0]-self.sc2x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='第一条直线的角度是：'+str(a1)
            print(text1)
            if (self.sc2x1[1]-self.sc2x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.sc2y1[1]+self.sc2y0[1])/(self.sc2x1[1]-self.sc2x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            text2='第二条直线的角度是：'+str(a2)
            print(text2)
            angle=round(a2-a1, 2)
            text3='两条直线的转角是：'+str(angle)
            self.result[7]=angle
            print(text3)
            g_fLineAngleForCobb=[a1, a2, angle]   
    def CalculateCoronal(self):
        if self.count[24] >= 4:
            l1=self.cr[0]-self.cr[2]
            l1=l1/self.scale
            l1=round(l1, 2)
            text1='C7PL的坐标是：'+str(self.cr[0])+','+str(self.cr[1])
            text2='CSVL的坐标是：'+str(self.cr[2])+','+str(self.cr[3])
            text3='C7PL的偏移距离（左-右+）是：'+str(l1)
            self.result[21]=l1
            print(text1)
            print(text2)
            print(text3)  
    def CalculateTVMUper(self):
        if self.count[22] >= 4:
            l1=self.tvmu[0]-self.tvmu[2]
            l1=l1/self.scale
            l1=round(l1, 2)
            text1='C7PL的坐标是：'+str(self.tvmu[0])+','+str(self.tvmu[1])
            text2='上方最大偏移标记点的坐标是：'+str(self.tvmu[2])+','+str(self.tvmu[3])
            text3='上方的偏移距离（左-右+）是：'+str(l1)
            self.result[19]=l1
            print(text1)
            print(text2)
            print(text3) 
    def CalculateTVMLower(self):
        if self.count[23] >= 4:
            l1=self.tvml[0]-self.tvml[2]
            l1=l1/self.scale
            l1=round(l1, 2)
            text4='CSVL的坐标是：'+str(self.tvml[0])+','+str(self.tvml[1])
            text5='下方最大偏移标记点的坐标是：'+str(self.tvml[2])+','+str(self.tvml[3])
            text6='下方的偏移距离（左-右+）是：'+str(l1)
            self.result[20]=l1
            print(text4)
            print(text5)
            print(text6)
    def CalculateSagittal(self):
        if self.count[25] >= 4:
            l1=self.sg[0]-self.sg[2]
            l1=l1/self.scale
            l1=round(l1, 2)
            text1='C7的坐标是：'+str(self.sg[0])+','+str(self.sg[1])
            text2='L5的坐标是：'+str(self.sg[2])+','+str(self.sg[3])
            text3='C7L5之间的偏移距离（左-右+）是：'+str(l1)
            self.result[22]=l1
            print(text1)
            print(text2)
            print(text3) 
            
    def CalculateLumbar(self):
        if self.count[21] >= 1:
            if (self.lbx1[0]-self.lbx0[0]) == 0:
                 a1=0.0000
            elif (self.lby1[0]-self.lby0[0]) == 0:
                 a1=-90.00
            else:
                k1=(-self.lby1[0]+self.lby0[0])/(self.lbx1[0]-self.lbx0[0])
                k1=-1/k1
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            text1='第一个标记点坐标是：'+str(self.lbx0[0])+','+str(self.lby0[0])
            text2='腰第二个标记点坐标是：'+str(self.lbx1[0])+','+str(self.lby1[0])
            text3='腰椎前凸角（顺正逆负）是：'+str(a1)
            self.result[18]=a1
            print(text1)
            print(text2)
            print(text3)
    def CalculateRs(self):
        if self.count[18] >= 6:
            x1=self.rs[0]
            y1=-1*self.rs[1]
            x2=self.rs[2]
            y2=-1*self.rs[3]
            x3=self.rs[4]
            y3=-1*self.rs[5]
            xmax=max(x1, x2, x3)
            xmin=min(x1, x2, x3)
            ymax=max(-y1, -y2, -y3)
            ymin=min(-y1, -y2, -y3)
            img_cut=self.rsimg[ymin-5:ymax+5, xmin-5:xmax+5, ::-1]
            cv2.imwrite('cutImage.jpg', img_cut)
            print('risser point')
            print(self.rs[0]-xmin+5)
            print(self.rs[1]-ymin+5)
            print(self.rs[2]-xmin+5)
            print(self.rs[3]-ymin+5)
            print(self.rs[4]-xmin+5)
            print(self.rs[5]-ymin+5)
            if (x1-x2)*(y1-y3)-(x1-x3)*(y1-y2)==0:
                if x1==x3 and y1==y3 and (x2!=x3 or y2!=y3):
                    self.result[15]=0
                    text1='Risser参数是：0'
                elif x2==x3 and y2==y3 and (x1!=x3 or y1!=y3):
                    self.result[15]=1
                    text1='Risser参数是：1'
                else:
                    self.CalRs=0
                    text1="三个点共线，无法计算，请重新标记"
            else:
                a1=((x1*x1-x2*x2)+(y1*y1-y2*y2))/2
                a2=((x1*x1-x3*x3)+(y1*y1-y3*y3))/2
                x0=((y1-y2)*a2-(y1-y3)*a1)/((y1-y2)*(x1-x3)-(x1-x2)*(y1-y3))
                y0=((x1-x3)*a1-(x1-x2)*a2)/((y1-y2)*(x1-x3)-(x1-x2)*(y1-y3))
                if x0>x1:
                  if (x0-x1) == 0:
                     a1=-90.00
                  else:
                     k1=(-y0+y1)/(x1-x0)
                     a1=-1*round(math.atan(k1)/3.1415926*180, 4)
                  if (x0-x2) == 0:
                     a2=-90.00
                  else:
                     k2=(-y0+y2)/(x2-x0)
                     a2=-1*round(math.atan(k2)/3.1415926*180, 4)
                     if a2<0:
                         a2=a2+180
                
                  if (x0-x3) == 0:
                     a3=-90.00
                  else:
                     k3=(-y0+y3)/(x3-x0)
                     a3=-1*round(math.atan(k3)/3.1415926*180, 4)
                     if a3<0 and np.sqrt(pow(x3-x2, 2)+pow(y3-y2, 2)) < np.sqrt(pow(x3-x1, 2)+pow(y3-y1, 2)) :
                         a3=a3+180
                else:
                   if (x0-x1) == 0:
                      a1=90.00
                   else:
                      k1=(-y0+y1)/(x1-x0)
                      a1=round(math.atan(k1)/3.1415926*180, 4)
                   if (x0-x2) == 0:
                     a2=90.00
                   else:
                     k2=(-y0+y2)/(x2-x0)
                     a2=round(math.atan(k2)/3.1415926*180, 4)
                     if a2<0:
                         a2=a2+180
                   if (x0-x3) == 0:
                     a3=90.00
                   else:
                     k3=(-y0+y3)/(x3-x0)
                     a3=round(math.atan(k3)/3.1415926*180, 4)
                     if a3<0 and np.sqrt(pow(x3-x2, 2)+pow(y3-y2, 2)) < np.sqrt(pow(x3-x1, 2)+pow(y3-y1, 2)) :
                         a3=a3+180
                angleMax=round(a2-a1, 4)
                angleCur=round(a3-a1, 4)
                fraction=abs(angleCur/angleMax)
                fraction=round(fraction, 2)
                self.result[15]=fraction
                text1='Risser参数是：'+str(fraction)
            print(text1)
    def CalculateRelL5(self):
        if self.count[30] >= 2:
            if (self.l5x1[0]-self.l5x0[0]) == 0:
                a1=-90.00
            else:
                k1=(-self.l5y1[0]+self.l5y0[0])/(self.l5x1[0]-self.l5x0[0])
                a1=-1*round(math.atan(k1)/3.1415926*180, 2)
            self.result[26]=a1
            text1='腰4腰5连线的角度是：'+str(a1)
            print(text1)
            if (self.l5x1[1]-self.l5x0[1]) == 0:
                a2=-90.00
            else:
                k2=(-self.l5y1[1]+self.l5y0[1])/(self.l5x1[1]-self.l5x0[1])
                a2=-1*round(math.atan(k2)/3.1415926*180, 2)
            self.result[26]=a2
            text2='腰4腰5连线的角度是：'+str(a2)
            print(text2)
            if (a1-a2)>90:
                angle=180-(a1-a2)
            elif (a1-a2)<-90:
                angle=a1-a2+180
            else:
                angle=a1-a2
            self.result[26]=angle
    def CalculateTVS(self):
        if self.count[26] >= 4:
            level=0
            left=self.tvsx1[0]+self.tvsx0[0]
            left=left/2
            mid=self.tvsx1[1]+self.tvsx0[1]
            mid=mid/2
            right=self.tvsx1[2]+self.tvsx0[2]
            right=right/2
            if self.tvsx1[1]==self.tvsx0[1]:
              fCosMidline=1
            else:
              k1=(-self.tvsy1[1]+self.tvsy0[1])/(self.tvsx1[1]-self.tvsx0[1])
              a1=-1*round(math.atan(k1)/3.1415926*180, 2)
              a2=90.00-a1
              fCosMidline=cos(a2)
            max=abs(right-left)/fCosMidline
            if self.count[26] >= 5:
                if self.tvsx1[1]==self.tvsx0[1]:
                   d1=abs(self.ls[0]-mid)
                   d2=abs(self.ls[2]-mid)
                else:
                    b=-self.tvsy0[1]-k1*self.tvsx0[1]
                    d1=abs(self.ls[0]*k1+self.ls[1]+b)/np.sqrt(k1*k1+1)
                    d2=abs(self.ls[2]*k1+self.ls[3]+b)/np.sqrt(k1*k1+1)
                if abs(d1-d2)<=0.075*max :
                    level=0
                    self.result[23]=level
                elif (d1+d2)<=0.75*max :
                    level=2
                    self.result[23]=level
                else:
                    level=1
                    self.result[23]=level
                
            else:
                if self.tvsx1[1]==self.tvsx0[1]:
                    if abs(self.ls[0]-mid)<0.075*max:
                      level=3
                      self.result[23]=3
                    else:
                      level=4
                      self.result[23]=4
                else:
                    b=-self.tvsy0[1]-k1*self.tvsx0[1]
                    d1=abs(self.ls[0]*k1+self.ls[1]+b)/np.sqrt(k1*k1+1)
                    if d1<0.075*max:
                      level=3
                      self.result[23]=3
                    else:
                      level=4
                      self.result[23]=4
            text1='顶椎旋转等级是：'+str(level)+'级'
            print(text1)
    def CalculateUVS(self):
        if self.count[27] >= 4:
            level=0
            left=self.uvsx1[0]+self.uvsx0[0]
            left=left/2
            mid=self.uvsx1[1]+self.uvsx0[1]
            mid=mid/2
            right=self.uvsx1[2]+self.uvsx0[2]
            right=right/2
            if self.uvsx1[1]==self.uvsx0[1]:
              fCosMidline=1
            else:
              k1=(-self.uvsy1[1]+self.uvsy0[1])/(self.uvsx1[1]-self.uvsx0[1])
              a1=-1*round(math.atan(k1)/3.1415926*180, 2)
              a2=90.00-a1
              fCosMidline=cos(a2)
            max=abs(right-left)/fCosMidline
            if self.count[27] >= 5:
                if self.uvsx1[1]==self.uvsx0[1]:
                   d1=abs(self.ls2[0]-mid)
                   d2=abs(self.ls2[2]-mid)
                else:
                    b=-self.uvsy0[1]-k1*self.uvsx0[1]
                    d1=abs(self.ls2[0]*k1+self.ls2[1]+b)/np.sqrt(k1*k1+1)
                    d2=abs(self.ls2[2]*k1+self.ls2[3]+b)/np.sqrt(k1*k1+1)
                if abs(d1-d2)<=0.075*max :
                    level=0
                    self.result[24]=level
                elif (d1+d2)<=0.75*max :
                    level=2
                    self.result[24]=level
                else:
                    level=1
                    self.result[24]=level
                
            else:
                if self.uvsx1[1]==self.uvsx0[1]:
                    if abs(self.ls2[0]-mid)<0.075*max:
                      level=3
                      self.result[24]=3
                    else:
                      level=4
                      self.result[24]=4
                else:
                    b=-self.uvsy0[1]-k1*self.uvsx0[1]
                    d1=abs(self.ls2[0]*k1+self.ls2[1]+b)/np.sqrt(k1*k1+1)
                    if d1<0.075*max:
                      level=3
                      self.result[24]=3
                    else:
                      level=4
                      self.result[24]=4
            text1='下端椎旋转等级是：'+str(level)+'级'
            print(text1)
    def CalculateLVS(self):
        if self.count[28] >= 4:
            level=0
            left=self.lvsx1[0]+self.lvsx0[0]
            left=left/2
            mid=self.lvsx1[1]+self.lvsx0[1]
            mid=mid/2
            right=self.lvsx1[2]+self.lvsx0[2]
            right=right/2
            if self.lvsx1[1]==self.lvsx0[1]:
              fCosMidline=1
            else:
              k1=(-self.lvsy1[1]+self.lvsy0[1])/(self.lvsx1[1]-self.lvsx0[1])
              a1=-1*round(math.atan(k1)/3.1415926*180, 2)
              a2=90.00-a1
              fCosMidline=cos(a2)
            max=abs(right-left)/fCosMidline
            if self.count[28] >= 5:
                if self.lvsx1[1]==self.lvsx0[1]:
                   d1=abs(self.ls3[0]-mid)
                   d2=abs(self.ls3[2]-mid)
                else:
                    b=-self.lvsy0[1]-k1*self.lvsx0[1]
                    d1=abs(self.ls3[0]*k1+self.ls3[1]+b)/np.sqrt(k1*k1+1)
                    d2=abs(self.ls3[2]*k1+self.ls3[3]+b)/np.sqrt(k1*k1+1)
                if abs(d1-d2)<=0.075*max :
                    level=0
                    self.result[25]=level
                elif (d1+d2)<=0.75*max :
                    level=2
                    self.result[25]=level
                else:
                    level=1
                    self.result[25]=level
                
            else:
                if self.lvsx1[1]==self.lvsx0[1]:
                    if abs(self.ls3[0]-mid)<0.075*max:
                      level=3
                      self.result[25]=3
                    else:
                      level=4
                      self.result[25]=4
                else:
                    b=-self.lvsy0[1]-k1*self.lvsx0[1]
                    d1=abs(self.ls3[0]*k1+self.ls3[1]+b)/np.sqrt(k1*k1+1)
                    if d1<0.075*max:
                      level=3
                      self.result[25]=3
                    else:
                      level=4
                      self.result[25]=4
            text1='下端椎旋转等级是：'+str(level)+'级'
            print(text1)
    def CalculateMorph(self):
        a=5
        length_l=[]
        length_r=[]
        angel_l=[]
        angel_r=[]
        if self.groupcount>=18:
           for i in range(0, len(self.culpx)-1):
            if len(self.culpx[i])>=4:
              #根据x坐标的大小进行冒泡排序
             for j in range(0, 3):
                 for k in range(0, 3-j):
                      if self.culpx[i][k]>self.culpx[i][k+1]:
                          temp=self.culpx[i][k]
                          self.culpx[i][k]=self.culpx[i][k+1]
                          self.culpx[i][k+1]=temp
                          temp=self.culpy[i][k]
                          self.culpy[i][k]=self.culpy[i][k+1]
                          self.culpy[i][k+1]=temp
           #计算每个截断的左右边线长度和倾角
           for i in range(0, len(self.culpx)-1):
               if len(self.culpx[i])>=4:
                  if (self.culpx[i][0]-self.culpx[i][1])==0:
                       a1=90.00
                  else:
                       k1=(self.culpy[i][0]-self.culpy[i][1])/(self.culpx[i][1]-self.culpx[i][0])
                       a1=round(math.atan(k1)/3.1415926*180, 2)
                  angel_l.append(a1)
                  l1=math.sqrt(pow((self.culpy[i][0]-self.culpy[i][1]), 2)+pow((self.culpx[i][1]-self.culpx[i][0]), 2))
                  length_l.append(l1)
                  if (self.culpx[i][2]-self.culpx[i][3])==0:
                       a2=90.00
                  else:
                       k2=(self.culpy[i][2]-self.culpy[i][3])/(self.culpx[i][3]-self.culpx[i][2])
                       a2=round(math.atan(k2)/3.1415926*180, 2)
                  angel_r.append(a2)
                  l2=math.sqrt(pow((self.culpy[i][2]-self.culpy[i][3]), 2)+pow((self.culpx[i][2]-self.culpx[i][3]), 2))
                  length_r.append(l2)
           #计算相邻截断角度差和长度差
           for i in range(1, len(self.culpx)-2):
               sub_al=max(abs(angel_l[i]-angel_l[i+1]), abs(angel_l[i]-angel_l[i-1]))
               sub_ar=max(abs(angel_r[i]-angel_r[i+1]), abs(angel_r[i]-angel_r[i-1]))
               sub_l=abs(length_r[i]-length_l[i])/self.scale
               if max(length_r[i], length_l[i]) ==0:
                   self.morphError=self.morphError+1
               else:
                   ratio_l=sub_l/max(length_r[i], length_l[i])
                   if sub_al>=a or sub_ar>=a:
                       self.morphResult.append(i)
                       self.morphNum=self.morphNum+1 
                   elif sub_l>1.1:
                       self.morphResult.append(i)
                       self.morphNum=self.morphNum+1
                   elif ratio_l>0.25 and sub_l>=0.9:
                       self.morphResult.append(i)
                       self.morphNum=self.morphNum+1
                   elif ratio_l<=0.25 and sub_l>=0.9:
                       self.morphQues.append(i)
           #特殊情况1：索引为0
           sub_al1=abs(angel_l[0]-angel_l[1]) 
           sub_ar1=abs(angel_r[0]-angel_r[1]) 
           sub_l1=abs(length_r[0]-length_l[0])/self.scale
           if length_l[0] ==0:
                self.morphError=self.morphError+1
           else:
                ratio_l1=sub_l1/max(length_r[0], length_l[0])
                if sub_al1>=a or sub_ar1>=a:
                       self.morphResult.append(0)
                       self.morphNum=self.morphNum+1 
                elif sub_l1>1.1:
                       self.morphResult.append(0)
                       self.morphNum=self.morphNum+1
                elif ratio_l1>0.25 and sub_l1>=0.9:
                       self.morphResult.append(0)
                       self.morphNum=self.morphNum+1
                elif ratio_l1<=0.25 and sub_l1>=0.9:
                       self.morphQues.append(0)
           #特殊情况2：索引为17
           sub_al13=abs(angel_l[17]-angel_l[16]) 
           sub_ar13=abs(angel_r[17]-angel_r[16]) 
           sub_l13=abs(length_r[17]-length_l[17])/self.scale
           if length_l[17] ==0:
                self.morphError=self.morphError+1
           else:
                ratio_l13=sub_l13/max(length_r[17], length_l[17])
                if sub_al13>=a or sub_ar13>=a:
                       self.morphResult.append(17)
                       self.morphNum=self.morphNum+1 
                elif sub_l13>1.1:
                       self.morphResult.append(17)
                       self.morphNum=self.morphNum+1
                elif ratio_l13>0.25 and sub_l13>=0.9:
                       self.morphResult.append(17)
                       self.morphNum=self.morphNum+1
                elif ratio_l13<=0.25 and sub_l13>=0.9:
                       self.morphQues.append(17) 

         
          #结果：0=还没有计算，1=合格，2=不合格，3=综合临床综合判断
          #这里设置判断成角的度数，当前是5°

    #清除识别结果
    def MaskClear(self):
        # mask轮廓
        self.vert_set = []
        # 最小外接矩形
        self.box_set = []
        # 储存矩形的宽高数组，暂时没有使用
        self.rect_set = []
        # 矩形中心点的xy坐标
        self.list_center = []
        # 自动识别拟合出的脊柱曲线
        self.spacex = []
        self.spacey = []
        # 自动识别出的拐点
        self.inflection_auto = []
        # 临时储存添加mask的标记
        self.addmaskx = []
        self.addmasky = []
        self.addmaskcount = 0
    #重新拟合曲线
    def RecalculateCurve(self):
        # 拟合曲线
        # 将xy坐标分开保存成numpy数组，方便进行曲线拟合
        center = np.array(self.list_center, dtype=np.int0)
        centerx = center[:, 0]
        centery = center[:, 1]
        spacey = np.linspace(min(centery), max(centery), 500)
        f_yx = scipy.interpolate.interp1d(centery, centerx, kind='cubic')
        spacex = f_yx(spacey)
        self.spacex=spacex
        self.spacey=spacey
        # 计算拐点
        iflc = []
        self.inflection_auto=[]
        para = np.polyfit(spacey, spacex, 7)
        for n in range(1, len(spacey) - 1):
            fdrv2 = 42 * para[0] * pow(spacey[n - 1], 5) + 30 * para[1] * pow(spacey[n - 1], 4) + 20 * para[
                2] * pow(spacey[n - 1], 3) + 12 * para[3] * pow(spacey[n - 1], 2) + 6 * para[4] * spacey[
                        n - 1] + 2 * para[5]
            drv2 = 42 * para[0] * pow(spacey[n], 5) + 30 * para[1] * pow(spacey[n], 4) + 20 * para[2] * pow(
                spacey[n], 3) + 12 * para[3] * pow(spacey[n], 2) + 6 * para[4] * spacey[n] + 2 * para[5]
            bdrv2 = 42 * para[0] * pow(spacey[n + 1], 5) + 30 * para[1] * pow(spacey[n + 1], 4) + 20 * para[
                2] * pow(spacey[n + 1], 3) + 12 * para[3] * pow(spacey[n + 1], 2) + 6 * para[4] * spacey[
                        n + 1] + 2 * para[5]
            if abs(drv2) <= 0.02 and (bool(fdrv2 >= 0) != bool(bdrv2 >= 0)):
                iflc.append(n)
        for o in range(0, len(iflc)):
            iflcIndex = iflc[o]
            self.inflection_auto.append(iflcIndex)
    def MidClear(self):
        self.mid = []
        self.midcenter = []
    def EndClear(self):
        self.End = []
        self.Endcenter = []

    def RemoveMask(self):
        if len(self.list_center)>=2:
            self.inflectiontree = scipy.spatial.cKDTree(np.array(self.list_center))
    def RemoveMid(self):
        if len(self.mid)>=2:
            self.inflectiontree = scipy.spatial.cKDTree(np.array(self.midcenter))
    def RemoveEnd(self):
        if len(self.End)>=2:
            self.inflectiontree = scipy.spatial.cKDTree(np.array(self.Endcenter))

    def paintEvent(self, event):
        super().paintEvent(event)
        #rect =QRect(self.x0, self.y0, abs(self.x1-self.x0), abs(self.y1-self.y0))
        painter = QPainter(self)
        orange=QColor(255, 165, 0)
        pink=QColor(255, 192, 203)
        violet=QColor(238 , 130, 238)
        brown=QColor(165, 42, 42)
        oliver=QColor(154, 205, 50)
        Moccasin=QColor(255, 228, 181)
        Peru=QColor(205, 133, 63)
        Maroon=QColor(176, 48, 96)
        Bisque=QColor(255, 228, 196)
        Coral=QColor(255, 114, 86)
        Chartreuse=QColor(127, 255, 0)
        Chocolate=QColor(210,105,30)
        color_list=[orange,pink,violet,brown,oliver,Moccasin,Peru,Maroon,Bisque,Coral]
        painter.setPen(QPen(Qt.yellow,2,Qt.SolidLine))
        if self.pressed == True :
         painter.drawLine(self.xpress, self.ypress, self.xmove, self.ymove)
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        if self.scaleCount>=1:
            painter.drawLine(self.sc[0], self.sc[1], self.sc[2], self.sc[3])
        for i in range(0,len(self.x1)):
            painter.drawLine(self.x0[i], self.y0[i], self.x1[i], self.y1[i])
        painter.setPen(QPen(Qt.green,2,Qt.SolidLine))
        for j in range(len(self.px)):
            if len(self.px[j])>=2:
             painter.drawLine(self.px[j][0], self.py[j][0], self.px[j][1], self.py[j][1])
            if len(self.px[j])>=4:
             painter.drawLine(self.px[j][2], self.py[j][2], self.px[j][3], self.py[j][3])
             painter.drawText(self.px[j][3], self.py[j][3], str(j+1))
        for i in range(0,len(self.centerx)):
            painter.drawPoint(self.centerx[i], self.centery[i])
        painter.setPen(QPen(Qt.blue,2,Qt.SolidLine))
        for i in range(0,len(self.fx1)):
            painter.drawLine(self.fx0[i], self.fy0[i], self.fx1[i], self.fy1[i])
        painter.setPen(QPen(orange,2,Qt.SolidLine))
        for i in range(0,len(self.ilx1)):
            painter.drawLine(self.ilx0[i], self.ily0[i], self.ilx1[i], self.ily1[i])
        painter.setPen(QPen(pink,2,Qt.SolidLine))
        for i in range(0,len(self.isx1)):
            painter.drawLine(self.isx0[i], self.isy0[i], self.isx1[i], self.isy1[i])
        painter.setPen(QPen(Moccasin,2,Qt.SolidLine))
        for i in range(0,len(self.t1x1)):
            painter.drawLine(self.t1x0[i], self.t1y0[i], self.t1x1[i], self.t1y1[i])
        painter.setPen(QPen(Moccasin,2,Qt.SolidLine))
        for i in range(0,len(self.cox1)):
            painter.drawLine(self.cox0[i], self.coy0[i], self.cox1[i], self.coy1[i])
        painter.setPen(QPen(Moccasin,2,Qt.SolidLine))
        for i in range(0,len(self.sax1)):
            painter.drawLine(self.sax0[i], self.say0[i], self.sax1[i], self.say1[i])
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        for i in range(0,len(self.rsax1)):
            painter.drawLine(self.rsax0[i], self.rsay0[i], self.rsax1[i], self.rsay1[i])
        for i in range(0,len(self.l5x1)):
            painter.drawLine(self.l5x0[i], self.l5y0[i], self.l5x1[i], self.l5y1[i])
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        for i in range(0,len(self.cc1x1)):
            painter.drawLine(self.cc1x0[i], self.cc1y0[i], self.cc1x1[i], self.cc1y1[i])
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        for i in range(0,len(self.cc2x1)):
            painter.drawLine(self.cc2x0[i], self.cc2y0[i], self.cc2x1[i], self.cc2y1[i])
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        for i in range(0,len(self.cc3x1)):
            painter.drawLine(self.cc3x0[i], self.cc3y0[i], self.cc3x1[i], self.cc3y1[i])
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        for i in range(0,len(self.sc1x1)):
            painter.drawLine(self.sc1x0[i], self.sc1y0[i], self.sc1x1[i], self.sc1y1[i])
        painter.setPen(QPen(Qt.red,2,Qt.SolidLine))
        for i in range(0,len(self.sc2x1)):
            painter.drawLine(self.sc2x0[i], self.sc2y0[i], self.sc2x1[i], self.sc2y1[i])
        painter.setPen(QPen(Peru,4,Qt.SolidLine))
        for i in range(0,math.floor(len(self.cr)/2)):
            painter.drawPoint(self.cr[2*i], self.cr[2*i+1])
        painter.setPen(QPen(Maroon,4,Qt.SolidLine))
        for i in range(0,math.floor(len(self.tvmu)/2)):
            painter.drawPoint(self.tvmu[2*i], self.tvmu[2*i+1])
        painter.setPen(QPen(Maroon,4,Qt.SolidLine))
        for i in range(0,math.floor(len(self.tvml)/2)):
            painter.drawPoint(self.tvml[2*i], self.tvml[2*i+1])
        painter.setPen(QPen(Bisque,4,Qt.SolidLine))
        for i in range(0,math.floor(len(self.sg)/2)):
            painter.drawPoint(self.sg[2*i], self.sg[2*i+1])
        painter.setPen(QPen(Coral,2,Qt.SolidLine))
        for i in range(0,len(self.lbx1)):
            painter.drawLine(self.lbx0[i], self.lby0[i], self.lbx1[i], self.lby1[i])
        painter.setPen(QPen(Chartreuse,4,Qt.SolidLine))
        for i in range(0,math.floor(len(self.rs)/2)):
            painter.drawPoint(self.rs[2*i], self.rs[2*i+1])
        painter.setPen(QPen(Chartreuse,2,Qt.SolidLine))
        for i in range(0,len(self.tvsx1)):
            painter.drawLine(self.tvsx0[i], self.tvsy0[i], self.tvsx1[i], self.tvsy1[i])
        for i in range(0,math.floor(len(self.ls)/2)):
            painter.drawPoint(self.ls[2*i], self.ls[2*i+1])
        painter.setPen(QPen(Chartreuse,2,Qt.SolidLine))
        for i in range(0,len(self.uvsx1)):
            painter.drawLine(self.uvsx0[i], self.uvsy0[i], self.uvsx1[i], self.uvsy1[i])
        for i in range(0,math.floor(len(self.ls2)/2)):
            painter.drawPoint(self.ls2[2*i], self.ls2[2*i+1])
        painter.setPen(QPen(Chartreuse,2,Qt.SolidLine))
        for i in range(0,len(self.lvsx1)):
            painter.drawLine(self.lvsx0[i], self.lvsy0[i], self.lvsx1[i], self.lvsy1[i])
        for i in range(0,math.floor(len(self.ls3)/2)):
            painter.drawPoint(self.ls3[2*i], self.ls3[2*i+1])
        painter.setPen(QPen(Chartreuse,2,Qt.SolidLine))
        for i in range(0,math.floor(len(self.morphx)/2)):
            painter.drawLine(self.morphx[2*i], self.morphy[2*i], self.morphx[2*i+1], self.morphy[2*i+1])
        painter.setPen(QPen(Qt.green,4,Qt.SolidLine))    
        for k in range(0, len(self.curvex)): 
            painter.drawPoint(self.curvex[k], self.curvey[k])
        painter.setPen(QPen(Qt.red,4,Qt.SolidLine))    
        for m in range(0, len(self.inflection)):
            painter.drawPoint(self.curvex[self.inflection[m]], self.curvey[self.inflection[m]])
        #painter.setPen(QPen(Qt.blue, 4, Qt.SolidLine))
        #for i in range(0,len(self.vert_set)):
           #for j in range(0,len(self.vert_set[i])):
              #painter.drawPoint(round(self.vert_set[i][j][0]), round(self.vert_set[i][j][1]))
        #if len(self.vert_set)>=1:
            #painter.setPen(QPen(Qt.red, 4, Qt.SolidLine))
            #painter.setBrush(QBrush(Qt.red, Qt.SolidPattern))
            # points=[]
            #for j in range(0, len(self.vert_set[0])):
            # points.append(QPoint(round(self.vert_set[0][j][0]), round(self.vert_set[0][j][1])))
            #polygon = QPolygon(points)
            #painter.drawPolygon(polygon)
            # painter.drawText(self.list_center[0][0] + 2 * (
            # max(self.box_set[0][0][0], self.box_set[0][1][0], self.box_set[0][2][0],
            # self.box_set[0][3][0]) - min(self.box_set[0][0][0], self.box_set[0][1][0],
            #   self.box_set[0][2][0], self.box_set[0][3][0])),
        #  self.list_center[0][1], str(1))
        #if len(self.vert_set)>=13:
            #  painter.setPen(QPen(orange, 4, Qt.SolidLine))
            #  painter.setBrush(QBrush(orange, Qt.SolidPattern))
            #  for i in range(1,13):
            #  points=[]
            #  for j in range(0,len(self.vert_set[i])):
            #  points.append(QPoint(round(self.vert_set[i][j][0]), round(self.vert_set[i][j][1])))
            #  polygon=QPolygon(points)
            #   painter.drawPolygon(polygon)
            #   painter.drawText(self.list_center[i][0] + 2 * (
            #  max(self.box_set[i][0][0], self.box_set[i][1][0], self.box_set[i][2][0],
            # self.box_set[i][3][0]) - min(self.box_set[i][0][0], self.box_set[i][1][0],
                                                          #self.box_set[i][2][0], self.box_set[i][3][0])),
            # self.list_center[i][1], str(i + 1))
        #if len(self.vert_set)>=14:
            # painter.setPen(QPen(Qt.yellow, 4, Qt.SolidLine))
            # painter.setBrush(QBrush(Qt.yellow, Qt.SolidPattern))
            # for i in range(13,len(self.vert_set)):
            # points=[]
            # for j in range(0,len(self.vert_set[i])):
            #  points.append(QPoint(round(self.vert_set[i][j][0]), round(self.vert_set[i][j][1])))
            # polygon=QPolygon(points)
            #  painter.drawPolygon(polygon)
            #  painter.drawText(self.list_center[i][0] + 2 * (
            # max(self.box_set[i][0][0], self.box_set[i][1][0], self.box_set[i][2][0],
                             #self.box_set[i][3][0]) - min(self.box_set[i][0][0], self.box_set[i][1][0],
            # self.box_set[i][2][0], self.box_set[i][3][0])),
            # self.list_center[i][1], str(i + 1))
        painter.setPen(QPen(Qt.green, 2, Qt.SolidLine))
        for k in range(0, len(self.box_set)):
            painter.drawLine(self.box_set[k][0][0], self.box_set[k][0][1], self.box_set[k][1][0],
                             self.box_set[k][1][1])
            painter.drawLine(self.box_set[k][1][0], self.box_set[k][1][1], self.box_set[k][2][0],
                             self.box_set[k][2][1])
            painter.drawLine(self.box_set[k][2][0], self.box_set[k][2][1], self.box_set[k][3][0],
                             self.box_set[k][3][1])
            painter.drawLine(self.box_set[k][3][0], self.box_set[k][3][1], self.box_set[k][0][0],
                             self.box_set[k][0][1])
            painter.drawText(self.list_center[k][0] - 2*(
                        max(self.box_set[k][0][0], self.box_set[k][1][0], self.box_set[k][2][0],
                            self.box_set[k][3][0]) - min(self.box_set[k][0][0], self.box_set[k][1][0],
                                                         self.box_set[k][2][0], self.box_set[k][3][0])),
                             self.list_center[k][1], str(k + 1))
        for k in range(0,len(self.list_center)):
            painter.drawPoint(self.list_center[k][0], self.list_center[k][1])
        for k in range(0,len(self.spacex)):
            painter.drawPoint(self.spacex[k], self.spacey[k])
        for m in range(0, len(self.inflection_auto)-1):
            colorindex=Xcopy.deepcopy(m)
            if colorindex>len(color_list)-1:
                colorindex=0
            painter.setPen(QPen(color_list[colorindex], 4, Qt.SolidLine))
            for k in range(self.inflection_auto[m],self.inflection_auto[m+1]):
                painter.drawPoint(self.spacex[k], self.spacey[k])
        painter.setPen(QPen(Qt.blue, 4, Qt.SolidLine))
        for m in range(0, len(self.inflection_auto)):
            painter.drawPoint(self.spacex[self.inflection_auto[m]], self.spacey[self.inflection_auto[m]])
        painter.setPen(QPen(Qt.green, 4, Qt.SolidLine))
        for i in range(0,math.floor(len(self.addmaskx)/2)):
            painter.drawLine(self.addmaskx[2*i], self.addmasky[2*i], self.addmaskx[2*i+1], self.addmasky[2*i+1])
        painter.setPen(QPen(Qt.red, 2, Qt.SolidLine))
        for i in range(0,len(self.mid)):
            painter.drawLine(self.mid[i][0], self.mid[i][1], self.mid[i][2], self.mid[i][3])
        painter.setPen(QPen(Qt.blue, 2, Qt.SolidLine))
        for i in range(0, len(self.End)):
            painter.drawLine(self.End[i][0], self.End[i][1], self.End[i][2], self.End[i][3])
        #painter.drawLine(self.x0, self.y0, self.x1+100, self.y1)
#        pqscreen  = QGuiApplication.primaryScreen()
#        pixmap2 = pqscreen.grabWindow(self.winId(), self.x0, self.y0, abs(self.x1-self.x0), abs(self.y1-self.y0))
#        pixmap2.save('555.png')


#class CloudyPoints(QWidget, CloudyPoints):
#    def __init__(self):
#        super(CloudyPoints, self).__init__()
#        # 子窗口初始化时实现子窗口布局
#        self.setupUi()

# class plot_map(MainForm,plot_map):
#     def __init__(self):
#         super(plot_map, self).__init__()
#         # 子窗口初始化时实现子窗口布局
#         self.plot_contour_map(self)
class CDataPreprocessing(QWidget, CDataPreprocessing):
    def __init__(self): 
        super(CDataPreprocessing, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)

class CResultShowFirst(QWidget, CResultShowFirst):
    def __init__(self):
        super(CResultShowFirst, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)

class CResultShowSecond(QWidget, CResultShowSecond):
    def __init__(self):
        super(CResultShowSecond, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)

class CResultShowThird(QWidget, CResultShowThird):
    def __init__(self):
        super(CResultShowThird, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)
class CCoronal_c_bend(QWidget, CCoronal_c_bend):
    def __init__(self):
        super(CCoronal_c_bend, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)

class CResult_dynamic(QWidget, CResult_dynamic):
    def __init__(self):
        super(CResult_dynamic, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)


class CFootPressAnalyze(QWidget, CFootPressAnalyze):
    def __init__(self):
        super(CFootPressAnalyze, self).__init__()
        # 子窗口初始化时实现子窗口布局
        self.setupUi(self)




if __name__ == "__main__":
    
    app = QApplication(sys.argv)
    myshow = MainForm()

    # global g_strDataFolderNameOutHospital
    myshow.show()
    # g_strDataFolderNameOutHospital = '数据邹可'
    # g_strDataFolderNameOutHospital = myshow.HistoryDataList[0]
    # plot_contour_map(myshow,g_strDataFolderNameOutHospital)  # 画点云图
    # plot_Coronal_plane_with_background_plate(myshow)

    sys.exit(app.exec_())
#    app =QApplication(sys.argv)
#    ui = MainForm()
#    ui.show()
#    sys.exit(app.exec_())
